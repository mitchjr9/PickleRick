"""
Pickle Rick — Your Pickleball Rules Assistant
==============================================
Single-file Streamlit app — pickleball rules AI assistant trained on the
2026 USA Pickleball Official Rulebook + 2026 Rulebook Change Document.

Tabs: 🥒 Home/Chat | 🎬 Video Analyzer | 📝 Quiz
Run:  streamlit run app.py
"""

# ── Standard library ──────────────────────────────────────────────────────────
import base64
import datetime
import json
import os
import subprocess
import sys
import tempfile
import urllib.parse
from io import BytesIO
from pathlib import Path

# ── Third-party ───────────────────────────────────────────────────────────────
import anthropic
import streamlit as st

# ── OpenCV — auto-install if missing ─────────────────────────────────────────
try:
    import cv2
    OPENCV_AVAILABLE = True
except ImportError:
    try:
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "opencv-python-headless", "-q"]
        )
        import cv2
        OPENCV_AVAILABLE = True
    except Exception:
        OPENCV_AVAILABLE = False

# ── fpdf2 — auto-install if missing ──────────────────────────────────────────
try:
    from fpdf import FPDF  # noqa: F401
    FPDF2_AVAILABLE = True
except ImportError:
    try:
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "fpdf2", "-q"]
        )
        from fpdf import FPDF  # noqa: F401
        FPDF2_AVAILABLE = True
    except Exception:
        FPDF2_AVAILABLE = False

# ── python-docx — auto-install if missing ────────────────────────────────────
try:
    from docx import Document  # noqa: F401
    DOCX_AVAILABLE = True
except ImportError:
    try:
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "python-docx", "-q"]
        )
        from docx import Document  # noqa: F401
        DOCX_AVAILABLE = True
    except Exception:
        DOCX_AVAILABLE = False


# =============================================================================
# CORE KNOWLEDGE BASE — 2026 USA Pickleball Official Rulebook + Change Document
# =============================================================================
# This is the permanent brain of Pickle Rick. Every rule citation here comes
# directly from the two project PDFs. Update this string when new rule changes
# or personal notes are added.

CORE_KNOWLEDGE = """
# Pickle Rick Core Knowledge Base
## 2026 USA Pickleball Official Rulebook + 2026 Change Document

> **CRITICAL INSTRUCTION FOR ALL RESPONSES:**
> The 2026 USA Pickleball Official Rulebook (effective January 1, 2026) is the
> authoritative source. The 2026 Rulebook Change Document explains every rule
> that changed from 2025 → 2026. Always cite the specific 2026 rule number
> (e.g., "Rule 7.C.1") and reference the 2025→2026 change when relevant.
> If a player asks about something using 2025 rule numbers (e.g., "Rule 4.A.7"),
> map it to the new 2026 rule number using Section 0 below.

---

## 0. 2025 → 2026 RULE CHANGE QUICK MAP

| # | 2025 Rule | 2026 Rule | Topic |
|---|-----------|-----------|-------|
| 1 | 4.M.4 | **7.E.1** | Serve must land in correct service court (clarification) |
| 2 | 11.K | **10.C.5** | Net post — ball that bounces in then hits net post = winner for hitter |
| 3 | 12.C.4.a | **15.B.4.a** | Round robin withdrawals/retirements/forfeits |
| 4 | 9.B.1 | Definition | "Act of Volleying" definition added |
| 5 | 12.C.4 | **15.B.4** | Round robin tiebreaker process |
| 6 | 13.C.2.a | **8.J** | Spectators must NOT be consulted on calls |
| 7 | 3.A.5 | Deleted | "Cross-court" definition removed |
| 8 | 13.M.2 | **22.L, 22.L.2** | Ejections for assault |
| 9 | 13.M | **22.L, 22.L.5** | Ejections for willful damage to property |
| 10 | 13.G.3.e | **22.A** | Penalties before match starts |
| 11 | 11.A | **10.D, 10.D.1, 10.D.2** | Double-hit rule clarified |
| 12 | 13.C.4.b | **17.C.2** | Pre-match briefing |
| 13 | 6.C.8 | **8.H** | "Doubt" → "Conflict" on partner disagreement |
| 14 | 12.B.1.c | **4.B** | Game potentially won by a technical foul |
| 15 | 4.A.5 | **7.B.2** | Paddle adding spin to ball at contact during serve — **NOW LEGAL by paddle** |
| 16 | 6.C.7 | **8.F, 8.F.3** | Prompt line calls |
| 17 | 10.B.2.C | **21.C.9** | Doubles player retirement, partner continues |
| 18 | 13.J | **20.J, 20.J.1** | Rescinding a Head Referee call |
| 19 | 10.B | **21.C.4** | Rescinding medical time-out |
| 20 | 12.B.1.c | **4.B, 14.A.2** | Rally scoring — winning point: **receiver CAN now win on game point** |
| 21 | — | **25.A** | NEW: Wheelchair player rules (extensive new section) |
| 22 | — | **25.A.4** | NEW: Wheelchair is an extension of the body |
| 23 | 4.A.7 | **7.C** | "Flat serve" — "clearly" added — stronger enforcement |
| 24 | 4.K | **6.F** | Wrong score called — process clarified |
| 25 | — | **25.B** | NEW: Adaptive standing player rules |
| 26 | 7.N | **24.B** | Extra ball visible/dropped = fault (stricter) |
| 27 | — | **8.J** | Spectators may NOT be consulted on calls |
| 28 | — | **19.G** | Officiated line-call procedure clarified |
| 29 | 10.A.3 | **21.A.2** | Time-out must be called audibly/visibly |

> **HEADLINE 2026 CHANGES TO REMEMBER:**
> 1. **Rally scoring receiver CAN win on game-point** (Rule 4.B, 14.A.2) — change #20
> 2. **Spin allowed on paddle contact during serve** — but NEVER on hand release (Rule 7.B.2) — change #15
> 3. **Net post: ball that bounces in then hits net post = WINNER for hitter** (Rule 10.C.5) — change #2
> 4. **Partner "doubt" is now "conflict"** — disagreement = ball is "in" (Rule 8.H) — change #13
> 5. **Volley serve: "clearly" added** to upward arc, paddle-head, ball-height (Rule 7.C) — change #23
> 6. **Time-outs MUST be audibly/visibly signaled** (Rule 21.A.2) — change #29
> 7. **Extra ball visible or dropped = fault** (Rule 24.B.1) — change #26
> 8. **Spectators must NOT be consulted on calls** (Rule 8.J) — change #6
> 9. **NEW Wheelchair (25.A), Adaptive Standing (25.B), Hybrid Doubles (25.C) sections**
> 10. **Ejection/expulsion** for paddle/ball assault or willful property damage (Rule 22.L) — changes #8, #9

---

## 1. THE GAME — Section 1 (p. 1)

Pickleball is played on a 20 ft × 44 ft court with a perforated ball and tennis-style net.
- **Standard scoring:** Points scored only by serving team. First to 11, win by 2.
- **Rally scoring (provisional, Rule 14.A / 15.C.2):** Point on every rally.
  - **2026 change:** Receiver CAN now win on game point — point awarded on every rally regardless of who serves (Rule 4.B, 14.A.2; Change #20).
- **Two-Bounce Rule:** Receiver must let serve bounce. Then serving team must let return bounce. (Rule 10.A)
- **Non-Volley Zone (NVZ):** 7 ft × 20 ft area on each side of net where you cannot volley. (Rule 3.A.4.c)

---

## 2. KEY DEFINITIONS — Section 2 (pp. 2–4)

- **Fault:** A rules violation resulting in dead ball + loss of rally.
- **Hinder:** Transient element NOT caused by a player (stray balls, insects, foreign material). Permanent objects are NOT hinders.
- **Distraction (player-caused):** Physical actions not common to the game that interfere with opponent's ability to hit the ball. Examples: loud noises, stomping, erratic paddle waving. (10.F, 20.I)
- **Volley:** A strike of the ball before the ball bounces.
- **Volleying, Act of:** Begins when the ball is hit out of the air (i.e., volleyed) and ends when the player's follow-through momentum stops. (NEW definition — Change #4)
- **Momentum:** Property of a body in motion that causes a player to continue moving after contacting the ball. Ends when player regains balance/control or stops moving toward NVZ.
- **Permanent Object:** Anything on/above/near court that can interfere — ceilings, walls, fencing, lighting, **net posts (including connected wheels, arms, legs)**, net cable/rope on top of net post, stands, seats, spectators in recognized positions, referee, line judges.
- **Imaginary Extension:** Continuation of a line beyond its physical endpoint (used for serving area boundaries, etc.).
- **Live Ball:** From start of score call until ball becomes dead.
- **Technical Warning:** Punitive warning for minor unsportsmanlike behavior.
- **Technical Foul:** One-point penalty for extreme unsportsmanlike behavior.
- **Verbal Warning:** Non-punitive caution (NEW classification).

---

## 3. COURT & EQUIPMENT — Section 3 (pp. 5–10)

**Court (Rule 3.A):** 20 ft × 44 ft (6.10 m × 13.41 m).
- **NVZ (Rule 3.A.4.c):** 7 ft × 20 ft. **All NVZ lines ARE PART OF THE NVZ.** The NVZ is **2-dimensional and does NOT extend above the playing surface.** This is critical — you can lean over the NVZ without faulting if you don't touch it.
- **Service courts (3.A.4.f):** Each service court is bounded by and **includes** its baseline, sideline, and centerline.
- **Serving area (3.A.4.g):** Behind baseline, bounded by **imaginary extensions** of sideline and centerline.
- **Lines (3.A.4.e):** 2 inches wide, contrasting color.

**Net (Rule 3.B):** 36" at sidelines, 34" at center.

**Ball (Rule 3.C):** Approved ball list. Indoor balls have larger holes; outdoor balls have smaller holes.

**Paddle (Rule 3.D):**
- 3.D.2 — Combined length+width ≤ 24"; length ≤ 17". No thickness restriction.
- 3.D.3 — No weight restriction.
- 3.D.1 — Must have brand/model marking + "USA Pickleball Approved" seal/text.
- Must be on USA Pickleball Approved Paddle List for sanctioned play.
- 18.A.2 — **Match forfeit** if non-compliant paddle discovered DURING play. No penalty to switch BEFORE match (18.A.1).

---

## 4. SCORING & WINNING THE GAME — Section 4 (pp. 10–11)

- **4.A — Standard scoring:** Point only when serving team wins rally.
- **4.B — Winning the Game (2026 NEW WORDING):** "The first singles player or doubles team to score the winning point wins the game." — Change #14, #20.
  - **Implication:** In rally scoring, a receiving team CAN win the game by winning a rally on game point. Also: if the serving team is at 0 and assessed a technical foul when receiver is at game point, the receiver gets the point AND the game.

---

## 5. PLAYER POSITIONS & SERVING SEQUENCE — Section 5 (pp. 12–13)

**Singles (5.A):**
- Score 0 or even → serve from RIGHT, into opponent's right court.
- Score odd → serve from LEFT, into opponent's left court.

**Doubles (5.B):**
- Both players serve before side-out, **except** at start of game (only starting server serves before first side-out).
- 5.B.2 — After side-out, the player on the correct side per team's score becomes the **first server**; partner is **second server**. The first-game starting server is designated as a "second server" for the team's first service rotation.
- 5.B.3 — Starting server on RIGHT when team score is 0/even; LEFT when odd. Partner is opposite.
- 5.B.4 — Except while serving/receiving, **no restriction on player positions during rallies**.

**5.C — Player/Position Errors:**
- 5.C.1.a — Replay if correct claim during rally.
- 5.C.1.b — **Fault if INCORRECT claim** that stopped the rally.
- 5.C.1.c — **Fault on incorrect receiver** even if rally was completed (must be called before next serve).
- 5.C.2 — If positions wrong but rally completed and no one stopped it, result stands.

---

## 6. PLAYER READINESS & SCORE CALLING — Section 6 (pp. 14–15)

- 6.D — Server has 10 seconds to serve after score call.
- 6.D.2 — If serving team changes serving areas after score is called, replay/repostion + re-call score (restarts 10-sec count).
- 6.E — "Stop" or "wait" before serve is hit will be recognized.
- **6.F — Challenging the Score Call (CHANGE #24):**
  - 6.F.1 — Replay if incorrect score caught BEFORE return of serve and ball still live.
  - 6.F.2 — **Fault** if a player stops a rally to challenge a CORRECT score.
  - 6.F.3 — **Fault** if challenge made AFTER return of serve.
  - 6.F.4 — If incorrect score discovered after rally completes, rally result stands; correct score before next serve.

---

## 7. SERVING — Section 7 (pp. 16–18) ★ HIGHLY TESTED ★

**7.A — Server Positioning (when serve is hit):**
- 7.A.1 — At least one foot must contact the correct serving area. (Fault 7.A.1.a — Server Not Grounded)
- 7.A.2 — Neither foot may contact the court. (Fault 7.A.2.a — Server Contacting Court)
- 7.A.3 — Neither foot may contact playing surface OUTSIDE the correct serving area. (Fault 7.A.3.a — Server Outside Serving Area) — bounded by imaginary extensions.

**7.B — Ball Release:**
- 7.B.1 — Release ball using ONLY one hand OR only the paddle.
- 7.B.2 — **No manipulation/spin upon release** by any body part or paddle. Exception: ball may roll off paddle face by gravity.
  - **2026 CHANGE #15:** "Spin may be applied to the ball upon contact by the paddle." Spin via PADDLE at contact = LEGAL. Spin via HAND on release = ILLEGAL.
- 7.B.2.a — Replay if receiver determines manipulation/spin (must call before returning serve).
- 7.B.3 — Server's release must be VISIBLE to receiver.
- 7.B.3.a — Receiver may call replay before return if release was not visible.

**7.C — Volley Serve (CHANGE #23 — "clearly" added for stronger enforcement):**
- 7.C.1 — Paddle must be moving in a **clear** upward arc at contact.
- 7.C.2 — Highest point of paddle head must **clearly** NOT be above highest part of server's wrist joint at contact.
- 7.C.3 — Ball must **clearly** be no higher than server's waist at contact.
- 7.C.4 — Forehand or backhand allowed.
- 7.C.5 — Fault if 7.C.1, 7.C.2, or 7.C.3 violated.

**7.D — Drop Serve:**
- 7.D.1 — Release from natural (unaided) height.
- 7.D.2 — Ball must NOT be propelled in any direction or any manner before being struck.
- 7.D.3 — No restriction on # of bounces before strike.
- 7.D.4 — No restriction on bounce location.
- 7.D.5 — Forehand or backhand.
- 7.D.6 — Fault if 7.D.1 or 7.D.2 violated.
- **NOTE:** Drop serve has NO upward-arc / paddle-head / waist-height requirements.

**7.E — Serve Placement:**
- Must serve diagonally to opposite service court. Must clear NVZ. May or may not touch net.
- 7.E.1 — Fault: serve lands outside correct service court (CHANGE #1).
- 7.E.2 — Fault: serve lands in NVZ (NVZ line counts as IN the NVZ → fault).
- 7.E.3 — Fault: serve hits permanent object before landing.
- 7.E.4 — Fault: serve hits server or server's partner (or anything worn/carried).
- 7.E.5 — Fault: serve hits receiver or receiver's partner before landing → fault on RECEIVER (point to server).

---

## 8. LINE CALLS — Section 8 (pp. 19–20) ★ CRUCIAL FOR REC PLAY ★

- 8.A — Players responsible for "out" calls on THEIR end of the court. Either doubles partner may call.
- 8.B — "In" ball: served ball that clears NVZ line and lands in correct service court is "in." Returned ball that lands on opponent's end is "in."
- 8.C — "Out" ball: served ball not landing in correct service court (including landing ON the NVZ line) is "out." Any other ball outside court is "out."
- **8.D — Code of Ethics:** "All questionable calls must be resolved in favor of the opponent. The opponent gets the benefit of the doubt. **Any ball that cannot be promptly called 'out' must be considered 'in.'**"
- **8.E — Line Call Certainty:** "Players must NOT call a ball 'out' unless they can clearly see a SPACE between the line and the ball when it lands." (Figure 8-1)
- 8.F — "Out" must be promptly signaled audibly OR visibly OR both.
  - 8.F.1 — Any "out" call after the ball lands is a line call.
  - 8.F.2 — Line call → dead ball.
  - 8.F.3 — **Out-call timing (CHANGE #16):** If you RETURN the ball, the "out" call must be made before opponent hits ball or before ball becomes dead, otherwise play continues. If you DON'T return the ball, a prompt "out" call counts even if ball is already dead.
  - 8.F.4 — Partner communication while ball is in air ("out!" "no!" "bounce it!") = communication, NOT a line call.
- 8.G — Players may override own/partner's call to disadvantage themselves.
- **8.H — Partner Disagreement (CHANGE #13):** "When partners disagree on a line call, then conflict exists, and the team's call will be 'in.'" (was "doubt" pre-2026)
- 8.I — Asking opponent's opinion: opponent's clear in/out decision stands. If no definitive call, your call stands; if you made no call, ball is "in." After asking, you & partner LOSE the chance to call (except to override in opponent's favor).
- **8.J — Spectator Involvement (CHANGE #6):** "Spectators must not be consulted on any call."

---

## 9. DEAD BALLS, FAULTS & HINDERS — Section 9 (pp. 21–22)

- 9.A.1 — **Fault** if you stop a rally before ball becomes dead, UNLESS you're calling a hinder, correctly identifying position/server error, or correcting a wrong score before return.
- 9.B.1 — Faults only occur while ball is live, EXCEPT NVZ momentum violations (11.A.2) that occur after dead ball.
- 9.B.3 — **Players may ONLY call NVZ faults & service foot faults on opponents.**
  - 9.B.3.b — Disagreement between teams on a fault → REPLAY.
  - 9.B.3.c — Disagreement between partners → benefit of opponents.
- 9.B.4 — For other faults (besides NVZ + service foot fault), you may mention to opponent but cannot enforce — final call is on the alleged offender.
- 9.C.1 — Any player may call a hinder.
- 9.C.2 — Called hinder = REPLAY.

---

## 10. RALLY SITUATIONS — Section 10 (pp. 23–24)

- **10.A — Two-Bounce Rule:**
  - 10.A.1 — Fault: receiver volleys serve.
  - 10.A.2 — Fault: serving side volleys return of serve.
- 10.B.1 — Fault: ball not returned before second bounce.
- **10.C — Returned Ball Placement Faults:**
  - 10.C.1 — Lands out of bounds.
  - 10.C.2 — Lands on hitter's side (failed to cross net).
  - 10.C.3 — Hits player or anything worn/carried EXCEPT paddle/hand below wrist while holding paddle.
  - 10.C.4 — Hits permanent object before landing.
  - **10.C.5 — Hits permanent object AFTER landing on hitter's court (CHANGE #2).** Net post is a permanent object. **If a ball crosses net, bounces in opponent's court, and THEN (due to spin/wind) touches the net post → POINT for the hitter.** This is the "net post winner" rule.
- **10.D — Double Hit (CHANGE #11):**
  - 10.D.1 — Fault if ball is hit more than once unless the stroke is **continuous and in a single direction** by ONE player.
  - 10.D.2 — Fault if both partners strike the ball.
- 10.E — Missed shot (whiff) → ball remains live.
- 10.F — Distraction (player-caused) prohibited; see 20.I.1.
- 10.G — Damaged ball: play continues to end of rally. Replace if all agree (10.G.1 — replay if all agree it affected outcome; 10.G.2 — soft/degraded ball: prior rally stands).
- 10.H — Injury during rally: play continues to end of rally.
- 10.I — Equipment problem during rally: play continues.
- 10.J — Item dropped on YOUR side: ball remains in play (even if it hits the item) — UNLESS item contacts NVZ as a result of you volleying (then NVZ fault per 11.A.1).
- 10.K — Between rallies: quick hydration/towel/equipment OK if game flow not impacted.

---

## 11. NON-VOLLEY ZONE INFRACTIONS — Section 11 (p. 25) ★★★

**Allowable Contact (11.A):** All volleys must be **initiated outside** the NVZ. A player or anything in contact with the player may contact the NVZ AT ANY TIME except during the act of volleying.

**Faults:**
- **11.A.1 — NVZ Contact While Volleying.** When a volleying player (or anything in contact with that player, including their partner) touches the NVZ. Includes hat/sunglasses/paddle that fall in NVZ during a volley.
- **11.A.2 — NVZ Momentum.** When a volleying player's momentum causes them to contact anything (including their partner) that is in contact with the NVZ — **even AFTER the ball becomes dead.** (This is the famous "you can't fall in the kitchen" rule.)
- **11.A.3 — Failure to Exit NVZ Before Volleying.** After contacting the NVZ, you must establish BOTH FEET completely outside the NVZ before volleying.

**Critical points referees miss most often:**
1. NVZ is 2D — leaning over without contact is FINE.
2. You can stand in NVZ all day if you're not volleying.
3. Momentum fault counts even if you've already won the rally.
4. Dropping your paddle in NVZ during a volley = fault.
5. Partner pushing you into NVZ during your volley = your fault (or theirs if anything they wear/carry contacts NVZ during your volley).

---

## 12. THE PADDLE DURING PLAY — Section 12 (p. 26)

- 12.A.1 — Fault: more than one paddle during a rally.
- 12.B — Must have possession when hitting; one or both hands; can switch hands anytime.
- 12.B.1 — Fault: not in possession at contact.
- 12.C.1 — Fault: catch or carry ball on paddle ("scoop").

---

## 13. NET & NET SUPPORT SYSTEM — Section 13 (pp. 27–29)

- 13.A — Ball contacting net itself: ball remains in play. (13.A.1 — replay if ball gets caught in net or contacts a draping/deflecting net.)
- 13.B.1 — Fault: served ball contacts net support system on either end; or returned ball contacts net support system on hitter's end before crossing net.
- 13.B.2 — Replay: returned ball crosses net then contacts crossbar/support system within court boundaries.
- 13.C — Ball may be returned around outside of net post (Aussie shot — legal!).
- 13.D.1 — Fault: ball hit between net and net post.
- 13.E.1 — Fault: ball hit under net.
- **13.F.1 — Fault: ball hit before it entirely crosses plane of net to your end** (no reaching over).
- 13.G.1 — Fault: player or anything worn/carried contacts net or net support system while ball is live.
- 13.H.1 — Fault: contacting opponent or opponent's court while ball is live.
- **13.I — Plane of Net (Reaching Over):**
  - 13.I.1 — May cross plane IMMEDIATELY AFTER hitting (follow-through OK while executing strike).
  - 13.I.1.a — Fault: cross plane BEFORE hitting.
  - 13.I.1.b — Fault: cross plane and DON'T hit.
  - 13.I.2 — Spin-back ball (e.g., severe backspin causes ball to cross net back to opponent's side without being touched): receiver may go around/over/under net to play it — but only AFTER the ball has crossed back.
    - 13.I.2.a — Fault: crossing plane BEFORE ball crosses back.
    - 13.I.2.b — Fault: failing to hit it before second bounce on opponent's end.
- 13.J — Replay if net system malfunctions during rally.

---

## 14. RALLY SCORING (Provisional) — Section 14 (pp. 30–34)

- 14.A.2 — **Point scored by player/team that wins each rally** (CHANGE #20).
- 14.B — Mini-singles: half-court singles variation. Server's side determines diagonal court; receiver's score determines their side.
- See also 15.C.2 for tournament rally scoring formats.

---

## 15. TOURNAMENT MATCH/SCORING OPTIONS — Section 15 (pp. 35)

- 15.C.1 — Standard scoring formats: 2-of-3 to 11; 3-of-5 to 11; one game to 15 or 21; round robin one-game to 11 (6+ teams).
- 15.C.2 — **Provisional rally scoring formats** — TD's option, NOT permitted in double-elim doubles, 2026 USAP Golden Ticket events, or 2026 USAP National Championship events.
- 15.F — Two-match minimum per event entered (except single-elim w/o consolation).

---

## 16. TOURNAMENT DIRECTOR — Section 16 (pp. 37–38)

- 16.F — TD **must NOT** implement any rule not in the Rulebook. Exceptions for physical court limitations require pre-approval by USA Pickleball Managing Director of Officiating.
- 16.G — TD chooses the tournament ball from approved list.
- 16.K — TD may require apparel changes.

---

## 17. REFEREE DUTIES — Section 17 (pp. 39–40)

- 17.D.2 — Call score after confirming correct server.
- 17.D.4 — Call "second server" after first server's team loses rally.
- 17.D.5 — Announce "side out" when singles player or second server loses serve.
- 17.D.8 — Call faults when they occur.
- 17.D.10 — Assist with line calls upon appeal or when line judge signals blocked view.
- 17.D.13 — Call hinders / determine validity of player-called hinders.
- 17.D.16 — 15-second warnings during time-outs.

---

## 18. PLAYER PRE-MATCH — Section 18 (p. 42)

- 18.A.1 — Non-compliant paddle BEFORE match → switch, no penalty.
- 18.A.2 — Non-compliant paddle DURING match → **match forfeit**.
- 18.A.3 — Discovered AFTER match → results stand.
- 18.B.3 — Apparel approximating ball color must be changed (referee time-out used).
- 18.C — Starting servers must wear ID determined by TD (visible during play).

---

## 19. TOURNAMENT LINE CALLS — Section 19 (pp. 44–46)

- 19.A.1 — Referee responsible for service foot faults, short serves, NVZ faults.
- 19.B — No line judges → players make all "out" calls on their end.
- **19.C — With line judges:**
  - 19.C.1 — Line judges call their assigned line + assist as requested ("OUT" + outstretched arm).
  - 19.C.2 — Players responsible only for centerline of service court on their end.
  - 19.C.3 — Player line calls invalid except centerline + override to favor opponent.
- 19.D — Only RALLY-ENDING shots can be appealed. Appeal before next serve hit (or before scoresheet initialed for match-ending shot).
- 19.F — Referee may overrule line judge's "out" call (19.F.2 → replay).
- 19.G — All officials uncertain → referee may consult; if all unable → replay.

---

## 20. TOURNAMENT MATCH SITUATIONS — Section 20 (pp. 47–51)

- 20.A — Any player may ask referee for correct score before serve hit.
- 20.B — Any player may ask referee to confirm correct server/receiver/positions before serve hit.
- 20.D — After score is called, if serving team changes serving areas → referee stops play, allows reposition, re-calls score.
- **20.E — Service Faults (referee-called):**
  - 20.E.1.a — Volley serve: no upward arc.
  - 20.E.1.b — Volley serve: paddle head clearly above wrist.
  - 20.E.1.c — Volley serve: ball clearly above waist.
  - 20.E.1.d — Volley/drop serve: not clearly released from one hand or paddle only.
  - 20.E.1.e — Volley/drop serve: spin or manipulation on release.
  - 20.E.1.f & 20.E.1.g — additional service faults.
  - 20.E.2 — Replays when referee uncertain whether serve requirements met.
- **20.H — Hinder (CHANGE: referee must validate):**
  - 20.H.1 — Fault if invalid (referee doesn't validate).
  - 20.H.2 — Replay if valid.
- **20.I — Distraction:**
  - 20.I.1 — Fault if referee judges player distracted opponent about to hit ball.
- **20.J — Officiating Decision Challenge (CHANGE #18):**
  - 20.J.1 — Rescinding challenge after referee acknowledged → standard time-out charged (or technical foul if none available).
  - 20.J.2 — Challenge UPHELD → technical warning + standard time-out (or tech foul if no TO).
  - 20.J.3 — Challenge OVERTURNED → reverse ruling or order replay.
- 20.K — Referee may remove a line judge for cause.

---

## 21. TIME-OUTS & BREAKS — Section 21 (pp. 52–55)

**21.A — Standard Time-Outs:**
- 21.A.4 — Two 1-minute time-outs per game in 11-point format; three in 15- or 21-point.
- **21.A.2 — Requesting Time-Out (CHANGE #29):** Must be called audibly by voice OR visibly by hand, OR both, directed toward opponents and referee.
  - 21.A.2.a — Fault: time-out called AFTER serve is hit.
  - 21.A.2.b — Verbal/technical warning if not called audibly/visibly (delay of game).

**21.B — End Changes:**
- 21.B.1 — In 11-point game: at score of 6.
- 21.B.2 — In 15-point game: at score of 8.
- 21.B.3 — In 21-point game: at score of 11.
- 21.B.4 — Same server continues after end change.
- 21.B.5 — Max 1 minute.
- 21.B.7 — If end change missed, execute when detected (no fault, score unchanged, same server).

**21.C — Medical Time-Out (CHANGE #19):**
- 21.C.1 — Cannot be called before match starts.
- 21.C.3 — One per match.
- 21.C.4 — **Rescinding after medical called and BEFORE arrival → standard TO charged; medical TO not charged.** No standard TO available → tech foul (delay of game).
- 21.C.6 — Max 15 minutes (off-court treatment travel time excluded).
- 21.C.7 — Timer starts when medical assistance arrives.
- 21.C.9 — Match retirement if can't continue after 15 min (CHANGE #17 — partner may continue in doubles per rules).
- 21.C.10 — Invalid medical condition → technical warning + standard TO charged (or tech foul if none).
- 21.C.11 — Resume with 15-second warning.

**21.D — Referee Time-Out:**
- 21.D.1 — Referee can call for safety / potential medical.
- 21.D.2 — **Bleeding/blood**: stop at end of rally, must control bleed and clean before resuming.
- 21.D.3 — Foreign substances must be cleaned.

**21.E — Equipment Time-Out:** Reasonable duration if needed for fair/safe continuation.

---

## 22. PLAYER CONDUCT, WARNINGS & PENALTIES — Section 22 (pp. 56–60)

**22.A — Verbal Warnings:** Non-punitive caution for minor unsportsmanlike behavior. No effect on rally/score/server.

**22.B — Technical Warnings (non-exhaustive):**
- 22.B.6 — Excessive line-call appeals.
- 22.B.7 — Coaching during rally (also see 20.G.1).
- 22.B.8 — **Lost ruling challenge (TO available)** → tech warning + standard TO. (See 20.J.2.)
- 22.B.9 — **Invalid medical TO request (TO available)** → tech warning + standard TO. (See 21.C.10.)
- 22.B.10 — Other minor unsportsmanlike behavior.

**22.C — Verbal/Technical Warning Consequences:**
- 22.C.1 — No loss of rally or point.
- 22.C.2 — No effect on server.

**22.D — Technical Fouls (extreme unsportsmanlike conduct):**
- 22.D.1 — **Paddle abuse** (negligent/reckless throw not striking person or property).
- 22.D.2 — **Ball abuse** (aggressively/recklessly throwing/hitting dead ball).
- 22.D.3 — Extreme objectionable language/profanity.
- 22.D.4 — Threats.
- 22.D.5 — Lost ruling challenge — no TO available.
- 22.D.6 — Invalid medical TO — no TO available.
- 22.D.7 — Second technical warning during a match.
- 22.D.8 — Other extreme unsportsmanlike behavior.

**22.E — Tech Foul Consequences:**
- **22.E.1 — Score adjustment: −1 from offender (or +1 to opponent if offender at 0).** Player/team must move to correct position based on new score. **Game-winning point can result if adjustment makes the score game-winning, regardless of who was serving.**
- 22.E.2 — No effect on server (no side-out from tech foul itself).

**22.F — Game Forfeit:**
- 22.F.1 — Combination of one tech warning + one tech foul (or three tech warnings) during match → game forfeit.
- 22.F.2 — Late reporting for first game (Rule 18.F.1).

**22.I — Match Forfeit by Referee:**
- 22.I.1 — Dangerous paddle/ball abuse striking person or damaging property.
- 22.I.5 — Non-compliant paddle (Rule 18.A.2).
- 22.I.6 — Improper deliberately aggressive physical contact.
- 22.I.7 — Second tech foul, OR tech warning/foul after a game forfeit from prior tech accumulation.

**22.J — Match Forfeit by TD:** Venue rules, improper between-match conduct, facility abuse, improper apparel, other rules.

**22.K — Forfeit Scoring:** All forfeit games reported as 11-0, 15-0, or 21-0.

**22.L — Ejection or Expulsion (CHANGES #8, #9):**
- 22.L.2 — **Injurious paddle/ball abuse** or other physical violence — ejection from tournament.
- 22.L.5 — **Willful damage** to venue property — ejection.

---

## 23. NON-OFFICIATED & UNOFFICIATED PLAY

- 24.A — Non-officiated tournament play uses Sections 1–14 + Section 24.
- **24.B — Extra Ball (CHANGE #26):**
  - Must not be visible to opponent + must remain in player's possession.
  - 24.B.1 — Fault if extra ball visible to opponent OR falls to playing surface while ball is live.

---

## 24. WHEELCHAIR / ADAPTIVE STANDING / HYBRID — Section 25 (pp. 65–71) ★ NEW IN 2026 ★

**25.A — Wheelchair Play (Changes #21, #22):**
- 25.A.4 — **Wheelchair is part of player's body. Large rear wheels = legs for positioning.**
- 25.A.5 — Player must contact seat (with allowed exceptions per 25.A.5.a / 25.A.7).
- 25.A.8 — Server must be grounded with at least one rear wheel in serving area; rear wheels must NOT contact court at serve; rear wheels must NOT be outside serving area at serve.
- 25.A.9 — **Two-bounce allowance**: wheelchair player may let ball bounce TWICE; second bounce can be anywhere on playing surface. Fault if not returned before third bounce.
- 25.A.10 — NVZ rules: front (smaller) wheels and rear stabilizing wheels MAY contact NVZ at any time. Large rear wheels may not contact NVZ during volley.
- 25.A.12.a — Transition standing↔wheelchair allowed once per match (treated as equipment time-out).

**25.B — Adaptive Standing (Change #25):**
- 25.B.3 — Eligibility-based two-bounce allowance for above-knee amputees, significant mobility impairments, neurological conditions (CP, stroke), users of crutches/braces. Must be declared before match.

**25.C — Hybrid Doubles (Change #25):**
- 25.C.1 — Each team = one wheelchair player + one standing player.
- 25.C.3 — Only eligible adaptive standing players may use two-bounce allowance in hybrid play.

---

## 25. COMMON PITFALLS & PROVEN-CORRECT CALLS

> **Pickle Rick's Top Ref/Player Pitfalls — Memorize These:**

1. **"My foot was on the NVZ line during my volley."** → FAULT (11.A.1). NVZ lines are part of the NVZ.
2. **"I volleyed; my momentum carried me into the kitchen after the ball was dead."** → FAULT (11.A.2). Momentum doesn't stop until you regain balance.
3. **"My partner was standing in NVZ during my volley with no contact between us."** → No fault. Partner's NVZ presence only matters if partner was IN CONTACT with you during the volley.
4. **"My hat fell off into the NVZ during a volley."** → FAULT (11.A.1) — anything in contact with the volleying player.
5. **"I called 'out' but the ball was clearly in."** → If you returned it, opponent gets the point (you stopped the rally; even if you returned, the call ends play). If you didn't return, your "out" call is the call (subject to opponent override per 8.G).
6. **"My partner and I disagreed on a line call."** → Ball is "in." (8.H — CHANGE #13 — "conflict" → in.)
7. **"I asked the spectator if it was in."** → 8.J — Spectators MUST NOT be consulted. Repeated → verbal/tech warning per 22.B.10.
8. **"Ball bounced in their court, then spin took it back to hit the net post."** → POINT for hitter (10.C.5 — Net Post Winner — CHANGE #2).
9. **"I called time-out silently with a finger."** → Tech warning per 21.A.2.b. Must be audible OR visible hand signal.
10. **"During my serve I added spin with my paddle at contact."** → LEGAL (7.B.2 — CHANGE #15). Spin via the paddle at contact is now allowed; spin via hand on release is still illegal.
11. **"Receiver was at game point in rally scoring; they won the rally on my serve."** → Receiver wins game (4.B, 14.A.2 — CHANGE #20).
12. **"I dropped my second ball during a rally."** → FAULT (24.B.1 — CHANGE #26).
13. **"My paddle face went over the net during my follow-through."** → LEGAL (13.I.1) IF you already hit the ball; ILLEGAL (13.I.1.a) if you crossed before contact.
14. **"Around-the-post (ATP) shot."** → LEGAL (13.C). Ball may go around outside the post; doesn't have to clear net height on the return path.
15. **"Ball hit the receiver's body before bouncing."** → If during a rally: FAULT on whoever was hit (10.C.3). If on a serve: FAULT on receiver (7.E.5) — point for server.
16. **"Volley serve where paddle head was at the same height as the wrist."** → If not CLEARLY above, it's legal (7.C.2). 2026 added "clearly" — benefit goes to server when ambiguous.
17. **"Player stopped a rally to challenge a CORRECT score."** → FAULT (6.F.2) on player who stopped.
18. **"Wheelchair player let ball bounce twice."** → LEGAL (25.A.9). Must return before third bounce.
19. **"Player carried/scooped the ball with their paddle."** → FAULT (12.C.1 — catch or carry).
20. **"During a volley, player's partner was touching them, and partner was standing in NVZ."** → FAULT (11.A.1) — anything in contact with the volleying player contacting NVZ counts.

---

## 26. RESPONSE FORMAT REQUIREMENTS

For EVERY answer Pickle Rick gives:
1. **Start with the most relevant 2026 rule citation** (e.g., "Rule 11.A.2 [Non-Volley Zone Momentum]" or "Rule 7.C.1 [Volley Serve – Upward Arc]"). Include the year if a 2025→2026 change applies (e.g., "Rule 8.H [2026 change — was 'doubt' under 6.C.8 in 2025]").
2. **If game context is missing** (rec vs. tournament, officiated vs. non-officiated, singles vs. doubles, standard vs. rally scoring) → ASK before ruling.
3. **If user references a 2025 rule number, map it to 2026 using Section 0.**
4. **Reference personal notes** when applicable. (No personal notes are loaded yet — placeholder for future.)
5. **End every response with:** "*Not official USAPA interpretation — confirm with a certified referee or tournament director if needed.*"
6. **For video questions:** First ask for transcription/key timestamps if you don't already have them. When analyzing frames, use "Frame N" format and cite the exact 2026 rule for each observed action.
"""


# =============================================================================
# SYSTEM PROMPTS — System prompts are concatenated with CORE_KNOWLEDGE on every call
# =============================================================================

# Verbatim custom instructions from the Project (with Pickle Rick brand)
PROJECT_INSTRUCTIONS = """You are Pickle Rick, a straightforward, hyper-precise USAPA pickleball rules assistant. You ONLY reference the uploaded documents (the 2026 USA Pickleball Official Rulebook and the 2026 Rulebook Change Document plus any personal notes loaded into CORE_KNOWLEDGE). Cite page/rule number every time. For video questions, first ask for transcription or key timestamps. Never hallucinate USAPA rules/cases/mechanics. Always ask clarifying questions on game context before ruling.

CRITICAL LAYERING RULE:
- The 2026 USA Pickleball Official Rulebook (effective January 1, 2026) is the AUTHORITATIVE source.
- The 2026 Rulebook Change Document explains every rule that changed from 2025 to 2026.
- If a player references a 2025 rule number, map it to the 2026 rule number using Section 0 of the CORE_KNOWLEDGE.
- Always cite using 2026 rule numbers as the primary reference. Note the 2025 → 2026 change number when relevant (e.g., "Change #15").
- Standard pickleball uses the rules as written; sanctioned tournament play modifies them via Sections 15–24 (e.g., a referee calls service foot faults under 19.A.1 / 20.E that players cannot enforce in non-officiated play).
- Wheelchair, adaptive standing, and hybrid doubles rules in Section 25 modify Sections 1–24 — apply the 25-series rules whenever relevant.

YOUR BEHAVIOR:
1. Start EVERY response with the most relevant 2026 rule citation (e.g., "Rule 11.A.2 [Non-Volley Zone Momentum]" or "Rule 7.C.1 [Volley Serve – Upward Arc]"). Include 2026 change # when a 2025→2026 change applies.
2. Reference personal notes when applicable (e.g., "From your 2025 pickup game notes…"). If no personal notes are loaded, skip this gracefully.
3. End EVERY response with: "*Not official USAPA interpretation — confirm with a certified referee or tournament director if needed.*"
4. Temperature = 0 mindset: maximum precision, no guessing, no hallucinating. If the rule is not in the CORE_KNOWLEDGE or attached files, say so plainly.
5. If game context is missing (rec vs. tournament, officiated vs. non-officiated, singles vs. doubles, standard vs. rally scoring, wheelchair/adaptive), ASK before ruling.
6. For video/film analysis: ALWAYS first ask for a transcription or key timestamps if you don't already have frame-level information. When you do analyze, use "Frame N" format and cite the exact 2026 rule for each observed action. Always include a VISIBILITY CHECK section.
7. Be straightforward. No filler. Cite, rule, disclaimer.
"""

# Main system prompt — used for chat + general Q&A
SYSTEM_PROMPT = f"""{PROJECT_INSTRUCTIONS}

---

{CORE_KNOWLEDGE}
"""

# Video-analysis system prompt — adds video-specific output structure
VIDEO_SYSTEM_PROMPT = f"""{PROJECT_INSTRUCTIONS}

VIDEO ANALYSIS OUTPUT STRUCTURE — use this exact structure when analyzing pickleball video frames:

## 🎬 Pickle Rick Video Analysis
**Clip:** [filename] | **Frames:** [range] | **Date:** [today]

## 👁️ Visibility Check
For each key element (server, receiver, NVZ line, baseline, ball flight, paddle contact point), note: CLEARLY VISIBLE (frames N…) / PARTIALLY VISIBLE / NOT VISIBLE.

## 📋 Play-by-Play (Frame-by-Frame)
Walk through the sequence with rule citations on every observation. Use "Frame N" format.

## ⚠️ Faults & Violations Detected
List every fault observed with the EXACT 2026 rule number (e.g., "Rule 11.A.1 — NVZ Contact While Volleying — Frame 4: server's left foot clearly on NVZ line during volley").

## ✅ Correct Calls / Legal Plays
Confirm what was within the rules.

## 🤔 Inconclusive
List anything you can't rule on definitively due to camera angle, frame rate, or visibility.

## 📝 Summary
Concise ruling: who wins the rally and why, citing the deciding rule.

---
*Not official USAPA interpretation — confirm with a certified referee or tournament director if needed.*

{CORE_KNOWLEDGE}
"""

# Quiz system prompt — strict JSON output for quiz generator
QUIZ_SYSTEM_PROMPT = f"""You are Pickle Rick Quiz Engine — a precise question generator for USA Pickleball players and referees, based EXCLUSIVELY on the 2026 USA Pickleball Official Rulebook + 2026 Rulebook Change Document.

ABSOLUTE RULES — violating these will cause failures:
1. Respond with ONLY valid JSON. Zero preamble. Zero markdown fences. Zero trailing text.
2. Multiple-choice: EXACTLY 4 options (A, B, C, D). Exactly ONE correct answer.
3. True/False: EXACTLY 2 options: {{"A": "True", "B": "False"}}.
4. Mix types roughly 50% multiple_choice / 50% true_false.
5. Questions must be CHALLENGING. Use specific 2026 rule numbers and realistic scenarios.
6. NEVER repeat the same topic, scenario, or rule in a batch. Cover wide breadth.
7. Distractors for MC must be plausible but clearly wrong to a careful student.

Single question JSON structure:
{{
  "question": "Full question text — be specific and scenario-based when possible",
  "type": "multiple_choice",
  "options": {{"A": "option", "B": "option", "C": "option", "D": "option"}},
  "correct": "B",
  "explanation": "Thorough explanation: why correct answer is right, why each wrong answer is wrong, what the 2026 rule actually says.",
  "rule_citation": "Exact 2026 rule number (e.g. 'Rule 11.A.2' or 'Rule 7.C.1, 2026 Change #23')",
  "personal_note": "",
  "topic": "Serving|NVZ|Line Calls|Scoring|Faults|Tournament|2026 Changes|Wheelchair|Mechanics|Conduct"
}}

True/False structure:
{{
  "question": "True or False: [specific statement]",
  "type": "true_false",
  "options": {{"A": "True", "B": "False"}},
  "correct": "A",
  "explanation": "...",
  "rule_citation": "...",
  "personal_note": "",
  "topic": "..."
}}

For a BATCH of 10 questions: JSON array of 10 objects. Include:
- At least 2 questions on 2026 changes (rally scoring receiver wins, paddle spin on serve, net post winner, conflict→in, audible time-out, etc.)
- At least 2 NVZ questions (Section 11)
- At least 2 serving questions (Section 7 — volley vs drop serve, foot faults, spin)
- At least 1 line-call question (Section 8 / Section 19)
- At least 1 scoring question (Section 4 / Section 14)
- At least 1 conduct/tech foul question (Section 22)
- The rest from elsewhere — vary breadth

{CORE_KNOWLEDGE}
"""


# =============================================================================
# PAGE CONFIG
# =============================================================================

# =============================================================================
# EMBEDDED LOGOS — base64-encoded PNGs (no external file dependency)
# =============================================================================
# Both logos are embedded directly in this file so the app works in any
# environment without needing companion image files. The PNG bytes are
# decoded once at startup into BytesIO objects that st.image() can use.

_PICKLERICK_LOGO_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAAASwAAAEsCAMAAABOo35HAAADAFBMVEWbqJ6cqJycqZHc1JTb4dzg"
    "6N9gkKkwWWLQ19ff4GPm1iJZUyFoameWkWGZoZfZ4bZhY2Iia6RhZWTU4qtWe5VqnGKxxZ1wlq3W"
    "3d2x3qo1NDQcIBunky2Ys8QlKSGOqHe7xMWzyNGntdHM25wpN0RlZQC1yKK7xMVFPS2oqFkhISBe"
    "X1/DtIttbamOoXyxzmvDo1j//wAFdAWQscN7laOHcC07PEA7UCKo5mW5y6Czxs4heXkA/wB//3/F"
    "pjVBQT5roqWs0DcAAHZCP0J9fYF1gmx///+3vsA9PUE8QUE4gpF/AABckjF8gn2/vz/Q0HIAAP8+"
    "QT4A//9BPz9BQT9tkW2BfX2Ba0CqVarGvLzCv7wAAAACAwJvl0uKtzMVFxJFWi5oiDhNZjNlikov"
    "dqU6XDcnJycyNy44Ryf+/v47YDpbhUtYdjcrNxuOuEUwcp2LqGobJBN5o04ubJkjKBlvljyUwTtG"
    "SUlWeEaGp09/f3+uuLm4xMWWwkRRVlI+Pj5VVVVDRzdzpDxRaEYXWI2YtHEzRB1agzllekpkezY1"
    "NjaGmm1Mhqlxd3WtxYlHR0cZY5ooKCl4mmWjunlVVVU3NzjKp02pqap2dnWPlpeDqTxlamvWuWan"
    "uYby5iyrx04bHSEbIidGRkZXV1eLpnX5+/RISElnZ2cZaKI6Q0e3uLfW19V4h2qnpqfkx3Syly+k"
    "q6x8sTsiHRSCm1Ht5BhpaWnP19CWl5XRslgzPEV1ZzLYxit4eHiIiIjo5+cZNk9VVlZoWC/RuBny"
    "6E12dnaEiot6o2SnxHfCy8zK09SLpnW0yKaJdzC3ubbKycqKiYm2x6pnZ2ebpKZISEhbYmOkuJV6"
    "g4R/v3+0vsDo1yxmaVEXRm1/fz+yllC/v3/Z5dA4ODm0pS+nulCkuo/LuCrcwnDDy81UVlSSiC6m"
    "uZUoKClZWVmWmJbCmzc7ga1ISEhwmLCbtIili0/HxsnI1rb//39HOhdVqlWOekeRq4OkupIZUXqb"
    "soG4x6utxM+21VDW1iG4AAABAHRSTlNkmtz5XCbd+6D8//ii9xwcGP7mZPnymp3WEWHh/2IRmaIl"
    "Ip3lAmbN+gOjZP8KZfD/AQKUa/+p8gTvbgMBAv+oFPwC4lKaAsphtP8C/3AEBQFiAbhdBzn/AxtQ"
    "AP7+/v7+/v7+/v78/f4H/v7+/v7+/P7+/v7+//v+/gL9/P79BAP+/v3//v7+/v/Q+/r9/rD/0fr+"
    "kLH/B1D7//z//v/+/P3QcdErkXD//Q4R/S7///z+//7/VSsx//z//zJMEP6t////a/3/+/z4sm3/"
    "LBQzTov9dPxv/gT+//7/BP8EL5L//6v//8/M/0+xVU3//lawsP8pTwL/A/+sjP/1MVL+beGHNQAA"
    "Ub9JREFUeNrtvQdgE1f27y8bE3oJgfTey/52s/23ffdXX/m/9/69zRV3NIOmxCOhEZIluchFcY/x"
    "GuzYxnGwKeYHmIXAAksNZCEBQgmBkJDeNn2TTS+bZN85d6pkyZgANtlwQ0CW1eajc879nnObhzvX"
    "htw85xCcg3UO1jlY52Cdg3UO1jkE52Cdg3UO1jlY52Cdg3UOwTlY52Cdg2U1lStur6ioaG+HW+dg"
    "DQaqojz9jvLic7CytmIDlKoe8Gzr+BfPge3q2cprxGFVMFAdR34k99AeCq2Xym2tt62wf3cOlhvV"
    "gcVtCIlGAoGAAi0QoYTShmUbzjpcnhFGpXa0gT1FZJ4XBEHzsSaKGi/rlP7oLRb3z8HCWAUkdjVQ"
    "oisIyiJl8vKuEyspvf6ts8q4PCNpVh0tlASAFJhVvS+ziaHGFCXN6llkXCMFCwAc+BGlOVEZuEJJ"
    "InWcPbQ8I+aCuyQaMVBpvhxNA+uK0FaVK/4mwwKV3kqpzGKVkAtVPa8E+HXdSXTF4m8urApuRQPG"
    "dT4zrqfDAruSebG7jrxxltAaCVhTuQMtJGD3gVFvNCssbyI0PyTX++ZPQlrfUFgV3CqZyDYrka/x"
    "azlgheanBF999STSpp4NUd4zAvFqlURAqJusNL2yro4XstNKVLOHhPppH1f+DYRVzB10s/JFnuye"
    "P3+dVp+d1jp2v9CdIsvOAnU63LDAnRrAB21WfGV3KBRa5xu0aUK3TFeNfJAfbljlXB/EdoUJLLQa"
    "XpwfCiW8g8PyCWI+bVO/aZZVwXVQnbFSApU8mJawLpQ4ISswrVATWT/iYcszzAFrhURZwFKU0Px1"
    "TEwl1p2QFTAVuqUWteIbBauCa8aABYYVaZxvep93CKx8Gt84iY64aXlGwAkBFq9ApFrnG3KrF/gk"
    "uXmkxdbwWhb0hAYrXmkMDcX9XLhSCtnxDYJVzi2jEaUGckJo2smxgqwolCIdIyu2PMNrWBDdaxir"
    "QRLonC1E274xMQsjlmNYA/3MJ4r1ojiIvYXky75BMauNGhFLyJILagKvU0p1XozmhiWPbIgfTlgq"
    "zQ1LEyiRQdYrEV2I5oQlfVNglXPraSCXF2o8lRvnV4PyEoWaXLXTkHTrN8cNn8fw7heywAK74ufO"
    "RZmqafWiCSvqTa8KekO0+RsDS21hfWFWWAF5bvd8tKx6TXNFMc2NKy6TxSOr4YcNVjF3bY9uCYeM"
    "mKUJemjU3PnzQ+t8Jixv1CtiAItoDq1UJWnZ/s2IWRXcVhAOihGyjPKMPT6oKfTJ/tBcDFkMltfr"
    "XQcNskZed2EN9ZORLZgOo2VtozjvQzBNy6mNigKNyLJOK0OgsTQGK8Ea0PLWRxxYtdWVdMVIlgCH"
    "DZbVGZqweFs/iIoszu/uDtXpAVFjCOu9Idaw0OXlBa87xI9o1Bo2WD8D/W4rB2Sl+EXGihfmQ2gP"
    "hbz1AT3Fhly1uvnV2IyyoCZakc3nq5ZbviEBPi1mCQG+kZmMoFUbViRqcv/8EJKpl+ezVg1BzNUb"
    "Cl7v/KQ0kuOtwwYLtDcrklqsFOj7mM1412F8gmAupLpDCXZPDT/XoBUKOz4oKmJifoqu/CbAAlzP"
    "27AEJRKqtsYpvKyBfohb94g63z137kMovKLRdSzUw50BQZxbRw9+I2BN5d6XqGzQEvQn57JBHUdB"
    "sG7QdDqxTtfrQnPnzg15xXWhUDUYWciryXV8SvlmuGEFt7GFMPGAw2B6aH4aLEu4a6bMEhVCUE/o"
    "EWgyjvXXRSq93jpZWsUVq3/vsIpxggObX6tUgqaKQF8HMat+wEC0KIp1IN0DEZyS2yNZjc1jjsh8"
    "P/1nRr747xrWP6ktRAblAI2CM5ohywUrCgmOJvhlAxM0SZYlp8nQEBmhrduYZY3ENPlhglXOHUEX"
    "5JGVLmhCZB0algkLOfkEnmGy6eDNaW5c0CorJUaxef37hmf/XcJSue2oGyC0KzrFaQ66gD1cPQh2"
    "AFUPnSOuFkDzMbFEIggrg5Uk81KPybFt2ft/r5ZVwXId7AdlWgddXL0SEHE0WqzThBrGxXE6w7xk"
    "RAJhyiEVsX5rWV9zh/p3CetKVvgzxOhcUE+ipmhgV8zzjHDkgJLlWDKZjMVks7FfsB/xAXKMYTRa"
    "w+Th1RHDFeAbdENgRZ4EWGBbdTwfoLrbxxgo5ASgKo0Wg1YpRxgsu8HvZQOXLLX9PVoWwIpUGmnO"
    "pFHdjXVKAEdypJijDODSGQeLk9Nk6oJlIkzGJIar4e8ywKsNEVaV4QOVSdmIOxHLmehaNKlYTHEA"
    "KVazYFWChcnooAwXuwG46KK/O1hsPeHzEV6RmYTqtQN1D4tHFCAAHkSh4K1K4wdo5j8IC6glm5qS"
    "yZIm2aAlx7oAV/PfFyxj5eU1b77BrMkI2BHd0p3IygnlcnJ3l0XI1TA4yXJXSUnJGjAoYKgYVgZO"
    "fOvfEyzWV21c1mZJcMOizFZJZLSrJP4q2d/ZuVvZvWe3LOeAFetck590340B7/nhrZueacvaeKTF"
    "zFYQFKFy5c1rCXUsizK7agKrWVPSJcfknLAAqJJ5v5Rl/KKinCsvr2Ct+OsF61Cz0cWbFkU+eua9"
    "Tz55bzMlDixmJIBqTUmnnL0xwZp5nwHrMMJSueLicqNdWf71taxWgwfjRAglN98zYfOknWs/eWYt"
    "IRYsvGwlHwJS5wAiJheaG9biiisz33Nb3619y6YeOvR/b924neOeLgaS2M5yWBXcMpaxGE2PyErk"
    "mY/InyZFPtl88yTK9AM1SIBpTdrdJBtxLZOLlAnLyIxkZq9XMlnyDxvf6njznXcuxdbs2CyVdqxM"
    "S+YrzmLLUiWTFQFQKLImfPJRZOf+ne/tJJOejBATFnWlhbLZEaQHJmqhsYOf/eBth1qbG2TLdNFc"
    "idQ6sa+5oeWyy8YkJdqwzXPl1KlTD03ZbvJSz05YKrfBgIXfcQSHC/2Bne9NCry3ecL+93a+Rxkt"
    "E5YLlytRtPMg8/+BrUdi7o1WW5lMpVJ1CpFXmW8/9uLOGKNH2GdoWdTWyhLv8lMwsDNuWexyCeDi"
    "+cgna3dOeu+TtT+c8NF7EZmZQQYG2SjPwJ+IZBUcIrlgwasTOdXfmN9ttPlzP9elVdyV5eXc/zX5"
    "6ifgDYBhJY8SF5QdQmtYrHKnUgg7kzHrecMNL8MLhs8q082byc2T9m9eS/6fT9gXTrASs9YNIMIa"
    "PH5txF3YchdvjGoNIKRSpLF7PhuKNVpYIcuge1SL27mrS3qJ4GWjbMbodmP/kzFULs8/33rIEoBn"
    "Eaxi7hAaBDHNgkWVSZvXfjJpM915+9pKakV4amOJuLHQdFhrXbDQ8nSWLoW83mjUa7Hy0gY2M3Al"
    "d8UemcThV9GoOdOrjiVbxAhsi779Fa3rDLphOdeHAavNA9/oWskQpZvf27xz53vPRCxdirB0C1XE"
    "zQrsCylDlk1ld/TCf/DB2AJedwsLdCITqSvabzpGeMbKaGBizLqgTepPwZs3b/xKVfwzGbMgahlX"
    "zkYh2HWSyM7Nm28mLg3PuLDmwFqL5sZcMWI+Js1VacTMLpV0WAp9S12pqtdwX5ZKVIymN28YXbK6"
    "urq7uwn8v3mbelZZFlj6rh68UjYOHTAHbFj/ZbYItayIwXKNViC3tWstfwTTNNwYTRQd18zE+TRY"
    "3gBdoR7cvv3giqsfJQr6p2VX6S0xtzsFrzjmtvfPJliWI5pj9nJEgmDvTqenwc8Mls6aYVruurtk"
    "W18PgAJXZqCtokUmLF16ZQNoVFW9qQmiu9sLmeFZDXA91F0HLyUvO8lAf4YTabUBOsGIUVFWGrtF"
    "xTIf2w+ZvDdYmTrDaqZhMf0BV9bfPXfu3PxK2yxpDxVw7NpCJYq6VJp3zYbRGzyl09JgmbREc5cb"
    "hqu6uzFJpK0nF+fPdIlmVQsYjp9ZltI4f36oDv2RBR4XLEaLWrQwZLE+MrLWNiy58Sej5mL7ScqJ"
    "d1QzIOGgPxtfoyWlJdhKk0TQvA4qX63GGltrrPnQyrzh6rlPkj7uyrPIsqZy26iuGLD83mocWBU1"
    "RXeKWlIPJTSjrTX01tq1Zh9ASN3cn8y1mmw+nki6iIPYAptLiBxk0mnA+sk0wgu1Vsjyscf4lQAL"
    "dZBMaBozx5BIdpxV9azy4g4qK8Z0bl4zF62KIh9xdYcDYDFgaFZGqtJDJo2yUXXPzbc7CBlQGZhY"
    "03hy1IBV0kRknBaOId6LqJRAgIl4yCTgf8SliaIYoctOrnh4pmFxt+HoqmDO5xbtWXw82obpZFlo"
    "EbOqg48glY5ZPfRQ9VwzbPUQSDgNswKzAXWi64TuLmXtUQq/9LJYFQb/R8tSdMwhqMIrlLABX7ij"
    "YftZZVnF3L84sJyVAppfT4UaIXyZmoDmbj0S6TcwsVadmNuPj8enBdCeeN3QFGYAjB1NHu06BqYV"
    "EXwsXiUUGmDiHZ4g8xHIJgX4mz3lnydzZxmsQziTNBOWQMW5ONSaUkyzyo0LZDvOAQRO+Nd3wlXV"
    "3SjYJJrUEZaMCV/Lor4juzqmbt14RDJLNRKhzA+j3mqFoZR+1XFEIgHgBE4IrpmkdNdJr5o6w7D+"
    "V1Dxuj2R1F5aGEixiX/1dWIdhC/SkzaQkQELons1zpisnv+d6mA4HBW7I9B7yoIIsQdZL9rluuT2"
    "9f/nO5cexmI2QMF1GtFQHdm07E2m1w+2kEqeQuyH+F6dIs0nPdpxxscN1TY6AJYQCOFcbrGuTqv3"
    "euO8TK3UJjN0ReCiE+EEm+kNqKqqwpoIlqUIgJkqkOU1mKjKny53jVCoq3ZRKgis06uWIWU0RjL6"
    "MMwRnf0iRKVDJ5tNn2lY5dybELTYUjB7lSH8gHbFYOGcyIQAvZTZNxraykqHZBnyvypvFCABKGhR"
    "TQg1EqWOlxXwNCovA1QVqnuYEhq79X9A1ALb8kYTIdqzDDQMLtHGMKfAX6AdwLR2cGebZRkzunnW"
    "xxtRSzBEhBcnHNXjlx8gRn8VcccuY4qkLBgpCtIKe6MgLkXWGxq9pfz/c9ldSZ1csaGS6LyAsjXR"
    "SJlUx9VDYFXxAJFBO3hFkjGtRP1wxGHh1CyDFgJh/RcLGsCqXsM10cCK11BZg8LmA7wA6hFUUSQS"
    "CIDPKCA2jPqBj8lwQQTdTWU+FY/XpWQibcvlSZA7NDBFhY5YR4yNbFQJAlacJwQ3GwzQIwNIT1FH"
    "2LKKcURMNrbqsfQWkBJ99RrIU59XgQsQ4Cbjp3lZpmtmvomEoPOau3lDSUKRtVioaXUpAnGnONfb"
    "bu/DlcSMVqQHZ3KVc4sh+YasyBxaak+zK3XDBYhrpGGpMmH7hFhim/eZSS2IVK+GPggxJBGyqsNm"
    "Qsf0d6IeBILDKpWkTBOwuiiY2ZPkenWQqQPrsU/Ero8nu5gVqW/A88mipa0Nsrzj/TTO3/3w4HWj"
    "LlC5Gz4cSVicusic021ELnA1kVmWGPVFfV4dwghcfai/UhFdtNi/QCThDWA6x7K7gE6YB8HD6/xK"
    "XVgT+vlBFrcCxg7oEzWmE8yHqRMb5OZV8O/obJvFXjBqFFjXJSMZsw6zBU6KUlNTg6z0uhDzOGPX"
    "HgGdEPqmGAQYntmMVVJhiV04Gg4LVvGcsDwF7dAP3STRvILQiGX3QVobUcBqQ7yTBarw+Kl2rcHp"
    "SUGh3sJdMG7cdUu5W0bOsrZbvSGkY0BLr5sbYhsbgXF5Ra9C0bDgq6epOoEJIKe2Ga1itbowLgCW"
    "AwHZ8mSI83V1OtG9tUKjQgdRS7jGMYI+rtNtXHuG3amq7a6GH16iPsstHfW3cRdwz468zkJaSqBy"
    "LpNYPsz7AVYEvTD8uY5WAz0VGpQLl1Xb9AnupnkTIvT/OphWY4q8mVuIq9xBiQpx6P9k1YlknLpx"
    "cfMbsiw3NBsLEPC+Kbih+odTkNZ3ctI685bVZhsWtkg/bkgg+ljQAlh6QBNqQQlVYuSlAY3VMW1Y"
    "qENNXJrNSgx7cWVPnYKwniTvDJK1VLCeOECk9eYm/fD/gcMtVpUR/21dhaP6xSsvuO4DjvvHKaZt"
    "TRmRsjJ0hXa6wybgNlbj2kJfLdACJCJFWOFGCsL65hhmxYqhHgxeJiqjcs6WQWHdjif0n0GbKaIo"
    "gJLoGAQWvP0irGwsNpwNgB1olYxJ0caEXpyZguNiKkb370zh/jvS+g3QumREYI12EmnmiF7cN0vU"
    "DGUFlhWBm95GCeQPuTn55JNgYaC86wwHrHJq6FHLKXG9WKxJpvASIB7ypRMtMVc7+iauYnZVfCWn"
    "LpOIlExWGpssgQhWYjGJRBbDLw9yUx566ALuhilcwbjfjBudGeKGB9YKmgYLOGB0F+IgQKFDBGEA"
    "juX1KgoILnnnzicnTWp8UoZ+T1cwMIfTmzfu1wmdNumTShLw1opeQewnJ9hkUnVFT24biPpkUhG8"
    "wTKjBRMa4qI7VK4YbOqhv12HtK4b95tR6i3qiFmWHbN4DWKVTwQDw6X1oiUd6hRMeybs/OGkZ+4Z"
    "NWpU/pM3r2VlqIC5BTOmjiizaLLxW8/s31mJExm8mtYYoSccoKkox6FneNDBPkpjSUVLlAWdVhas"
    "V5IyaQD7bOeWfmfhhdyz/10d9Rug9uFIlGhaqLWnCpOkAoMFoSsUErHgxEwLVSpktpGdOz955u57"
    "ENeoUd2NTU8mb7YLNxGp8slJ+aNG3XP5/v2bJQI+rAmNMmkdUplFLefUxTjZt5IPu1ExXAk+FiMt"
    "7VzxaG70Qwuv4/4HN/o7vxm3NAutM98bPk/NzYGNHFpg3SCPS3hE1KAijaQEI4Hxk8qdmwHW7353"
    "zx/uvvvue/7EoI361rfy87/1LXbzd3+4e/Xll3+yea0OrOqA1dA2qAaehxZR9EBfWSYrbALQWqRi"
    "2eG6hQsv4P6Ru+5vf8umH868zsJ9/mxfYioJOkJRqQytw2JW1FsHutGgBdpx5879t9/zJ2D15z//"
    "+e67/3DPn/D26tWr8Ta7//bVl783aSfxezWxP0KGtCsuHizSSgl4oABeh38Wmu1Fy7iE2HKw0XL1"
    "Q3XUwofUG767FGJ8lq9heOpZNUZ2x8b3cD8xqss0IKCGx8AVMAtPEOPpRzvf+909dzNWfwATw9ur"
    "L18NP/zhT+YPz0z6RCKgaJNkSAc24BWzPjBmemAZUkokEi8u/NvChQasmUE+KdH/xpV/CEEeHBFN"
    "axxa2PDnhm+yvUnRrFjcMtZoBigJ1LGxdzAvIYAFc2gKFvV2vgewfnn33YwVsysEZ7AChu9VQl8Z"
    "EilpWT8ky+YmL2KB3WeAWbgwIQSMM5Em+F60cAWVJJVWcsXPchf87SH1uxC1soX4M29Z5aCidWvv"
    "aWdQlFcoNYb2UEcJfqNgrtOmJCVrP/rhM4BqHMODrCDoG4Fs9Q9BjndRBUeT1SF64GVETi73sz6w"
    "rGxhcAINbF79hz+Me2bnBMi1TF8s8yox0qxWqL9QfwLh6h9V0PED/XA4SjQ7sPhXw7JgS8gbI34g"
    "AFCaMrFu1GdEIk3KnwS8SGTzMyxI/fKX4IK/+x0E93ue2fwRIfKkbzVBXxiiQ5inACq0Q0YPHDPm"
    "QdOsfDSwc/W4e6FtWf3aa/WRyDqTlgZyC5KiW7gLvnOh+o9gYN9ZOgKw1DdIxCxnmaVSY49SSPZ0"
    "EgcvTJtBpZHIe8/c/q38JpTyayMfgfTaP2nSpE82V+LAu9zUPeqeZ9YGvN5qqUU9cR+4kXngpr/O"
    "Og/iUrDsxYUC2bn/tfHjx69+7eV77z2+5Y/jNlPfQiNsKTGK+y9+V/3OQ1jr+s7fRg/wwzPfGy6G"
    "kFVjTjpSXHvYiWK9QHUxfa5ZFBK/jy6//JnbUZlOerISV92x4eaInGzKR/Xwh9t/CIYVra6kg6sG"
    "sw8EDxwzZ9asGUADfHAzeWHLA8d/f3xnJBJ5Aazr+L3jLjdozSyLyjHcywyC+0OjuZ+r47Kk02c+"
    "3YHOMGDmOwH38KEo1kGeZ85IcDWe3PzM7bff/QdTZYHM+papssAV/3TP3c+sBbUfDaVOePRARwN6"
    "4P3nz4L2IGPlIy/cO/73x/84DiP8i+P+eO+9l1/+x8vpuoUzWY8Yk6TJ3C3clFEXcP9DvS5LXctz"
    "5g0rYu0QLAfchuVbB2qe6t4BDWwLwhR0f39i2tRocIepJHaieo8m6uhbuWGBya1qNj0Q22w0rIUv"
    "UvTAe//4x3F/+qQe0oRxfzw+oXLCuM16cOZMrAWBI+7g/um76k9u5C7hLhh+WFjOkq2tjSIuL8Rx"
    "nVBAUAjTDyKTXFGbVmTS78zu78+gslb/8m5DdOE9k1haiLC25YRVzKSVnNxkmNUsCFkQ3P/GfzRp"
    "ArICWtiA1rjN+ze/Ni4ivDiTDXbH5N6N3GTuOoQ1OovQOtOw1EXOplm6GxabTUZlwgMnNn5Ta49X"
    "QHZN3xvH7Op2EA5/NuQ7Y/UH5oQnsKxiTv0Vlhfuv3OW2eYhrAR9bf+ELff+cdwf/2gCg78m7Nx/"
    "/DUanLmgKlwV9sdoMzeV++ADbgWnjhpY1DrjsGIDYAkWLCy9E0WLa342zuXTfIaAEL1xney8HfU6"
    "U6QuVndHIl7sE6LVuWNWMff+GDSrpyxUs2Y/iKqB/+iFnfX33jtu9TiGy2A27oUt48dF6ssWVDHT"
    "kuBVudGjARPCmjLcsNqcHcZ02YAlmqxw0gbbUNm0LF40xle1uLcOtDzk1H9OY3XPPe+xDBotq1rJ"
    "1RsWc6MXQSJ4/2+BktsLX4zsfwEMa8uEnR+NQ2esj0wAXvfeO/6BezdPWLiAzaSAqNXKCq8/V9VR"
    "U4ZbOqB8t2NWICJYJS1rSnooQioBVi2wqhUVrNvgUKuS8tYJEUIr9z8Dkd1gNe53v5tUSQKiV6tl"
    "mCM5x8BUZPWUTYrBmgnp4NoXtrw8fvxrm18MrB5/7/gXaIROAFb3jh//8muBIIOFfmisaOF+zl03"
    "ZbhFaQV3qMeSDgIfgaTQMC0rOlU3UlLJY0oNzc/PnftQCOO9zodETQiAjq/8BIQ86Pc/3f5JJaG8"
    "iFtxRqPRRCNpzV57r+BawQefcqHCkFX2YkJc+/L4B0CO7nwtsgUQXU7lSACk/PHxWyZsjiSCVZYf"
    "bjS8Wx07ZSRGd/A8FMUcCRMYN83rswZSq0XIYDQcjQdGslKXYgOtcV2e563TNJ7Nb4wEApG1kD37"
    "67z9lbzIvLAyR8gq5lZJvcn701jNerDM63tx0tqXtzwAmPZPWI3ZznhI5TejLh3/8oT9kcuDbGyk"
    "VpGoObSmbleHHdY/ce0SlQOyOb1bUHACCyTNUcu25tcRwrPNN70i81GscuH071RjoxiPx60BVl6o"
    "8zY+KfNsOh8YVs5taJ4nyU1z3Khm3fWgwCcWrlt7+RZQWdh+//t7fw/B67XjhhdO2B9YF8TeMBxV"
    "ZFDxIzl8P5Ft9WcKU56yabOulTQP1REKwQrCFa/ENYheWNuCzDECuIzZgZoWr6sTvY11oGqNKmE3"
    "2/svu2FRKWY5oWleYxRt4cKFIn3h5ctfRlJmGz8eoL3wwvgHVr+AbmjA2mQXFId/wAJnkOEeklbC"
    "o/BshpF5zQat+Smih+ZWV4dk5mJeNjbo0+Lgg7hQ1Wx1SQmnHxmdKM4HqcjVocQ2nX++yw3n3K+E"
    "Qb2/Vk9Xr66csOX37jb+chJ54YHjL0SCCeaGCCumjpxlsUPnavzWfsoBLGMJguBzaIXmCwatQF0C"
    "QldYNFaZYBEHj96O4LLpiE4DPLM5jHNJ0pz77WjSFlgAbdZfnzp/RllQ25x4MVL/Qv3my5lljX/Z"
    "MLDjAcqvHn98c43ZG6IbyiMGq4ItR7HtCq+fnVVhTo61aPGE9s+dq8TBF8MQsBhJttIEjzNip7yD"
    "FhMcVrlq79Dt63LsTgfWnL8CrwejNQLoLOEj6Pf2HwfnuzwQmfAyIPv98bU317/8wJaIwDrDKgjw"
    "IwmL4/qw8MdbVQdcP6KzKaQuR2S2ReRKqsyvxtnxorMpbloTjS6hMffxv/jdxJY74f383wKrOf4f"
    "RxcufPFFkax+ecILkEqPn0BlOuHe8b8ff3z/ZlDwr1HvgipLOqAbqiMES73eSXd4Wa7TeL4S5yoK"
    "bNmWTas7hdPU9Ea2nbLdcK6NPXWGrU6CntDbT7fmmt5Qwd0GsNK6wlnnj/EFUZTWVwY+2rLlAWjH"
    "V6+ldD+IBtBdW7asfjmimFOhw3ysZ7ARo+HIDWV7+F7nRcHPyzgVPS1sgW0lQXLyge75BizRFzZr"
    "zWBnPkiF2NAG+yuU7D2QK4Uu5w7T2Jg0VHc+NSOYCJb5JvDrEnT/8Qe2bAG/e2H/a1vGI7cHtry2"
    "n8clwsiKpTt0BKWDu+qAi4sYLWMFjaO2GK0U0JIjjSHDpgTcIjcsuiIbkgsnQlpIzl1QLueez4D1"
    "2/PPnxcMJhS9fmEwCB3i8R/+8OWXgREzsQcQ234qRGujYUAVFnA3wUEOBjzjMavZzg2NsAXiARyR"
    "KNoAWk8CLeWyuMHFn/o85BXTq4LhEKgxoZ+2DjJ9rS0D1qzzZweDgl6TwIGddZV09WsTJryAmF5+"
    "4fJfPvDy8ZfrKe+tBVigHLxKklIQcOoIwTIGpGXXBC32pxJUuzEfPurSW6DllZRmlAJFYCpkwAJW"
    "3UKdLK0YpPjeIJmdIcqG86E7nP2gNxDw4XiFV6hP8HT/6v0/BLN6eQK0Lce3ICsfwgJafFKWBj27"
    "bVhq8I5pgRjXRLApBazIWj1g06quA1lRZ0yIZ1NF0prora5u5FMKzX1IMtscYfmY/zLbaP/lvPNm"
    "zACz0oIzZwa9NTV4johAP7p83L3Hj0Oy89rqLasDVMBFwbU401eTkz1sOGzEYhZK6ogttPh6NnkP"
    "onsAg3y6bYliHQ53RXAVk3lHWsjCxz8pQwCuyD2i8w9U5sPBGfYMmTJfgE+AWSV4XYDQlViHU8U/"
    "mrD6l1t++dr+zQGqiN7CWoAVjYZr5SQoikE38zzjMet7KlbhTUcU6rRCnEeDq79xIVK6bYUgqWF1"
    "BqLLvCbaQ68iExliXKgL4W4zgw+PxISZdisL+2u8ZeCB2mX+cDDMjn9IJMTKCNtKQo/ghPxaXOUC"
    "lqVBwJIHN6zhSKRb0VpkC5boF+C6MaPGJVrG1HdXCIdOsD8pm8Diorm8Hv+q1erEUIp2DDYABslO"
    "DJ2OtWDQ50eJFQwrATE4w5o9GPWKfh6zc5wAbLDSvGEB7EqSyeB7eZ5x6cAdgERaDxiw+LhPVPgQ"
    "rjFEWjJ2iVrmhh7V8+eHRN4AZrkkhvw4XFQjmTjY9/IWlSujFquwpiWC94FZ1WjpqGrYvLDa2lrw"
    "QMaqUFSWA6sYkbcPOhY5HCPSYFdywKxn1Yl8BPfQERgtJRstlFMArNEBFgfHheeCoQGs8itzvxcI"
    "B940rGC0tgoDu88PyGbMgTYGW9wv+9GigBOeEYWsCuE+MGY5dsJNMc58pbSZYNkgYJ2opgk02R0K"
    "eXEDAUL8caQlDhxpRQOb2x2qY4u/qY6rdgBbfz95J+cbTeU6qLzcZ7CqikaRWkLTwsHgjN/+B7Y7"
    "77zzqevv/+uc8w1sVvPLsVhvT4wtyCvnRhgWlbHUaZoWj7lhEsSlaNBiqbWQjRZG/Oq5CAwXFDTs"
    "aGY7QJDmXYfYKGpxccY2keU4fTVpGVbVArSuKp8XfkZWv8V251OA6rfA7M6/Wm3MmF45iTunS0Te"
    "eqJF02ce1g4KrABWgHdoRfg6L0JCuYW2xVYQZuWVAGApIq3HiS27mlvYwvPmyRneZwSsVYtoUrEj"
    "FiNWFQzeN3PGGDQrxuqpOxkqpz11vRSLQbSCNKf5fW5q+8jCAp1FAoyWHHDOCRMEyI2RlozLB/FG"
    "rTdXS/AkdsDqLia/2feqRHC4avSUjVunTp26kW16iKHmNolsigkAx93uu+8+0wl/i5gyUN15f28s"
    "plPARaRl6Mc4vjqCsNgJ7swNsYZnz4UXCzXQTQ4trFZlN62wQOT3WdxVK+BqCi7e2wle03KZubtB"
    "T0vzMrbvxzLaI1PLCV2sZoZn/4dpWb/NNKveTSywS4S2rmTzud7yrCweUTdUZWLAkq0UUYC+KB6v"
    "FTU28sposemTmi8LL5FSK+6u4K4Zu7fJWHQOkhL372abl7UsVjc2UzKtqSfDsO67r8owrN+aLd2s"
    "5GSM9sY2UTphI0O1daJnhGMWW1AuoxfalgWwCguxRqWxiqBNKysvP9lh9uf/pr6/pxRrYWIjiA/c"
    "j+7zz7u7Q/2VlDZIRKJS0zQaDabb1X3h8BjbrtJYXd8bg8C+PNZD5F0qbvqgtr7awY30gvIKbgcx"
    "aVluGC9kOYZBCzyRgJRySqK1aRkhlccaiwRHr1hxU6lE/F57xXmiMc7zKTH0JMgKKSmT5HQ54JjW"
    "fawZXeEAWE9dL0NQB3lFpD7DAztir7Zj1BpZWExq0YCcFrNQCxZatBSUULybl2YGMFzb01RgfN8r"
    "uav3yoRPr9k0poyddGNSchPt7eqiAgpRyHHuMxvrCgd44f1ybDmVksspXeRhgf1Ac2XribfjHI5j"
    "GdRf0fQAz9fFC7GJFi3dXAdmltxZADOABciaArWdBawr9nYRZYB0/XxSJSEACxd3yU0xymYthMNY"
    "cwAnDD6IXWGGZT11veGB4LnLjP0eDldWdgzhSoZjHvw2yqSWAysggHCu1RguzVhJEMDMJiLzJjBL"
    "YIhUKr0CQHFc+8qbSiity6Jccfc+kEpYuUt20QDOPanEFYVRwDVjxoCu8K/3G30geGDrBmZN22I3"
    "t0KP6llxonVAw+GGzSYrm1aE1+J1WjzOXFGzVqmwTBBXSju8eNBYpfmgqtonc569R4mSVV4kuvtx"
    "WjKe09Alk+VsM+vlsZjCa+EZpm74DxPWb++8/36mQqEPbNhmrG1tvjm2jXWFk4tH3g1xR0kTVsCa"
    "AhgHWoVxcLXaWrYXncVLZ1vNYQDDCAYmR5KlpVdvQP1+0xO9VMghXKv7SQ9bwxtLmsufe9nxNMqY"
    "O112hRr++uWxXvDAXsI8ECTJMlk+rHIrJybXnwVuqOJomG1Z9oFhdfF4HWaxInNFe+UFX2lsgQxC"
    "FoHFEVZJaelNV3tu2tPJ5ibloJXCTa2BUFKm9gbgvZISW34/pDiWgn/K8MAkqNDmg+aqArn5AKcu"
    "iy3aOIQ9S4cDFlqWEeArzWo8jdQZoxKoT2uxV3RWqvDWntEIDN0Q99LEBuE9Jyzv5xCCZFnCJfy4"
    "ntPeyhrMa9Om+5+Cdv/9m2LQ/0nQD9IGCOdTQVodji1CD1wUWzyknamHI2a14akeRsiyhi4oJNBx"
    "nBivFdbif6y251rZ42yyLZWYrTRJspYnNDYTNdRoEJIaFnsOTJ58oGNxm2QeiISnRcWWs7X2tBf6"
    "QCLhtvC4VU0yth67wtjEA0Pb9Hw4esPFNGL4YQ30etZIPstZsN5Qi6ZlUBDjtoUxYBDDKGna84QB"
    "azkRssES5bg185ku2rXR6dDUjmaJbQ8o9ZrnEuDBRuCBqxiajROTGLa2TZu4kRvihnbDkBsac/+M"
    "+F7paC3eWN6KMb5W9MdD1oBOnHWG1onvlYQ2PYpuWFIaI9mqqt4wrxt7tOrkMPty2oFEewW7/I19"
    "aF6se9y0PLYc941oWM+Ospi8ayKiWuHJm8oNeS/9YXDDCty2yhJa7g0xjBbHHlGQQ8bOduCTxmzc"
    "uCkf8AKlWNOjpXtwO81stLw6UxSJFHmem+q6bBWj0MrFDZKz+alxgGSx2n4gz1OMB82sPHAyp+EO"
    "S7qjNhOzAqg4Q4h+K7kpFIFPpPJznG7EpngbQ9LYBLbolR1AkXyUBa0BhS9RDPPUcE+2BC5jIBHa"
    "qtv62M4zb7SuH82+O05d+f4K66iUkzn3aThgFePUdJzBx6alBeyZy/a8K6AV1+UUjidCrPeuwwmA"
    "DJZmPLZSwb3Noa/T0yd2aWiVLNk2pjjTrQMcSjU6OdXa1MjeA6rdChLc2QUL5UMrJTpzRVZ9UGrY"
    "/GM3LU2RZQVY1daxAzrCjJZfMbhWRgib5I2DZ67xjbhiDJQpBO7UBJoFFmcdhMfCWXGayZ10G57D"
    "IOGz7aK6EeQNGYGVO4dWHEunbABPE/m67mrwSGZZ7GxSRVEikGb7cQdXHDxz4QoY+Y/AZl7yRMo5"
    "aQ8M67unfhnDdo70pTj336QVYEc9umkJWGlmQUoQ5GR/Y2McV13IOkOFm7cqNYIOstPcUNkqegkY"
    "rURvHCf1gnZr5b42B6ux3rq8Qs0Jy84QZROWa89EtvcmSivcTxt+y1jhHtuBSAQ9GFwUj0lhm0YJ"
    "muA3QpdImajXdBKBrHKRekPx1wNWRZZbbli3sQgvRwImMllxVpmjRvCbP/txLQbck6rUGStQr7jR"
    "QcAvBAg7Jws3dE8p5hyJAGXRKoC74bae9MbSI2dZ6vq+RYsW9a1Xs4VP3L0bB1tllBA1xnaJSo4G"
    "2AJ4okUAbSwgG9sr4t4/mCvH5B48+0MPCHGjQs9sk9KJ297nzjSr0wdrW4O1///i7Ke9NeAwfsRZ"
    "XZ6Tld/PVxruiJvhOrCECNkk4emiMq7Ir2TaSiBKCjeIWHTSMmDEYDFpQKU+z1RP3iJK21YO/OTl"
    "bJBHTtt5TMndeL8fsaHL4q4QNQxWJend1BOL4QlPvSC5wBELBRwUI0TeyFUMw9nup8mydlDcRYcb"
    "3dqCA3ly68GBMY3JeBJw79KmDN6YYUXYCjxj658IHnm7CU2LSBDEoE/g2dEOi1YNz7HuntMS2yH5"
    "gxdSl0kB3GVNhiSMDfGqmTKe7dx6Qlp+9gcadgTsQChj+1w2pQvNk8D3wZs7tEh5uzZyXDv3NYFV"
    "wW2U6ETuSnViRMBRlYfmNkYIkQ5nRHr4sQ+cJw2WfzDDYn2mbm01zGgpBCt68CoBPLwQRzl0smk0"
    "Nzx2dVpgFasNpBnNS/ex9UJeITUt/3pCGzrS+/Jy7ghR4n5eOLFpsW3OA8wJQWNZ9/J8HA+kIJE6"
    "r+g3jvnguytJy2SO+7rAKuc6qPT+VNUT0ZBVQpOnjYU2EY+XmJxuWxJJsSKDXRPlc7ghi1agXXXj"
    "nFDZRUvXFWNdAZvXlZjfLZMfqRVfG8vi2ugONC8/W2fM/9hTMPHVCROaWw9fRuQVxaoTtLZR3V17"
    "MQ/MyWFYoPdBa0lTO5qNU0oVY+GwENeszNDY6nX+55TexlV8PWBVcKskuqqduy0SraoK+uQdS/ta"
    "duSNHfsKWNcb9Ij7gc2EDxmll1qblqK7g5Qhs/AWG5WXduCpeqtaGyibvGxNtEwvaM3tJw3q18Sy"
    "sMTerBarLfx9VTMmyGsK5EVjSzubknhafR7u0ehoMYk0hozKnjm0ygqhcsCJ50YShP8S+fq2xVgs"
    "Z5NCOtokuydNL8SHvaHuSM+24TGtU3dDtYHuUtWOSDToC8ReKWh59eJjPVgJjhGSJDF7LABDm9xt"
    "ztIWfaK1qE6302ungTqgHrvuZJQ7W1n9z9gHL33G7twkOXKGyw2nCVYxd0CSVk7m2nivX1/+itrX"
    "MLaLLO9cc3VBwS4ZBMQuzk5+msmT8x1Yhh8qWLhxqhFWfdBY9efU6sxqOmWj+1ihKXTBmt9/grn+"
    "Zw0s3O+9QS1eKclyjdA5WpXzdpOmPWPZ1Mzth2XzFCHEpUo0FLJgmSEeKwkug6Ks5oUn6zVnVhDY"
    "PMldbQ4ux7iqQ6RF/XrELO5XtBX34l3TpMXzuA55rCzvKZ3CTV6xAjeSexV853qPinsVLiPKXK8D"
    "yxpslWUXK/Pgw5ZsC3SYPoC+EZxRSFuRHuqmssqpXwdYagvJ38A1dO3hfU0F6rJFN9GmvR42SQh+"
    "9zyVcbYZO85rEemv9tq0jLkzCk6bYWsKQCgQ2ndoYvMbb+zoULNXEFQMTNt+1EN0vxtWojtCJ38d"
    "LAt36KB7Ct6X1vRr9fkqd/jVfLobYE3GgzXUVykPTlIHuOTFG6nenciAhYZlHdjHtrgfrH5o/2KV"
    "xDJDV4SXezzDkh16TjVkdRB5b8GR5aVCNPV9DmCVkN17frIdLAOPRGBXVM2O1u0hqbleBxZOJwLD"
    "alDX75DZuXxt7Ngv1qmVq4NXZNuokrYxxFxl8LViZw+sPtJVWnBZf6cWnVhQrE7t2kObSktvuoab"
    "rHqsQ7EhJ+mHUB8KObAKWc0AV2+Dd61Y5Vk19JAzdcUONpgjurXDO8OiHTynHrIevalDWsNXCdsK"
    "2rldy/dKyb2le0sLcN29fT3h6mpZ6fY6sDTLsCqcQb2co1jf5Yrbb7nl2Wcff9q45wjb/McFK0UO"
    "P1txS/tZXoNnSrO04NamY1r4x0vzuI2LGkq6oIMD4/rptb2Kez1OoD/k9kLDsDrYYDqwKM4EpH64"
    "4vHHb3n646dvyHzPpWCyehqs6jrSav7ye08/XfxssaqenZbVRo698i9yJ1/lm6jmqWPyFjdNpDFK"
    "kk+U7iZxFyyRd7HyajxbJr1ItSF9F/D8ov3Zxx+v+N4NN3wvA5w6+rkvJj98Yd5V7+YvWRIKzWN7"
    "/xS6J0lOuPDGyc+pDqTvfe+Gx59t/8VppnZKsCq4jbS3ZGxLV6r2Pm0bN7a1aembTdOkvRe/So7t"
    "OUrjrhXhorn5oduwpLdumQp8nn46i+utnPLBI69/+8KXrtr3ORvODyXgT8IoNHgjuD9L3IHVSJTq"
    "RAKXXex76TPPB1+MdjN6/OnH2y85PcZ2irCaSdeePKlLqIqmVC6v65W2lqbY8pKSPXJvaQyH9AYs"
    "jDN0g8AK6u6TK9Wlz02ZfOONnm+/e+FV+/YhHdws0YATDmfMyJLZKX+uhehECVtrYIFZOBTad9VL"
    "3374EaDmMLrh44+Ln20vLh4hWHgUZW/JK7LMR+8DL1QXlRz9kRSju0tKnsgnXTFdEAYsUhXtrlAh"
    "0pQpH7z++oXgXfngXdXMegzjwQPCvLlbWGHH/LnWJJJA5kPCYWZpoUn5V3124esXbEizLPDQ9nZV"
    "/fmwWpbaQJpKLyJKbVVVvIPb0V/QsngbOfpESUl/SZL0RIQBtETHCXVShxeD2wdY3jXUFuZJmtBK"
    "5NNIzsca1KDtexeNbcpS1+f/+OmnH3+2uFgdBlisnlySRwO+aLiK3/7r5L9vXSYvbUiW9Me1uhJc"
    "7CVk4BIdJwzYnnPyLZwictyBJcYVog/hWSa2JfkvXXXh65c8t93dHTz97OO/AOnx3dMPC/unKe3F"
    "3wPDOvZKs+6rjc6ITuSu75wmvdp79OJNAe1/iWpNJTFibMHmOIxR+NMMJ6TiV4cVZwcdGqe44sRA"
    "ieon8WyvZWtL9l3oef2DS9I89Omnb7nkEjVHdzBUWB9+WNy+4vGnP77hBkf4vEnkJ/KoUFtbdZ/4"
    "1m1dF/f+4FUiTd8j10SBVrykqdfIeF3lJ4sVT4nwlVl5wyImSObBwFJD3/rRMv1qLxSGvmAeGttV"
    "j4H0WOruDm74+PHHiz/8xZBgAdyft1/S3v74448/naVj31DwiOf1BrL7lR9BnxcOVwmjl5csbZPH"
    "NsmQ7SzXhWjYK6zJhyQ5kOaLJivoCRXvqTQqlzz66MT/dGTxrq24VIVrod5E+BToIzQ0tUn7rvrM"
    "80jBQbdlfZwblvps8dMff/zxAEBLn/vikQse+/ZLL+3bxzouPL2KyKV5UjIp+6NVE49Mz3sjT4pd"
    "3EnJtNJjvThRbzklik6o4kxXM1EJMtFPiZVX7y0p3VvgyrnYTkpi+vWbBIxZl+EhY2PGtmTfu3mP"
    "PXzJ6KVDtayDG7Bn/+yqq/InLZkUCj3I3tPpuMJ+iFitsT0X53fxNf+6qCTWjNsjlECs6tldMl3C"
    "FYTgauBwupK+z+GpBSzWInRNaenV5VPLKyCjfKvPmLxNdDy2DtGgZy3Jz/M8/PAX4FtLp1zysOfd"
    "/M9DjewCToKaFz30MyeCpcNSOXjdzy58d99VkFV4URUmcr+8TteMbdldukJdesU06dgrUl5BS+wo"
    "pIWdMSonISOJe9muo7hO1e9mFSA0fmqsQGg9CrC4X3PF3MEdEpV+dHjZDw7HSGX+pxc+9voHH0yx"
    "Lk8dfc2XV3x5DfsROqUPXv/spX1wZfDFD/0DhEZz2WE9zj1W7baeQaMsWQ5eWHI1Hn+jxkrG5vfu"
    "GLtbjsnLj5WskUgkbr5GOByXCYnY1sVHTpmVV5RJJ8JqVysONtCGZaO//a8zwMmpNNok9Osrvn/R"
    "X86DdpfR4NZFV1xxrblpgzrl9Qvz8gFa41D8s3FKDljFnGeo1xEWSFfpxNgez9KlKndk2li5S5Jj"
    "a57owtKnTiJpfXU8YJxqzPNKBMB9VR80g1AdvBrCyuNW4FZmbSp3m1iFm4QoZDogAjSA5+2333bv"
    "/Pf2rLeRGVD7y/c9X442Te2SRx777KUln5tzyXO0eacDlgJ94atHS0s7p+WPfdWzqLkVp00lS0o7"
    "j8lEFjOkjSZbZ0Hr/Enx8Zq9Owahq/LAyQr+lZLlMtldWnqhevDXaosEcb5DZNtdaGT5XbNO1N5m"
    "MMHS8q64wki4VXX0B54LP0VTywrtktxuOGRYMnl07BvTS0ukZugQC7pa8sbmN4HDyV2dMhlgO2Gv"
    "KOAkZD7uDQ+9K4ccJf8lCNOTv1iq2v1OC2na2wSwSkpxXEOWLy7YwE2own1UwrT3xLBcO58z/7zo"
    "+z+4wghq4L+TH34s710sAoUedPwzl2U9yz0cOokeKa+l85WOGKcu27T34um9ctOa0jXHkgCMz0HA"
    "mzMYst+FHnwQ48jn+fteuvAxT8FKuxtSl1577U89P/j+Xy46767pVCrtxJhVUnLgy7F7j5KmR0vG"
    "enxlwSqfGCFjZs86yfb22wa0v+T94KcWNOgKHgNTw5gWCoe254D1Mff6kOMJdIb5cucr22LXctx/"
    "Xt41bVr+UXnT0WP5j1I54T0ZVWNo6Px38y58ePIG18Y5o0dPhjh9kR2mjSA0h9A1j1KKsEr2rokx"
    "Eb/74mmCEpFvbZFp75xZX7XdZZmax+4Jnpt842ceNYfOaue+CA0dVkme3HnToVj7NQUXP9rV1dnZ"
    "eawriYaVOmGcZhYU+nzfuy95HrvRlZ1B+Pjplz/4AUNkEno745LGkN6uPRACwA1LkoTu8OQta6Oy"
    "9Mb6VZ631u8aQ+6fdarN7D6xJ7gmPUnMDPDqycC6SdldulTJ31t6LBmLNbTISRAOy7PBCls9mRmn"
    "b3y9YIOj9cDLvvyp56J8x4YGIHLactKULOknR0tKmijpGtvaEOt6lbZ5HtnR0NvDpGnPmLdnnZb2"
    "tmFqHu4fssP6OcctGSosmT5a2tBUuvT5pia5bdkh6I7VgleeeLQp2WslfuEwBiGW3r+U53kMi0mq"
    "M8EU2ooVKzbk5Z/H+noWmmfPxj+Dtl7y6LGjYFq90Ck+URKbtuaVTrlkbDMut29b5unIm0Z75sw6"
    "TQ0+yuy3L+J+liPdqVDfHbp06Cx9NblnbKu0eKVtlwWle0rXxIgSmrfk831XvXTh648UuIqU6tIv"
    "vngYEsx3l3y+5KWHP0B06rWTr2Di6C5EdeLwDCELMqrde5tiRx8tAa/vupq7NNa5XKZS6wZPawM7"
    "0J70njZc4JQ5LQu1wxCjc5gnTaXnySW7J6qcWlEMibq6fftK7gAGXplMdACp2wsewRSTCZl52Ckb"
    "LwBGtyTvsS/MTvvLK0xiswYHNobQZMmaWNeaErDhprFLWzf1FUyTJ06U3/c0SMslqaVv2eFXKZ3z"
    "9mmDdQ33bzlgqdyUIQYtSHdipSVyZ8lNG7iVK1duh6auOHgJ92UpVuDl0V/c6PG89ClkYo0Qyuc9"
    "mEMzzHswtOTdGx8xeh8gdpFFbHZOWHxNsrSkKZk8mj92KXzgvK5YfoEq7TgsQx7fV+BpXQTmRWjv"
    "mMHda8jtPHWQqoO65AQ92Twjp38pT6Jr9jTJYEk/3YDPW2nuADd2b0npJtr/YOjBIcv0eY1LPn3s"
    "keeMHOTLK/Iw0Of47PcTQQwcxT0x9o7luJUHCwoKOmMNWz3y8kdjDUvfWoTbYvW0TdwhkeXZKNlt"
    "yLD+4njhAFhPc3mZA0/hB+H7R0T73n33254bHznw3FJzJl9TaWksidsIXP1Ttl3Mt995551LJ3bi"
    "0m/xpDM/JPaZQYxTN1yRh245fYCN9RKttvayJraNyNUHABU8uKA/Nm2smid5WmUak9veeu4HE5+/"
    "HlzxK8Jy/3b2XVcMAqvd9EPoyLyGWISe3gM9/UqX5FBHfzH54Yevor0lJU8k5d0lENVL8/PGTsO9"
    "OZKLFslrjpK4KJ4kL/bwBxtDQGyDFcjy7NBvffheqmm1mn6s5Amwrb3oh+qV3FLwxf+kglXtki57"
    "a3SrvFyOLQJXvD9dzzNGg8HKgnL2Xe4JFAPc8Ltq/iQcNboRMrLRziMhTm8oeN2Dg3xLJjWG5nnn"
    "YaEEd9U5tlxOHutcA7pH/veLWcuHRNcv3IFrn+PxwpOCxh48b96SfZ95PjBtjAWy6WYgm91DgVWt"
    "QPlUU3/no2s6PKuuxW9xaV6yweMZ7ZHfWibTvAKVW5r/zzohPQYvF6TZaVDYz7OzNBuuSzhkrZSq"
    "LkTPHXjkYc9nL2ExEBxxnnVNhXHhDr/frzNaJZ1dseWblm+aVor7oMA3XgoJtZK6Ax8BDaClGLMT"
    "2proTIbwolnn53kMGwOBccVFF901ffr0OT26hvuxKEqttY2bkHp1Yuub7xfkJ//z1qUxWZ4Ya/7f"
    "drwaa35rI5s/b9hXmknNPnEzHzVn+jXcrwcvK4Nsmjxl5SUbLlyCFuQgsnb/F+MGhzF+SjZ1si2I"
    "SkrWlJTaW+zsTRJFuMNqJjOjmSXAlLmvsbmZnZh1lB9trHHJS54vTGLXXPH96dSApem8wLa8N46A"
    "8YrChP82tqm3OS+vYKncfPjw+hVs5v3Gw4sk1PSzswCaZRqV5YIDIOIDzuP+90Fr8O0/f+zdh1Xc"
    "N27KY/vmpW/ja1xDHDDh5ftTfhyQaAIfLHVaSUkn5Gx+h1UasyzNRKhZ9NzsDK9sXJL38BcsZsoE"
    "3RB41XSt6U/5/Ro7BKYqWFY201cZ2z1NXjQxxla1rpy8YsXkCuC1alkzqtTZs3MZVE5DmzV7zl1f"
    "ctcOCkvlVjY+GMp7hH2dz92Y3ziPGZS95QnYl2Bepf+OlLH3nLw8Oa0L21FID3ENs8uwHANjwFz4"
    "TJZZ2ZnkRMvsHpwHXvlwHpWSVGDbuMslpU+sOZb0az5mXjPLgoHdS7mD02jz6J+lXc7Gictj5PrZ"
    "Q3A+3Afd+QH/Ok/9mTr46E4x99g80TtviREw1OceW9LoQsVoxW1cbM8lapVBWaO67E8HksYmK8W0"
    "HzPApUxwQEwk8p7pvbpfq/Uq08GGS5/oVIRaw7hqOgu4yQUN05Zy6qE3L730UlAxl7b+qi2W7Mor"
    "aCFjZs3JiicdlfUz/It/pl/r0g05AnwFClOMGPseZopKLfAsmZfGi50ogVd5GbMrtjjJXtI1ZswY"
    "A4A/F7FsrPz+E7CD2yk/1q2h29V1JSA/wbx+DdCq9VaFf9w0lpu8tKmrQF3fEEsebWrqSjY1vdoE"
    "QazgILeY9s52m40L0SxmQ3PSac1htNIEaa4A3849Ms/sxhtNd1QfeTeUxgt/XRhPmbYk+6HzE1Ip"
    "YeAFDsopC5wciI3XU8hRtKfSNRCdJElONnWW7F2j4OkmmlK6coWal7yJ64P0x2pLzRF5VSYYtRiM"
    "OW5bwuMP58yak95MbuffNfqfik88yPox96l5Lgm4Y2OeqXge+XSexcvqrRTb9WTLPvxpnnQiWP4T"
    "m5/bKxXSxfrcJ0AGlzy6uwuUQrIzBfG+MLC7FC4lP5b3RnLs0rSJ4Or761tBoo6ZNWaMQ8NghahI"
    "z5zZc7K26QMm13uyLwVQG0VHJc5b8pmZhTz86TxXqPfG3ZFKTgl32ID8Q+V1opjmT0OvGPslmgoF"
    "u97OoxINxEVBwfRnZcHuWFMBJPbbV7K2ffsKdVmDLDXsmEh7NxGSjoWhIuT+WdlZzc50wlzD9+iI"
    "otdyN7SvT0FNqFNBUKOcsGApJC2wK6m4IAzUCJlSa9Bwn/6wzFcCYddUUvpEiatBjI/pgcCxPU9A"
    "NPvpWFaJULEIYuRnEyVcKXpxHuuC7p89ABXB0J+1naf+0xDPkX6c88xL7wHnMTXx/7YDr88w3Bda"
    "huUT7L4Qz3mGSBZHZgOuc6BqSBNhA2mlP5k1QaE41mbIOptZ6ZouSTZ3Brwpr+DLAxsO/tfi/8pt"
    "2FBQ8Iqcf/HFY9ckzQ/oDkv3m/eNye6F0BP+eqjzsx5XP03XC2hkzB2Lr+XUD/IasbCM71UTnLnA"
    "hctvb0KH0OIpIQu3rF7mNqZcsOBO3LOaSrGjTbvzH11j+CImD3s6Y9KmaRDvIbsuLd3zyiuQou6F"
    "9D6ZhFS1xTb+TaYVOagIzRmw/mHIk9l+wan5mbTAmOblo7i/BVzS824/ezN2wkZC0O03FzKfJhYy"
    "bizNcSmFbMreTS7L3XiTZ/txsncCiSfHunavQRsr3Zs3sa3lMrgjebSrqwlbKiZfNm1vp+yOFHPO"
    "T0dFQNxnDVgXZWGVe+ZfMXfJDHFAK0Q1caPKtU/GnXigKcbW/kVBB1eg0Mv6y0JRTOcGhpYl8vtP"
    "suEzFEWWcSMkRq03tpvlpaVjx17t2XW479bm5ra25lsvbc3LG5uGCtryu6anoXKMLSO4q8UndY70"
    "x9zokJgFF4r7x8AdO9h7ea2N/YuCmoWLpryF/ghcCm7B4DxRyO6O/q/SnGeyzb5xMJylpSUl/47L"
    "/qFdffUrJU0OKiUqsI+21JN07IyFrOyssi4DGmRO6SVcQcgr5uC172EWCdzHkbhwKXYQU5xn5bjm"
    "obLJcbeQ4lkci2E+707oodmsaqIQWdlnsieokICIs3HI+dl8MMdpwoNNwL2FmzLDK2bjBYma0RWu"
    "Szu5BZyRkoymF+Y0rZwYjFw6xZJCTTghzlTKb+bzsa6mY51mO9bUJTuooAnujxXwBRm9numzstnV"
    "ipOfrfw4d8kkb2F249IzDcs4UXBBzQBarmelB62B4Tstf06ZdYfCwuxqxP2ElN+O+yz0u78zpcz4"
    "bAuoG9XMmT681TDac9F506fPdpDNAlb/9lXmwT/LjV6SjVahl2WFNGEdNlVk4wry1ieOmE4p457v"
    "7KoNQeFcdYZR5XA1VupiZ14Yfeod2XAzXBIut5Z0o8EPCjMtasbVIsu0dERl/tjHSsMbPBcxZKgZ"
    "LvrKh26DBv00Cy3TsATzO5spCA6uMoOWlggGvSwqEGOHfNGqTjEpkbrjBIjuGECSVadtS0MJ5zyG"
    "6S/jpLrUGOOuMckxdySZMfFF6aYlLGQ/sg+X58woGH2F5y/n3fUD7quf9guBLm9eYWFGxBLTDSuK"
    "M5ODFrkgfgpfmX2TKN5Cq7FnF9rEhNykcoT+O4S4EwRNbPBCEcfr5FRaMoltQbpp+Rx0tH/ejE89"
    "Dz/njMtwPzuFE8o//B732DzLiYwPWGgZlvmNFbGPpNvnewK7gHnba0StwkI3Lm86ssxwdMcdJxAV"
    "Qtz9gvhKAXeQlONOExlFPi1qUYOdzwoRhd558+bty3vkOfWaazjuZ6e0dgdAb4DAVeh8m4WikGlY"
    "Rtg07yhTiHWKbJCd4OFcm1ehuqzwqbjDrFC0w9GAMmkuXOm0xAzlqYj292K4ADE/WBGLEMYXaZiZ"
    "334hcd68/ANcu3qqC50+5lTTFQuzGVbQLj7UmHLe58BiB1+KzkeKOLVnZFaYYSOGU6WEE+ivuEPD"
    "BoJLOcxPYg4axRXZ7Bb9lmm5QhizxpTrved5VhSfhlVhv+CwduqYh+CyZsewnOCQsN0wQdPdUMzQ"
    "YZSmLI8UXZfvCuJCRk4NQSvOdpCHR6Rw00B/ihmW7i0rKiqLUsO0ABfvimO2afnhdr0Bi4UsB1X/"
    "p6NP03pD9WPsFU0fzDCsmS5lFbURGQHe6BldAT6eqcIIwBIVPH9Pxvo9bwxPGNwKbbtLG2RE9eE3"
    "ttR3tWiRadRMCBdmqGMrakXtD8m+Yhq3nfCxtCVNp7besB2Ma4Zx0V4+LWItcKKr7tgTFUA6JAwV"
    "4bL1VCYr/Ljp8RlzNn+Gd5rfEpiSW3umKfKgO0jG5cy3WWBHLT2tZ1QYrnl5z3Gn89wdNC4P80XD"
    "lTRbWDkpoRWpjF45EDAuq+HC/HmWlxmWVs/bdsEcQc+8dGUgLNFPs1KynlHk0i3E0fJ6xHgnSxKu"
    "MwNr0Y/tClyhOOP1036UzC0c91xev2hcsJ6WQRv1mRr7vijvurCWFvQxHjfIL2QRmLnogmi93XcP"
    "gMVngaWQwVqaZdkcfAk8vi8KPmmZVlnAjO8LeNsfZqjclUNalH9Sy37hFb/Y159hWOwjBH0873Pz"
    "S/iUTEPAgo3sRLaiqG1DA0xGyOKGQqZnRWp4od5XLxgGbESFIq+LeKLM+JBFZRAPrBAb9VnqMGoH"
    "25aOE+wW9JXWSBc/zXETMw3L/EBlZRkViET9QGehrminWUInzu4Vfb56TTBsNJUlZNnhjgZqeM0X"
    "BaOB/g+aUdYO4E9lQTv8UZ/785RpAz4x9gbWo2nboTOx+r5iRbphDfwIoI1nDviW3S0SdCnElNVF"
    "mveW+XPBYkh1X3QBpO1FRelXnQAjifjC3nrbCfVE+kPKNK1sIK2g/QS6Y+WJz3U66dX3fW7Dgu/G"
    "m/kZyupNkkzUDKyYkJoil+SIWzZjBZ2ArSozG4MVLJqZrQX9GT4aHfAIbUGW55U5FTjpiHqi09VO"
    "0g25VemGBZfGh9NwlUX1hDsTq9GiiQUQzZ2Qb0UP1smDxDQUeE1ZGqwsuHS3YBnAwvAovcZwLcuM"
    "isoMIywqCmpaVs5lCZtzy3pu8B34Tw4Wbj7jjlg+I5KGy4rs0OWzcosiDFkBb9AOspbCMD+0VTCh"
    "uu6uo+iWVsS0J+7Oh1ydw8BWJOI+k9DHlGmuDxgUamrqgzOLoHvW3XF2gdtAy6JO6Np4OnvDjewK"
    "fWmiBs3CFyzDFvTW2NcThJtCMC1C8G6ZH81QC2mwUoo13iXbWbPifmvsUUD1hu2rDireoP26lvEG"
    "jF5RC5A0fygSdHf4LyoT7dD1/CH1NMGq4NpYeLF7E1eQwHMT/BEnUqPl1GfEszLW/VcNeLLLOQ09"
    "mya8qGz4pN/tw1FN8ONO1ZS3XzvqKK1omkq3AgC1vroi+B5rFjjlyqBmx4gWj1p8mtxwK3tRi0Ew"
    "kq2zE2wwwoD+B3NF0xuKMnST+bVXZZXpCnil0Q/4izKfbdIqquftfsWMbMEAyfqFGEk0Mau7aahu"
    "O31uqL5hTmlg34phG7Qh4xPZX7ATI5yuHr52s9srUqyu0vi3Pqtz2kogbiqMgElEcwzPjPlR64U1"
    "6/sY8GVaphWsMWvxZWmjw/IydfC0x3NSEUuyx+hBxhifpU9d30Cz5B1RJ7z4amr4RNCyN94dT0gg"
    "kYj6wKOoLxMWjSg1NRHqJNyG0Jrpkv9p306UGv49M2xFggGWZZmW/Wwl6liVvEw9naJUXdFhc6FC"
    "kAl0Cc9BOdTaYGsp053KeDtC8G7hE7UeEKTO6D+E6mAiHRYPaR12+wlfjTX+aGgHUy1VOQSqMl/Z"
    "sukiPs04eZCuwZlZfsOuY7F62kWpusvGZXRXO7gKJuRWdKxf5s5lgtQqIfnSLC5BTYOrSht3DPJR"
    "96NrEo4YmRnVzUpEmnagGeVZgGUOEBZZ35OrD9FBVRQBImHgGKKDSj296Q7OOlzW4o4CB9DJDT8/"
    "4iqVwOc0oVi+QI1wGgy4qNjKIihbiXA9IRmKA8KK8RKpNO0QySz8C3bX4bOJWuGoPohRE2BZ32Va"
    "99KDq0tPfyJt4Frs4HreNrlytcF1LWUKyfArlPIzQRspCTtCa67BM+rcPXCc24iOetLl5tD7Uz1Q"
    "I9ilDqYYisyqtpbRZ5pl1BqrZFrkGsunRw5yQzw66+R3ZnPjoged6L+VptVK9WBmt4ViTPMKdjyx"
    "dWFCtzsqvDzLr0B2FqWFc79bvEajC9IUb70Tv8sigbI01WwFuoARJooSTlinl67ghnzK2FfZxs5x"
    "xiOuzLPPVesGRFZHnjn3wbpYv8WqKOqiDLCsKwtqNTWC5WN+ewZMTdYMr4iNjATsLjecVvM2hyiM"
    "sZ2FCac+L+ExGdwZ3iCRWZeEx3oW22tYJLfIClgffED3beUcVg0g6KPpdRuTt6GFzHkKhmlRl3bI"
    "YFWluywIHu23PVz7sdV3GE5JnYyePr9q2LbeVFsX24ZVzi1jlzzTHojOBcuXWTlJ6yzLaizeWlpf"
    "Z4xD4PdBs8GqNwH4bFnidXJI61YivQekl7YPLayfpn1K3a3NLfoE26/KMkvn4fTLtAKtWaEpC1jy"
    "zOrtvC5vllyyKq0F0n0cHh0YUJUsSwsI9MgK7qTPGDsVy3K75WSjEF7k2EE4TesEAhnJbEYXbsGy"
    "YvMCv3GEk1DlepycXndwyrTWMIWT8dj1sazlQQPVsG4X7N7f9YghHH1sLk3UHZwE05uiQg11KhYz"
    "M3zDilS66yLLyuyk0tBfivt1fXyNbgsy62USToirCbqrMIlAhgOOHCzcQ9+SytFgWfpgGVyVZkSP"
    "BT6a+YVbiYdgRZuasqzFPWagMepyNd4dAa1sr95VFdPhmysyC49mB4gnVZ6ErjojsIq5tyRXKm2W"
    "RJ1am20fCadKYnWHC3LDKnI1NvjxK5d2MAQcDbt1llPBMb4CXfAtgES93m8VZQ8100sPfvXrPD2W"
    "pa7Y2iwNGJn3Fg0slDtFVrN3L/txWq2ZWHrSZ7V6bD7mzM0trh7XlzZjwIxq+oCypGuwRNrKTd3K"
    "ncJBt6fvaOT3F7dlFO70RNmAQqkpDEAXWeIzStJ0YyQ4YGTZbq0trlTdrDsowbTpmdGsiTJrbe2n"
    "evza6YLFPoZxuJ7LtnxlRemsLNkoONMPDCVmWlzUrnIqzjgapbpR4tjV7K476BnlT6YNNDNKDShl"
    "tbyJH/KU9s4/jZZVjNUgdWufRG1JhNW1oDkWFWSj5QFXWLc8yOdEaozTZhpkJeC+IO4Yz9jRrc+7"
    "g7rFw9YSiB0iWlkwYUcpe9iGFUFP9Wye0wfL5rUe3XHXrVa4F3zRqqqoVpM22cWdsjHTitqxxqxK"
    "mdKed88kUhe7y4tKRl2ZOR9N+PiAPW7a2nbrZZe1NS/eeHqu77TCsjTx/7d40WS1jaQvkknr9HS3"
    "hPSlDfnY5cOo8GNrsolBroV7h7qSdcFdyYYOswzyTPd7uYsiJ5nXDBMs8/gz7HEWZwmy9WUz0S3N"
    "+GulchE7lmmuFRHp4xzQFnNvuWoLM53ij78sUeXT+IzJgJQ2bDVJFZ+dlmXzAlqH2gbgwnlCmq/K"
    "l5bjgEZy1/7IgCkcRqpCV3Db3dUcr2O0kYFv1NLXsZLjzqJzd07YDvW1ZB0HtO4EbVHEpChNSxUz"
    "aJkj2X1mokCDOBo9c4HIRL2U5dWltmWrnKDw9YDFzp7dtri57dY3mvt2ZaOmK2ymVZmgs6lW1lIW"
    "wrsmDBUljEjOTsdkA5cgUZWAKRwuW5X5upSuV09fmBpGy3J9s6uknFMccfYpO/NYc2aXCdEgTp8I"
    "Bq0Zl3Qrvlpz5nMvU99xB6m2vnc6Jp/GKDWcsDjryEKOaz/05uLmBmnQWbTpVlfj5/019oyN9SAp"
    "y7lLB4DeyvXhaumW5iO7Dq1UBxSPvl6wBiSRb7a2tZwEM1MxsfPq8KSfDFY9Uyu4dzpcq0jKz+CJ"
    "0sMHq8I5ypBT2zveAWbZkUnygLsaNrLR4mJuq+NyLQ3PH35z47U2pvLyijN9euZwW1axCxmnrurY"
    "dWnzG8ZxHa4wdCStD6Ut76hWCNwuU+nW5r7FHRtdrlZ+EidYfZ1g2ZG/wjWxYMUqT8eb71x6uLXv"
    "+VvbFjW04fntzabZgV5a75rbom5Lo1RRMayfeoRgZbMz8z5VNVLMA1OvvPLKqeyw1gFIwOPUEfi8"
    "IwrLDmftFRDRysuzao50VCND6WyCld1Ti7nTUFX5hsA6C9s5WOdgnYN1DtY5WOdgnWvnYJ2DdQ7W"
    "OVjnYJ2Dda79TxSl6dyhSEB/AAAAAElFTkSuQmCC"
)

_CLAUDE_LOGO_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAeAAAAB6CAYAAABnXoAEAAC1cUlEQVR42ux9d3yd1X3+8z3nvHdo"
    "72V52/I2w2YYDEiMEkJI2iRSFmm6gpvRJm2atr827dVNxy9tkrZJf01iMiBpAonEHgZjQAJvW5Kn"
    "PCTbGtbe8673Pef7++Pea66FDQZsMO198iGSr6R733He83zn8wWSSCKJJJJIIokkkkgiiSSSSCKJ"
    "JJJIIokkkkgiiSSSSCKJJJJIIokkkkgiiSSSSCKJJJJIIokkkkgiiSSSSCKJJJJIIokkkkgiiSSS"
    "SCKJJJJIIonLC+zzCeYayTU1krlGJq9IEkkkkUQSSVxi1NXVqdcRck2cjJkuKfEzC2YWybuQRBJJ"
    "JJHE/w6Pl1kyM3FN1Nvt2LNnQfBY4we6G16+c+hI0x0zveNLdRwNGzdaANDQsNGqqUl63kkkkUQS"
    "SVx60Hv54cxMqK0V9IlP6LGtL35renjgXkvpWZCAcqcg4ujNtuU9wZk5/3fOVTd0XyrPu6Kiwhlk"
    "Ts8nmgSAmpoaWVVVpZPLI4kkkkgiif+xBExEPPTq5h97xob+qLO5EcHpYceyAGEs8ubkSyc9Bxnz"
    "FodEVs5nC6666VEARETmYhBveXk5A6De7Vv+VgemP1KSn/v8BFFD9pU3PcZcI4mSJJxEEkkkkcT/"
    "LAKmmpoakV9ZSQt2v3R/+ujg73fueCXiDY0rKcNCGBsqIhCULj3u8rI1e6GaddV14Lz8xQWLrj6J"
    "KAvzO/K8AWpshJw/9tzDFJr6WFfrQRQWFUAVlcBOS/100TV3PczMioic5DJJIokkkkjiYkO9J56v"
    "z0dUVaX79+9c7w1M/n53027tdiZdpByACGQsSEFQFqSLQtC9nYazS4Q74lpIi+lELB/8tgiYa2ok"
    "Eenp5oZrV04PbBw8efCKidOtNgcn5VDvKcedWyQLr7n+l8MHt7Sjuno319RISoajk0giiSSSuMh4"
    "byp/y8sFAEwND/71ZEeHxsQISx2ChgNNAhoWbJJwBKCEgWUH2Rno5/Dk+De5ocGqfyfHXVlpGhoa"
    "rPGR4W9MDfVcMdjWbHvsCStVauFlx2UPDWD01CkxPjT9N6iuFqisNHiPQ/VJJJFEEkkkCfiioL6+"
    "HgAwOT5iTwfGJWsbgg0kCMyx/0gChiEZUDByqP+0dsG5up/suyr8fs3naFu6EBARr1mzhqdHR24+"
    "3X7CEGvFILAEHDIggujv6+VgMHALAHcs1M3JpZJEEkkkkcTFxHsSgkY5AD+QnppiPB4XpglQguCA"
    "YJhgiOAQAQYgaFhCgE1QTPS2C5GaUs3MzwB4y2FhZpZEpHv21n/ZFZlyO2PDJp1YSeMgbDQkuUAM"
    "WCBSjjOZXB5JXGzU1NTIBQsWCABYs2aNvhgFhUkkcTmgrq5OpaenU2xtO++kTifpAV9S/i0HAKSl"
    "Z2SmpKSAjQEzYAyDmACS0GAIISCMARkHwoQoMNTNmBhdMH76dNbburn19cTMkoKTyzkw5ubQBFsw"
    "IBiQIJAgsNFwKQm3tJL9wEm8YzCzqKurU75YH3tVVZVeu3atvXbtWpuIjO8S9rcnkcSlXtcx3QQC"
    "gIqKCidhbSfJ93L1gGsHf8AAIKX1ZNBQhSMtIWDAYLBhSAIMCYAZCgpgAyk02ZMjdmp4Kj148vDX"
    "mPnvGu+/X67dsMG+oAXj8wmqqHDGTzQtJjZ/MNbXzSmSLDI2GAJCEIxjIIQFJgJFs770Nhdn7O+q"
    "CVhB9fXNVD64glFZaZIL89x4O0Tk9/vNZX4eIlZFbwCAiLB///6bLMu60RiDQCDw0rXXXrvX5/OJ"
    "y/FckkjiXPtadXW1TFzX8bW9a9euP5ISeVobXHvt9T8kovHkVbsMCbiystYAgKew+JFQ9+nvO+40"
    "kGNAZCA0A+yARTQEDRYwxGA2ICeseo4dpjTyfLnn+PF/X7thw1C8l/hCPzscDgOhgJSRECsT5W5D"
    "Fpg1CACTgLRckB63C0C8BYnwFvLACcfD51rE/9tJOH4NfD6fKCkpkQ899BD7/f633O7l8/lUb28v"
    "fepTn2IAKC8v17Fr+5bu10UAAcDGjRvVhrMNQnPs2LHVWuvSzs6ulQMDfX/88MMPl6SkpLi11hBC"
    "TD/99NN3fehDH9oGIEnCSVy2z+v999+viCi+tp0TJ04stixraWdnpzM9PfEHXV1da5577tn5zNGS"
    "mUOHDv/pc88997t33XXXi0lho8uMgInAXFMjh4I0xVbKE7ml83+7v+2Q4yZWFhEkDJzYPsqQ0CAY"
    "4UCxodDEqLECYxnuidP/zJ2dX6mtrYrELLEL23DDgAlHoNhAMsAQsAUBhqBiBOxOSwW5PX2JFt6F"
    "W4iERjSoNYNZ83o7T38kJ939W1OT49pyp2HEUV8govb/ja1NNTU1srm5mfx+/5ncUIxwDAAEg8H5"
    "4XD4LC84EonI6elp4XK5wMxsWZbjdrvhdrvh8XgcIuoAgPvvv/91twJ4TV+8oqLiUvRyk8/nkzHD"
    "gQFgw4YN9tTUVLGUMmXHjh23joyMfPSJJ55YFwwGM9vbT6GtrR29vb2Ynp6yQ6Gwc8stt6ReeeWV"
    "f0ZEWzdu3Kje6npLIolLgehzU4+KiqhRHHtebWZOnZiYmLV583N/+sorL38qEAjmdHZ2orOzE93d"
    "3Th9+rSORCLGspRz772/W+xyudMAoLm5OdlFcjkRMADUAqhatmxyeu8r/9dxnLtG+juVExpmNpri"
    "XEqwwBDQggDBEI5BqiAx3tNpL1i94vP9033PVFXVPhWXk7yQz404WphIBGQ0BBhMAhoSFgOSGQbQ"
    "aZmZaiI8/Q9EFLlQsmRmQnU1kR9mbHdo48hk8+8H+7owrhgMG5SeBU9azgu9x/bfSUuvbPvf4AnH"
    "w7FHjhyhuAXMzBYAN4BwS0vL9X19fR9pampy//znP//y1NQUgsEgwuEwwuEwHMeB1hpEBCEEpJRQ"
    "SiEtLQ0ej0d/4xvf+EFOTk6ksLBQzJs3DzfccMN3AAwCcCmlphPWBFVWVoqamhp+p0VPzExVVVWi"
    "trZW+/1+h5klAAuA+5lnnvnzmpqaL/b09OS1t7djdHQUJ06cQH9/nxMITJHjaFJKEkCWbUcoHA6y"
    "bUeSxX5JXDbPq9/v5/hzo5SCbdvWyZMnV7e0tHzin//5nysmJyfXdna24/Tp0+jq6nImJiYQCoWI"
    "mcmyLGmMkW63mxzHMV5vWlLE6N0i4Ng0IXGhylFVVVWamQUR7Rmr2/x83oJlH+k/1qQNJiTBQBkJ"
    "YoKBhGAAhmDYgISD0PQYTXR2GiqaHWafT9Sj/sIPNDgVcKanYWwNhxgkGNFqKwMGASSYpcVGugkA"
    "6vPzKXGBVldX80ziZCBGvn5zetuzD7pHOj53cs8eHZmaoAAcZmHgQOqCJSsXZ0jP5pHGxjuqq6tP"
    "s89H9D8v7EixkJVMDMfu2bPn9q6urjk///nPvzExMZE5Njamp6ens8PhsDp9uhMnTpzE2NgYB4MB"
    "BIMhaO2AYon4aFiLwMwgIqSkeJGdnSPnzJn9J1lZ2cjLy0VLy3E899ymP/R6vaHi4hL5wAM/2+9y"
    "qQcWLFgorr/+hv+ura3VRASfz6fKy8sTw9VvLXgT/RvNzK7a2trP/PrXD38jEAhk9PT00MjISG5T"
    "UxOOHz/Ok5OTzMxGCCGVUkoICZfrtbo+y3IhJSWdXC4rWYSVxHuOmENgiAh79+76vePHj9ta4w/v"
    "v/9Hqzo7T6dMTEykbN26FSdPnmKtbUNEwuVyKWMMpJQQQoCIEP+31+sRSqmk5/tuEXBjY6MsKChQ"
    "zPxWNjZu2LjRGpg1/xPZkXBTWn//knD/tLGghTQEAwJgIA0AYkAybKFhnDBF+vpFxJX5dfL7N/vg"
    "e/NPKi8X8PuN17L+WoAQNNoYQVIIhtAOiAwMKTgAUrNzyc7K9gKI1WtHUV1dfc6cLmprBVVV6cFX"
    "n3pADJ763IkDOx0r5CgXBIQUMGxDQMjJowciWRk5i4P55k+qq6u/Xl9dLf8nhR19Pp9KCDGblpaW"
    "0iNHDv/ByMjwB55//tlrR0ZG5aFDhzAwMICxsTHEyNYWQpAQQlGMcYUApLRixBst8IjnlgCC49gY"
    "HOxHf3+PY0z0l4gIXq83IyUlJSMvLw/FxcW3LV68+Lb+/gH8+Mcbv5qTk30yN7fgH2+++eaDfr//"
    "zPFWV1df0HqNRywef/zxrOnp6V9+61vfmj0yMrL6yJHDGB4eRn9/P4aHhx0ikkpJsixFxhghRFS0"
    "LX4ucZPNslxITU2FZbmTu1AS77nnS0TmhRde+ODw8MA/1tXVXxUPLXd0dGBgYACTk5M2EUkiCJfL"
    "JZkZxpgZz2f0eyKClAoy2Ufy7hHw2rVrbQD2W/kbIuIYYdtDW1/6btbceT/tnewxkdAERLQRCWAJ"
    "AoMIYE2QzHCzERNdnTq9oPTm7qatd8y6+qYtb5boj3vJjtHLWTswtsNKEYijHhYJAaPZuDPT5XQw"
    "0upN8dbXVFZK1NcnVvrN3Kjjnq8ebXj55+ju/N3OxkO2YrZgom1UIMDAQLKBhLDaW48alyv1S+3t"
    "7d+p8Pv7/ieEouOh5nhIdtu2bR8JhUJ3Pffcpg82NTWWNDY2YmRkBKFQyAag4g8vM8OyLAvAmYc5"
    "7vHGvd34v6Pf05mfxS6/InrtoQ+HwxwKhTA2NoYTJ07w9u3btZSS8vPzr169evXVV199deUjj9T+"
    "KiMj/egdd1z1PaLCKb/fH89P8/mKoJiZqqurqa2tLXPr1q11jz/++JVNTU2YmpqyAaO01ogVVSki"
    "gm07Z47pNQ8+8f0AIQRSUlIgpUzmfZN4T+H3+9HQ0GBt2/bq91966cWFO3bsDAMQzCyNMaS1AQBL"
    "CAG+wJ3K5bLg8XiSHvC7RcBAtNWnGkB11F3kCyMWYq6pkVhaUjuK4NcyRxcsHW87aixMiZjDEw1B"
    "Q4AhIA2QQiB7ehJ6sNdy0tL/mplfrq2tRSykfc4NrXxwBQOAtu2wtiMQAAj0WuWWYThM7EpNlxHm"
    "zryytSfrfD5Ffv95Sb2urk6ivJx6fmvtf+ueU5/u2dsUSbWNywgBJgnDAMNET0BrSAhyJkad9HDA"
    "RQMdfwXgzxobG9VbNVwup7BVdXV1vBAJmzdv/uSDDz74p8eOHVvX0tKCY8eOoq+vV8cISgghLACQ"
    "MdM4bjknkm389cSvcbJO/J2opyxmkhslvEZaa2HbNjo6OnR/fz/27NkjFi9e/JlFixbh0KHDf/rM"
    "M099++677/khEU0DbzyG0u/3m/vuuy/D6/Ve2draavf390uPx2MZo0H02jnFjzORfM917EQClmWB"
    "iNKT21AS7xXi9TNCqN+3LPecY8eOByORiDe+XpVSZz2Lr4/m4HXPctSwdsHlciUv8LtBwD6fT1Sv"
    "WEHxYiV/dMdCw8aN1pr77ntDRRQicF1dM1XkV02OH9z+3bzFq3462t9r68CUUGAQx8U5BISwYIyG"
    "NBpuMnL89AnOzsu7tX/vixWVlZUvtba2ugCE3/BgHYcIDEEA2IAhQCRBBBgp4bjcSM3IyfP5fGJw"
    "xZHzHjcfPuyilSsjgw2b/zE7Mvbp403bIx5Hu6QWMOxELQdmaI62VwEMBQ2pHVKRkCBtZwHAmsnJ"
    "96X3Gw9bAXCOHz++rLW19Vdbtmy5qrGxEW1tbTw8PKKNcaSUQkZDUtE80UxSSiAno7XW5yKvePg5"
    "FtpSsYsLY0ysX/vcZBf/PKWUdBwHAwMD6O/vdw4cOEBFRUUFq1ev+nZHR+eXdu3a8f158xY8VFRU"
    "1H8uEo61TKmSkpKeSCTyo6uvvvqPu7q6IrG3FjPDcTO93pnGRczzJymlUUrtBICenp5kf3gS7zrK"
    "y8s1AKGUeEwp9c0lS5aUDAwMhI0xUimltNZnPVsz13Ni1Cr+zCml4PV64XYn0yuXnIDjXqcfAI+2"
    "zweYkVVEgKeXiELYsCEqgvEGxUYVFX6Ha2okrb7xZz2vbF5duGjFVwYO9DsuCCViXqphhhEElgIw"
    "BoCGy55y3FMjEtN5XySiFw8frnnTTSza2MQQxCBiEMeaRplhlCSRnmkcKf7D7/efV6WIfT5FK1dG"
    "Tjzz+Brd2fkHXSeaHFdoUkmyEBYCJAjCRI0GQIMBMAHaGEiXhMvlcqTy9L9fvV4igt/vN0ePHi05"
    "duzoXzz88K8+dfz48aJt27Y7Q0ODQilLSCmUlGpGCBlICDcbrTU8Hq/IyspCaWmpyMvLFykpXmRk"
    "ZMDtdkMIAdu2MT0dQCAQwPT0FLq7u9Hb24fJyQk4jnYsyyUBJiIBYxwk7g1xUox/jVnzKhAIoKWl"
    "hdva2pzW1hPzWltb/m3p0uV/+vTTT3/vnnvu+Y9zkXA8V/zSSy995wMf+MC65cuXX3Ho0CHs2rUD"
    "g4ODhpkhpRTnCjmfa7MSgkgpBSllCABWrFiRJOAk3nXE0oBERKObN2/6+6qqqn+6+uqrC5ubm7Fz"
    "506MjY050SLCaGRJKQXHcc5JwnFIKWFZ6owHXF5ejnjdRRIXkYDjXlDHwT0L0kJTX+2sr/+ycWyk"
    "pKfTRCj00uj257Z55857mEqXHWf2CeANwtKVlcw1NbIrPeO/rGDeZ9LzSrMCA33GAxKCHTARNABD"
    "FBXQJIZgW04PdCE9J3/xWMO2hZmDdseb5VS1diCMgWADMJ+RumIGjLSQWlAkAl737tim+7qFEzM4"
    "nK76+qu8E/2bxk61FthDw0aRFloQbKWgmCFiwW1BEgYSmjVYWqzSMlRIuey0/NJvxVbn+6YfOPHa"
    "7tmz7cM7d25/aOfOnak7duxAR0eHISLldrvPIqBET9dxHAOAMzIyePny5Wr27NnIzy/Q8+cvlFo7"
    "e6enpzdlZma60tLSjNfrhVIKkUgEk5OTmJycNoHAlCs3N/erY2Nj7r6+PoyPj6umpia0t7c7oVCY"
    "LEuKKBm/Fg5LDFPHv0opQVFYhw4d4uPHW5ybbrp53kc+8pGvHT169Ilf//rXnTPVqeLnfdttt52U"
    "Ul65b9++zyxbtuyP16y5at3Ro0fl7t270dPTA8e5sM4LY4xxu93KGHMrgP8cHR2NWmtJJPHuk7CJ"
    "Pds/Zeaal19+8aulpbPL1qxZc+/Ro0fV7t27MTw8DGMM4h7xGxmXzElb8pITcLxsPdh2dF5ktG+L"
    "09+5YKJhh2HHEcMEZObn3ebYgdsCOvKV3j0v/Yzotq9xXbliZnOuPC0RmZqaGll19brW3l0v1eSW"
    "XfnF0bEd2nJCsBwHSgAOx4OPBECAyIiJgT6Tmj+0MuzJfZhuv+daPp98ZH4z+Xw+AScC49gAMwQD"
    "RgDMBKM1u9PTyVZWm3ClRWKyazzD8xUg4oE99VfR8NCLU21Hc8IDvdotpTRgaOZYuNnE5hrHmJ0E"
    "IN0I2YI96VmkU9IPBFyuyGuSlZc/amJzlA8ePJjd19f36EsvvXLzpk2b5L59+xwikpZlibNDxjjT"
    "lmCM4ZycHF64cKEoLS3F8uXLMWvWrONZWVkPrFlzzUOzZs0iAP1EFL6Adfc9AKK3tzezoaGheuXK"
    "lZXd3d3q1KlTaGk5jq6u08ZxHBEPeZ9rM9Ban7HmpZRkjLaampoid9xxR6mUcrHf72+Padye0+j0"
    "+/28evXqXzFzLRAsfvzx579YWjrr3h07dhYfOHCAA4GAeCMPYcZmNZHchpK4HDzhmEM1CeAfAGBq"
    "auo7zz23qXrevLm//eqrW3H8+HEOBAJ0vucqiXeRgOvrq2VdXR3Guk99PXN6dEFHw7agsgNeBSCF"
    "GU7vad032GeQW5hVcsXaP+/d9QLo+oqvAQDX1Sk6h3BGVVWV5oYGC2vW/Gl//TNLC8qWl48ea9IC"
    "WkpNsISENgKCBSQDRjqwyBGDJ0/YhSl5qwa313+xdTzwUy4oMJgxjaMe0UKazz36QNhtNMAa8bYW"
    "gEDG6KycbBWBfnLuqrUnZx5jrN2IJntack3z4ZcDJ/ZnBbpPaTcZ6bCANhKCDFyC4ZABYt3FgkU0"
    "fA6GsLxORlEp26lp/zBr1qxAnc+nKt6GBON74PkKItLMnPb8888//+KLW6598sknzMjICFQsrptY"
    "KCWljBMvvF6vmTVrlli/fj2tWLH8SGFhyd9fddWq0blz8/YTZY0kfs59991nFRcXc0lJCZWVlZ31"
    "hD/88MNUXFzMRNQbe6kbQFVzc/N1U1NTacePH7/zxInW32tqasxvamoy09PTYmZeNmGzOSs0Hcvh"
    "isnJSRMKhSJvdC3iXnFdXZ0iogiADgB/deJE80+Ushq7urpSp6amWClFF7JJEVGyDziJywLxtb1x"
    "40Zrw4YNOi0t7QCA32lo2PMXHo/326dPn9bhcFgldiwk8R4QcMz7ddraKotG9zXeN3CsgV3haa+W"
    "DEOA0gyLHWlpyNDwILfv2mYKliz/88iuzXND3tTv0RXrt0ajv/z61p7YiLb+Q7v/VYUDt3JHqoGJ"
    "AE7kTMiYTJSAWQhISdCRkBppP2VlZGb9V95dN/+qvv7AdPkMAqmurjYjRw+s0l0nrhiemjREJBIr"
    "+0hKsLRgpXg9zEyIzSw+g1iv78CWZ74W6enMmuhu1xnSSM0GGgIQCi42sGwbWkXz1uKMAy1gQ+p5"
    "S1a5TG5RU+H633oqXsT1PiFf09XVlfvYY49t2bZt21WPPvqoPTExZkWLoM4uMCIiaK0hpUR6eoa+"
    "5pq18p577rHnzVvwnfLy8r+Z6VVXVlaa6upqigmc2Bey9l4zAuvlihUrdsf++VJXV9d/7t276+lF"
    "ixZdsWXLFt3b2yuMMfT6NqCzQ9Px72PnekFRiYqKCifengRALFy4vN3tftZ2u92Ck+5BEu9jxAV0"
    "ampqXM3NzcblUvtTU1PORLSSuAw8YABIpTRylFvBaLicMLRygTmWyFIKDAOFCHmmR+V4004nb3Lo"
    "Y6PZxR8YPLz9r4aGyn6Rn0+TM3O2RGTY5xO0+vrNHS8+9UpO0aJbJk8e0kY6UoswNAgWpUBE3NAG"
    "IHJg6QCcsS5jTcyN8J7D6yoqKp6fsVELv9/vfOljty/LscScnmDQkcqlmDWkYQAGjuWCJ6MAYZXu"
    "IiLmurrXNuyYHGVf/RPf9kx2/UVPW7PjYih2GFIQmDRABjAcnazkCLA0MMIGpEDYuE1qwSLjFC3u"
    "Sc0v+X1mFo2NjZf9Jh0n3927d+e++mr9c08//fRVdXV1Tjgcfh35JpKalJJLS0v5s5/9Xblw4aKj"
    "V1xxxUcWLVrU6vNBHDlSSZWVlaiMToaKJ5T4Qos0ZhhsTjxUXFtbi9LS0tOHDx++Nj0969MLFix6"
    "4JFHHtFHjjQLx7EpMRc9M4TGzAiFAgiFgvB4PG8pbBc/ji9/+ZPpRIreowmfSSRx0ZGfn2/8fr9T"
    "WfnRlESTMlF4I4n3gICJiGtqamTn0NBQgRDfS83J+0pgtN8hGIVYYJcp6gUSGahoy49qP37ECeUM"
    "p+bm5/w/Oxz8YEtLy0cBRF4Xkq6u5sMrVrjmLL7irtFA4Lgz2l8yPRw0kqKNpMQGhjgqmQSCZENO"
    "OGTCY0Mek5bzf7i7+1U0NtpgdkDE5YODDAAmHApPT4xxtLJagGAgmBEhw5SaIscjznSqO/vbMdY2"
    "AMANDRatXWv3bX/5vozg4F+07GuIqEjEZeG1oh4mRqxTCg7H+n81Q0lCWBOQmsEZi5dZYynZ9+Ss"
    "XnMwlke8rHt/4zn+kZGRzK1bX9m8adOmNVu2bLFDoZCllDpvK4IQghctWsT33nsv1qy55r6VK2c/"
    "kpU1d/S1oqZa1NbWXrTjTKxWjh1zBMCD27ZtQ0ZGxgM///mD+sCB/YKZKTHkPBO2bSMSCV9wEdVM"
    "5OVlJV2DJP5HQgiRXNuX6tq+feuomdauXWuTQrtxe01AiJjPaWLE5kDAAGxgjAZLAVZKqfEhPvXy"
    "c2F9uv2DaT2tzzQDFlVUOJxQ9EJEvKK52aE5c4Kclf397EVLZER6GVpBaYIwDlg5UdIzUaF+l4Ts"
    "OdXqiMjUzf09p26ktWvtmtpaAUQHPwCAHQpLO2wTmSjxEkc1MhiAdrvIuFyRvFWr2gCAvvlNwzU1"
    "ktautU81vjQXw71f72lq0GIioNzaQBoHJKNFXFoAhgBHAJoYLA2EIDBbgExzytaskyYn68fz168/"
    "zHV16nIfO8fMVF9fLzs7O2dt2bL51WeeeWbN5s2bHa21JYR4XRgq7lUaY8ySJUv43nvvFatWrfzd"
    "9evX/zgra+4oM9O7cc7xUYQbN2601q9f/+DixYt/93Of+5xcuXKlcblcb/j5tm3DtiN4bQJlEkkk"
    "kcRlSsDlMZVkKyMnJXP2XGF7040DAY6rGDEg2EAiOuyAtQbBIM3YlDI17j7dsMO2u9tuz335qU1D"
    "jVs/SlVV+nBNjYuZiZkJ5eWipqZGRooLfmFnF57KLJ4HwS4jHIKgaHjXsIYAgQAoMvCYkBjraNUi"
    "NP3jrv17rqpsbmZO6OW17dA0GzumqWViHcYUnTksXYA3PQ1jY14AqHv5ZUVVVbq5cfvcjLC9WQx2"
    "LTKDvZRKEMI4MOzAkIEmcybT+5pPqKGJESTlFCxbrez0nJ+W3HjnfaivJ5SXX/bWZG1trVVRUeEc"
    "OnTgW42Njasff/zxiG3bKlHlKbHaOSZ4wYsXLxYf//jHxdq1az596613/Grjxo3WeyC1yRs2bLAb"
    "GhqsG2+88b+XLi379Kc//WlZUlJCjuPwuTz3eG7LcZJdQEkkkcT7gIBRUaGZWRQuX/zDcWFtzl2w"
    "3B1iYdusYIQCKDrFSOgYCbOBMAZsGJIAtw5afc279PSJptu8EwOPjry66Qsrq6oisc1aoLycKwEU"
    "L7pyYNLl3pg9r0xqlWEMLEAQNGwgXurEDDIabrZFsL9Le0Jjc1128Gry+03jPSWyubmZiQjScl9t"
    "whEIZhIxSuAombDHkw7h8uzvCWRFGho2WuWDgzzcemh2bmjqxem2liUDJ5sdhZAgaDAZGEUwMio1"
    "KZkhDaAMQTIgWCJI0rhmz5OuRUumI3l5/1pTUyNRXh6fKHDZwufziaqqqsgLLzz32b179/7OE088"
    "YTuOY8WrnOOI54Bj5MVer9f50Ifunl60aOGnbrzx5odjlZT2e6VzvXbtWruurk5df/2ND8+bN+/e"
    "O+64I1hUVOQkXv+Z4iBaazhJBziJJJK43AmYAEZ9vaCsuaMie843Mhcs781buNIKuzIQhIdtUmBY"
    "IJIQmqE4OjTVkQohqUDQyBYhGe44rLt2v8Ch060/GNu25Zme5ua5RKRra2tBVVWa6+rUghtu/9dJ"
    "uH+VPW+xChnhGKFAAGRMnN8QRQceQMNtB9RAy2FDUyOfnz56tGTtmg2O3+83xhiSLuvPJsfHIIxD"
    "ItZNbBjQzDojPYdTSX5v1iwKTE66JCorFfq7torO9kWDzc1acUgZsmMhZ4ZDsWIzBoRhCAcQRoAg"
    "YcPNacXzRMayVfZ0Wvo9Bcuva6msbGYi0pfz4IWamhrp9/vNvn37yk+davv5c889lzIwMKDi6leJ"
    "iGsfG2PgcrnMhz98j7VixYptv/3bH/u1z+dzJY4ifK9QUVHhbNy40frABz74q/Ly8h1333235Xa7"
    "ddzjTVToeq3CM8nASSSRxOXuAQOgWCtG8cS1+3XRoqVpC1a+WLrmBnBuCQVlKkeEgiYBkIz1Wmqw"
    "VNAiWiFNTgReE5E0NoDhI/t1pOvk3VZf54uhI/vLqqqqdE1NjcTgIDNA6cXF/yryiqaRmUMOSxYs"
    "QWwAEDQRNAkQM7wwItDdqb2h6esmxrurGUzs80WLwyKRKR0OAVqDwNH/CYJNAu7UdFKZGQXMTOXl"
    "q/TAzmfvcw/3zx1ubnbSjC1hHGhJMBTN+XJMEERCQRgBAQFDEkEmpsw8J2vximGTmXN78ar1dXU+"
    "nyK6/PO++fn5xMypra3H/+a5556jQ4cOaZxD2GRGH61ZvXq1uO22249ee+2639u4caNVXV192RSY"
    "3Xfffc7GjRutwsLizy1duuzYypUrhdbaSClfN0ItEolgenraBQDNzc3JMs8kkkji8iXg2ObFqCCd"
    "V1Y2kX3rPXemzC77rVlXXTORNnsBhZTLDpOK9sUqARYMsA1hItExgFBgVpAsSdlh2bFvtzN6dM+i"
    "UPvRpuFtz/+/eIVr48aNKmfNuoOR9LSnshcukhFWxmIXLBCYDDQRHFJgCFjGwLIj1qlDTY4ITnxm"
    "oLVxfmyikWDbFiZigxCVoWQADjN709NFUMpRSs84TEQ8tG/qSk8w8P2uQ40ORaakhImOkIM8U7hF"
    "QCzcDBALsLQQlhIBl4oUrF1rjadlf33W2tu31tTUuN4PYhuNjY2qoqLCqa9/+e9aW1vv2LNnjy2E"
    "UOeqGE7MBWdlZfEtt9xChYXF/zZ//vy+srIevpy8fCLi7Oxss2bNmp7Fixd+/5prrhFCRFt144Ih"
    "8XO07QgmJgIquS0kkUQS7wsCBuJzf6IzDdKvXrfFk1dwT/HKVcGFV15taZdyQqyNBsMBQxkbHuNA"
    "QUGTBxFKhQM3BAmkCltFeo6ZzgOvpDrjQ1/qqX/h/1FVlV67YYN9uKbGlTq74F/c+YXalZIBo4lh"
    "GCCGIRnViIaABOBmAz05TuHRgRRneMyHWN4vEg4DxsTGG0aPWDObrLxcGSB9Kqts9Zb+fk6zxyb/"
    "Zfj4EeOMjwiCTYYNhHBBasAygDICUhOkiQ1dQNSLDisy89dc4Q5lZDyTM2vpY1xXpyorKy/7UYPM"
    "LE6dOmW2b99+25EjR/+wtrbWCQaD6nztRtHXGMxsFi5cKBcuXNR28803/8Ln84lbb/3mZWdsfOIT"
    "n9DMLG6//c6fzZ8/r/2aa66RQggTz2O/1iMMBIPB8eS2kEQSSbxvCDhOwjEBC5VxxY2vitysm10L"
    "lt8/p/yDShQtECF2GQUXSFjQUoANYGlAGRObzashyYaCLXhimE/v3arReeRL4RefeHrwlZduXVlV"
    "Fcl+6LlDk3Dd652/kIxLRIf+kYBkDckaII6WZkkBRMIwvV3sHh1cwQ0N1hjG0lU4QpatAQk4MkYk"
    "LOCkZUDk59nM7BHHn3/JPdBTEew5zS5lhBE2NEUAQSCWMX1nCSOiHrdkA8kOQKQzZi02YW/JMyU5"
    "ZR/PWbhwHBUVl3XON47a2lqqqqrS3d1d39q69dW8jo4OiitInbvhnsBMMaWra528vPx/rK6udsrL"
    "y8Xl2KDPzKivrxfV1dX26tVX/fOtt97quN3umSMHKRQKQSn6IwDw+/3JEHQSSSRxSXHRw21UUeHU"
    "1NTI1GXrGwA09B/YdbJgpfWHJqu9bPD44YjRQZcUEi5BIGPAMDAUc1E5WsiUAkGR0JQcO7bPcQXG"
    "P+QUzr/jdP2Ld1H57XXMvKm/7nERGjgNe6g/SoAwINYAMxwhYQRBMstAd7dTNHvh6pPD/demt+vF"
    "hdnZszsCAVsIYdmCIR2GZEHGcsGdkpo5uPulZ9xjQ9d2HmhwXDqoQA5YRvPMrCUYIpb7Rex7DQcG"
    "kIqNJ4WQPWtaf+ATVUQU9r3JCMbLBT6fTzQ3N/POnTvnbdv2as6ePXu0bds0c8D8DEqD1sYsXbrc"
    "KitbcjIvL++//X4/+y/jUHtFRYXj8/nE9ddf//P9+/f97dy5c+ccPXrUCCHigxsoEomAme/43/Lw"
    "n2/cZiJiEqGXozFCPp+PqqurOS5lmiBp+q4fb1yWtLq6mmPXjS6Xa5d4bDHp1DP39nI6zkt13uda"
    "05fLeV+SfFdVVZVmZoHaWqIrrv9Xbjnw0KjbXVdSVLCoZ99ejckxhI0thUSspciADEGygKWjXbqk"
    "CDAR1Xliv22N9btT5654/vQLT7wwsWf7P6elZX5XlV35ta6xV2xPJGCxENEgOHS0sMpoSCJox5an"
    "T50i14rc71jB8JHg5BgZAcEU4wk2UFIKCoaQFYksmxoaXtZ2cJ9xG0cpNmAykLFhviqmbOlIwJAG"
    "wYEhRthyQVvpumD5lUqXzDpVjFbD7BOXe9FVHCtWrKCqqiq9bNmSP25vb18wOjrqCCFkXHDj3AQM"
    "CCF46dIlZtmyZd9Yu3atHZ8SdLmfLxFFnnzysYNLly6d09rayuFwVP3KcRxEImFMTwdG/6eSbUlJ"
    "iSwrK+OWlhbasGGDvhBxFL/fj40bN1rx4Rj19fXm3RaSYWZqbGxUjY2NSDh+x+/3m5iEKc/4eslQ"
    "V1enWlpazhoW0tLSQjEd88TP5/doBi7V1dXJNzg2Try38dd9Pp+KD0GpqKjQuMzbJS9kjcfP+1xr"
    "On7e8bXd0tJC9913n/NuE7K6hBudAYDDNTUuKruiq/P48fI8OXVvSWrqt8IdJ9F1/KixdJhcxiGL"
    "GQQCxQqjmAiSDYgcpLvZCo138cTRSVfO7OUfmsye9YG8OXObTVo2VHqG5JFgQmOniA5rYEBFFalo"
    "YrAX83Toem8kfH3v0ACgWBrDIGFAJKCYwZOTGDm8n0d7+tianhQeOCB24MCAWIA4mus1RLDJgNiG"
    "hIERbgSl1xQsu0J55pXtE6WzP9DY+JJZs6aagffFAGqKGUsFP/rRD7786quvGmaW8YEKMaI9158h"
    "OzsLK1asECUlJR1xIn8/GBvMTLt3b/OvX7/+niNHjmBwcJDdbjeFw2FkZmaBiHL+p5BuTU2NbG5u"
    "5hhRxeZknnk+0d3dPbenp4cBIBgMQghBbrebiYjcbjfn5OQQEQWLi4sHzvXe+fn5VHGOqWZv91jz"
    "8/PpXCQf2xTtGfsLjDFzo9XqIYyOBpGSkkL5+e7JOXNWjryTTby8vPzMop95LOc73/7+/qKhoSF3"
    "/Hrm5ORQTk7O+Pz588feJfJR/mixKc88xsnJycLBwUFPd3f3ma3S6/XGv1JmZqYpLS09fR5+uKwL"
    "SJmZamtrRVVVlZ65xnt7ewuCwaB3ZGSEg8EgvF4vwuEwFRUVcXFxMXm93r7EEagbNmxATU2lbG5e"
    "zu+WkXnJKz5XVlVF2OcTtGRJN4B/GW7dfcqblr6hNCP/tpHmQxBjQ9qlwxIw0JJhKw1ihmSGgYYx"
    "NiwGqdAUT548Zih9SKUqc4VtAV6vW9hCwhgNYkBKilYpI9ryRARYiCDQ22mMMwnlhIRG9HXjRL+S"
    "0ZgeHACNDJMVDpEbAHQYLA3i4xw4ppllQNBgWDGFL9uQSSkqZeSWNg25PR9cWLxowOfzibVr35kV"
    "FRskQYjl1S/Vvdm4caO6774NTl3dlr/u6OjwdnV1a621Fa8OTgxDJ0JrxyxcuJBSUlKaQ6FQbzyM"
    "fbkTUry1rbKyct/Q0Ogff+5zn/tRbKQiIpEIlZUtRmFh0aPxc77AwUiXZWg2HtmIG1Hbt29fLYT4"
    "6MDAAKampnQoFJr30ksv/cHY2JiZnp4mx3FARKSUYrfbTVJKzsrKIrfbPfDwww//v7S0NFlYWCgA"
    "PH799dfvi793fKLV212nsTC4qKqqcmY8AwoASym11jpjx44dfxoKhaxwOCyCwWBEa3v1E0889vH+"
    "/n4zPT1NgUDAKS0ttfLz834E4At1dXXqrRoHNTU1MmEjn3kskFI6O3bs+EMhROnw8DBFIhEKBoOY"
    "np7WL7744leGhoZyenp6jDHGWbFihSsnJ+fLAP7r7RzLW7h2OHLkCMXTP8zs2bFjx1enp6dThoaG"
    "2LZt56mnnvrK+Ph47sjIiAkGg6SUgsvlgmVZSE9Pp5SUFPvBBx/8bl5eTig7Ozt8ww3r/52IwhUV"
    "FSGfz6eMMeJyXeOxgS6aSODw4UMrQqFQZWdnJ42NjTlbtmzZEIlEZo2Pj5tQKETMDMuyKDs7W2dl"
    "Zcmpqam6hx566OXs7GxPampq8KabbvoeEU1djHV92RAwAJDfb+L6wrmLr6utYX7sHrXn9nRvRg33"
    "nM7oaW3WbE9LRgSCADY2BAtACtjCDbCAxZLgONKZGOT+5t0sMjyCJwMgEmA40bYYCDBH17mgaKbW"
    "BQ09NiiMdGAC0xDEYENQkECM5MkOg2wDxVEDypCGJoIRAHNU6lJHK8UgQbCEBW2YVXo2UmctYFkw"
    "966Fq68ciHmP+q0RLYEIzHV1CgAaZ4ROYvORL3poxOfziQ0bNjj33js96/jxB/9o8+bNxGxUfFLQ"
    "zIlBiRNQtNZm3rx5yuv1Hl62bFn7pdpgLhUJx+2PrVu3bol+G4IxLs7JyRGrVq06meB1ve883qqq"
    "qjMeW1PT3jtGRsbLOjvb//LXv34ol4hSx8bGMDU1haGhYQwPD4GZheM4ifeahJAQgkgIgbS0tILi"
    "4pJvZmZmIC0tHULQ13784/uH5s+f+y/Z2Xknrrrqqhfiodm3swbiXktra+s6AIVDQ0NmZGRo8dNP"
    "P/Vnw8OD5p/+6R/xf/7PX7lSU9MKx8ZGMTExgUAggNHRUfT19SESiQitNQKBgLjyyivxoQ99KO65"
    "vp1rpwOBkTmjo5NX9fQMIBwOYGpqeu1TTz3x+52dp80//EM1nn/+udmhUAjj49HrGAgEEAqF0Nvb"
    "h0BgGrbtCK2N7O3txW233XYTM/9kpnFxEe/1mb3m1VfrPjAxMfk3P/jBfy7s6ektGRsbw+DgIKam"
    "pjAwMIhAIACtHREdNkKQUkAICSkFLMuySktn/3V+fh6Kioqwb1/Tn2ze/HxLXl72v6xZc93zH/vY"
    "xwJxQ+4yMEwpFj7m2B6avnfv7k+dOnXqbx59tDYvFAql9vcPYGRkGN3d3ZieDoDZCK11TL1PwO12"
    "SSJCenp6RXFxcUVubi7y8vKxb1/TF1988cV/u+66lb9MTy/qjzkp1qUUFXrXeh5jG5rDXCMFkWZg"
    "88TOnde7l6RuKJ5V+JXWfTsMxoZMmtHKRYCBgTYSGgqAhBQCAg6kjhAmp8gJMIjdAKuof0oEZoNo"
    "MPs1dWbBNiKj/bBDk4iEw2DbQAqCgIiSrWAIinrcBANNAAsZLQyLOaJMDC0NyBhIA2iyEIbSi1au"
    "VVPZhT/KveKKIY4OaL+QB42Y6yTR2RXSZ02DIgE2egWAk0QUipP1JSAFPnWqZXF7e1t6W1sbxxWv"
    "Znq8ia8REYQQlJmZycXFxSMxb/19mS+66aabTp3vAX+fkq8GgFOnTt28c+f26ieeeLLixIkTOH36"
    "NLq7uzExMWHbtk2O48TvtUq8v3HDCzgjsgIi4oMHDzqWZZEQgrOyslL27WuaU1a25L+uv/56vPDC"
    "83+9YMGiBxYtikZ/LjR0x8yiublZnThx4i8HBwfLH3nkkVuCwaAaGBjA2NgIhoaGYNs2pqamEAqF"
    "MDExYYfD4ZhcqAOttSAiqbWGEALhcNiUlBTLYDD0lrcmZh8J8Qn9858/8O0f/ehnn7Ztu6Svrw+T"
    "k5MYHx/H0NAQwuEwwuEwRkdHTSAQ0I7jxAZ42HFSUkIIYgaIJNu2jdhkrou6luL7QFVVlT527Nj8"
    "cDi84MiRQ3/3wgsv3HL4cDO6u7vR09PjhMNhjkTCiB6HUERRgyp+bxOJlJnR0tLieDwe9npTqKSk"
    "uGTOnDklq1evLn/ssdpnifSuUCgEYwy9l+QbX+NEhIGBgaWHDh2682c/+8kfHzx4YOmBAwcxPDyM"
    "/v5+23EchMNhAKyIBCXuY4nnb4zRSinjdrvh9XqpuLi4ZMmSJd9paFjx5Wefffbn11236sd5eXO6"
    "o2uEL4lR/q6LDhC9FrrKWLfuKICv9h7abgrW3fInovu0Gj12jO2pcZbKCMOAgAKzgcM2JDRIRmf4"
    "uplgswZTdLAgwDBRfzI2YiFGGGDoSABOJAyQgBIEaaJ5XUBEPVsgWrxFFOsnjpY6xwU3mABDDEGA"
    "MEAQynhLFrKdXdxu5ed8DwAnDlmIF6ChsplRHX2tGkA5ICr8foeoIh4uSo/+dJiGDh/9GofCeaND"
    "/WHWvGRs38sfHO3t3zq28+XGzFlFPyWiwxeThOP50Geeebq6p6eHI5GIsSxLzpzzmxiCFkJAaw0p"
    "lVRK6ZtvvrmaiJiZ33fjyvx+vzlXFbDf73/fnUvc+2xoaLizra3tS//93/99z6FDB7Bnzx7T399v"
    "LMuSAKCUsogILpfrfJt7VHBGnDVwg4jIMsbAGIOhoSHu6+vjpqZ9Zvfu3aioqPjWsmXLvrxt26s/"
    "Xr/+5m9eiLHo8/kEEZlt27bN7+np+YcXXngBx48fx8DAgDM5OQXAQEoSQpxJhZCU0oqvQSJCrHId"
    "UkoQEWzbJpfLhdTUlLdqiBGR37z88os/qK+v/8Ljjz+OkZERe3JykoLBYDwsf+aBkFIKIhIxQxRu"
    "t/uM8RLvJQeI09PTMTU1tYeIIr6YEt/FMBqJyBARdu7c+ZGGhoaaw4cPu1pajmHPnj16dHSUmJks"
    "y1JRoRkFKdXr0khKvfZaQrGl0lpjYmIcY2Oj3NzczHv37uU77rjj7rKysrsPHToEZpbnSkm9m3le"
    "ZnY/++yT33/yycfv27dvH44ePYr9+/c7oVBIxtaClbg+zuEIJt4vCUDGDanR0VE+fPiwPnLkyLyT"
    "J9t9hw8f/uJTTz3143vuuefviMi8FQPzsiXgxFAg19RIVDYz0Y1/Hho8dv+E8PxttpV2b6SvgyY6"
    "TmiLWVpgEDlwOAKWDE0S2rEgjAUiBwQdDT2/zomJ0irYQBDDxdEZwIwoiSoWcES0BcowAIh4MzPA"
    "BMGANIBAdPawFgaSAJIKYXbx3CWrrAlXyn2zll/XwnV1CuXlOh5GPp8n7AfMOI/nBppPfNgZ7b9p"
    "Ysem3xkbGTYmHKZ0t5Ut7BAyJycRnBzD6abtnJWddZMdGL1pfGroj3tfeeIBAF+Oz+l9p9e/trYW"
    "VVVV/N3vfjvU1zdA8c1t5qjBxAc1voDT0lKRmZn5vh8d9H4k25nYuHGjVVFRYdfX19+1b9++Z+rq"
    "6sSrr75ixsfHWEopPR6PiN+3qF0VLbDTWnNMjtNJ2JQEwEJrosRNOlH5TEpJMYjOzk786le/0itW"
    "rCgVQvpfeOGFxXfeeednL9RIHB8f51OnTpmGhgYzNjYm3W63SknxIvHPz8zcPkNwr/1Ma33Wa1Iq"
    "WJb1VtcAA0AgMH33vn37dGtrK6ekpFhSSqSkpJyleX62ocJnChXP3thxxkBITU1Nu5j3ecOGDXZX"
    "V1fuwYMHn3zxxRdvfOGFLdzSctxMT08yEUml1FnPavxYE+ZgszH6zDHGn3khxBnP1rIsaK1JSknj"
    "4+N44okndH5+vgBAoVDorEjYu0W+sfWkm5qabn/kkd889PLLdfmvvvqq6e3tRTgchsvlUm63G47j"
    "vI50E738+LVIvEaJzkXsOqjW1lbT0nKC582bl19VVfU3gUBoJTNXEpF9sSOR76nsHsVCZrGTOgbg"
    "s6MHtr3izsr+usebVjbS0Q57epy9hkmQgWHAJoIWCoZdsCg68jB+NQRHFZ6j1EtnwrlE0dGDhqOv"
    "EkULquI6JPGNKc7Z8UpqEZtWGP0+ZimSMgULF9Mk5M+8Vtqetgce8FBFRTzu5QDAaMv+q9wcSus9"
    "3ROW0ro2KyPzD6ZGRwzsCA8//nhWRkbGonDvaXT3dCASCkFqg4lwUFvasGLNklhkEkmeHNODXSdN"
    "MCXLM++6m74w2FjXX7D2Vj/X1EhKyP+83VzhwYNNt7z44svrurpOa6WUOBf5JuZ+4gu3sLAQpaWl"
    "LpxDJzqJd5d8N2zYYG/evPmepqamRx95pBZHjhx1hIB6bbjEa+SgtYEQUufk5IjZs+eQ1+tBXl6+"
    "5XJZ0FpjdHQMQ0ND6OvrMyMjIyCCSNzEtNZn1QfENjS5f/9+7u3tsz//+fvufemlLbxs2YovdHd3"
    "R9auXeu8UQjWGMOx95PGGIq/f0xl7XW1CInGQHxdxskxbiAQvb16Ia31VGw9U5xY45/nOM7riHhm"
    "WiZxoycieL1eeL1eAwDl5eV4Jy1JlZWVcsOGDfaWLVsyn3766ef379+39vnnnzdjY+OCiCjx2BLb"
    "B2NGglFKaa83BWlpqVZaWtqZc7FtB1NTkxgdHbMjkbAQQkohzJn3ir2P7O3tPfO+ieNI3wXyFURk"
    "mDll69ZXf7lz5467nn32Gc/27dt1bM1AKXXGEEt0IBJFhGaMTT3rXiZMdDvzmlJKEEl0dnbyf/7n"
    "f9of/ehHP2zb4Rpm/njME75o883VRbhI79giICL2+XzinpISmX3F+p8w86/Gd79ak5U/+66xk0fk"
    "ZF+HdjuQlnZgGQJJA1tEYhctRqKMM2TKMT41RNFQs6EYIUcfLCdeYwxCdL3xDN+ZwQTY8uxXDbkZ"
    "6blCFRQGS+76yB/GfzrW0ZEdnhhY7rX0F0dGh5xw17HPKjC5h4dhpgOYtiMg2waHQ+DgNPqmJm3A"
    "EQRHuBgQ0TGGUoJiH8VwojdHZtu21GOj9vjwsHBll6wCAMTaNd4uYu0ezIz8wcHB9L6+Pieav+Lz"
    "hiYTFrVevny5tCyrFsB4ZWXlWyo8S+Ki5sPszZs333P0aPMjv/jFg6q9vR1EUFrz6wgBIKSmpvKV"
    "V14pr7jiCixbtmLY63WHmfmbQggNICKEuDsYDN7W2tqa29TUhAMHDppAYFpEIqGzNq9E4oltYjQ0"
    "NOj68Y83RlyuL342GAzX3H333c/MLBSaCbfbrYqKikRGRoYJhUJs2zZHCU8LKQUlbqiJU6tmrsu4"
    "x+l2e6BU1AN+q6Tn8XidWbNmiezsbK211uFwOBosM0bMJNuZ3tVr1/m1788XAn27RtaTTz751R07"
    "tv7hK6+8uvLgwQNOJGKrKKnImGElzpCQ4zgspTSZmZk0a9YsMXfuXDF79lwIIZCamjqckpIiAGBi"
    "YsL09vZ6tNapJ06cwMmTJ8309CRrrWViyikx2vAuer6CiEx3d3fK5s2bNu/ff2D9b37zGz59+jQT"
    "kUwk2EQ5WSLSSil4PB7h8XggpYSUElrreB6fg8GgASCjfxqNCMWv32v3WkMIUCQScj3xxGN2IDD1"
    "kVAo8Dgz/3Zsv7sotSLqrRAtZqqKVFcDqBe4CL1ifr/f+AHDNTWSiIIA7pk43nizysl8YvJUZvbo"
    "iSO2CgekW7AAayiOX4NzLwhOIOWL46gxjAKsTBdcXrKGm57/h4mgHpNhnRY41vQXZnw0leAQpsYw"
    "2NuO4XBYs9EQRkcD32yg2YGBISHJYkTzzcIALBjGcMwVNzCIqm85RDCwoK1UCW+aEMr9BABgcPBi"
    "mJ80PT01NTw8xFNTU0hNTT1r0z5XCDo+/ai0dLZIT0/fRkShuro6VVtbm2TEd5989YEDB+7as2fP"
    "kw899DDa2tqYiIQxrw/VEgGW5TK3334b3333h7bNnz9/43XXXff40NCQlZ+fP5lwn/8bgHvnzp2/"
    "d/3112/YuXPnlc8884zp6uokx3HO8gxnbsSx/LD1y18+ZH/1q1/90eHDh39r5cqVR8+VN4upRYm8"
    "vLyhZcuW7bvvvvuuam5uxuDgIEZHR9HR0Y6JiXETDofP5FrPNxQkTnJKSbhcrrdMED6fjwBQbm7+"
    "J2677fZni4qKFrS3t2FwcBD9/QPo7e3VkUhERot6Liz0SkSIFfa8o/sc93w3bdrk2717Z/Wjjz6K"
    "np4e4ziOSoxKxJ/NWPugKSkpEStWLJc33HAjSktLD2ZmZm4uKipS4bD98rXXXvsCAPfAwAAXFBRE"
    "tm/fPtu27S8ODQ1deeLEiVt37dqB/fv3OxMTEwIJUsUzJqBdUiKOp9laWloy9u/f9+wrr9Svf/TR"
    "R+3h4eEzo1ETa1VitQDsdrv1vHnzVEFBAebMmYM5c+YgPT0dXq8XgUAAfX196O/vp5MnT4qOjg4M"
    "DAwYrbVQSiHWhndWjjj+3pFIxNq0aZPt8Xg+5PF4nunv76984oknwhs2bHDeKQlfEAHHSPH16ihR"
    "C9PU1FTKBaO3izUJ6jAoB1A7yFRVpePUksCNbxiW5rh/umTNqx2N21amzVv256VzF36t80ADprva"
    "7BQDy8U6OpNXJFiklzQqQiA2ZIemMd7XoVKd0DcoaENMBzHV2wsxNY1gKOIYbcNSthSCJTG9NnlJ"
    "AEQGfKboi8BaxUyI6FAJJgPDDBPV5ARDYlq4ddrs+YLScg6UrL3pl8xMFPVY3jbq6+sNAO7q6l4z"
    "ODhM8U0ubuWeLwQdX5yZmRmUm5tbmKTC9yBtE7s/J0+ezNy5c+ffbNmyhY4fP+7En+WZm0fsr8yK"
    "FSv4tttu7/3kJz95W0LEIhQvECovL4/XLoQA/Ki/v/+XXq/3t+fMmfNAbe1vxJ49exCPkpyPEI0x"
    "1NZ2inbu3DFLKfU1AH+4YsUKca6IV2w9DZ48ebJi6dKlBWVlZV+1bXtJT0+PEwhM3Xnw4EHxwgsv"
    "8Pj4+BtW3iaebzy0+FYN/3gKbPv27evKypZkd3Z2+AKBqdyRkdHVo6OjRTU1Nbqzs1Mm5nzPfSyv"
    "7RVSyjPFTm8H8cK6urqXqnft2uV76KGH7N7eXiGllLF85VmhZiKCUkqvWbNGXnPN2uHrrruuoaRk"
    "1sNXXbXm0XhvawIiCd+fBPA1ZpZ79+76rUWL5v/VjTfeeMtzzz2H/fv3G2Y+k5q6GB79hZDv/fff"
    "r7q7uzObmw899sordetramqd0dFRa2bKgZmhlIJSyqxdu1bceeddKjU15VBGRuZpl8vlKykpmUxN"
    "TVWpqRYHAo7o6uoKM+u7pqenPtTaeuKGPXv2pG7dutWJRCIqMRSduI7iYX3HcaynnnoqkpeX94HM"
    "zMzvbtiwYUNM/OQdOZ8XtEKoqkrXMavy8MT8UGe/AUUIbjcmI0TpaZb2Fs09BdTq81PXuUm3rs6n"
    "UB8jawDl8W/iN6O83BBRD0j8xXjT1rqcJVd9Z3bp/KXtDbsjTjjoMqSjc4Xjv0+XNikpNYMmpmFP"
    "d6K39ZRjtM0uY8NltJKaSQiltBIIE8MhRIu2YiMMwQzF8jWdFmaQYZCgWPW1giZAMxwtCKSUYuWC"
    "7Urn/GVXkJ2Z+U/8938v6uurX2t2fpt7uN/vN2NjYzk//vH9f9ne3saWZYmZ+bWZm1wsVMMZGRmS"
    "mftWrlzzcwBUXl6eDD+/i/jNb34jq6qqdFtbW0lPT8/6nTt3aiJSiZ5JohXPzPB4PPzBD35Qzp8/"
    "/9tEpGtqalxVVVWRGPE4MSI6a/MvLCycAvDL3bt3h0OhwG+Gh4fNyZMnRazt7LyGgdZaPvPMM2bO"
    "nNm/Nzg4+O28vLzj50pTxUiYiGgcwDiAL8V/duhQ0w15eXl/zswfeeGFF8TU1JSYuS4Tw47xf1uW"
    "622RXjwFduONNw4AGADwaQDo7e2d395+4nPp6Wm+n//8FyYmb0jnf5/YWDgC3G433G73OyLfzZs3"
    "f2n//n2+X/ziF87Q0JBlWdYZ0o0bA/HJZFlZ2U55ebm6/fbbW1esWPaZq6++dm9iGHvNmjU4deqU"
    "qaqqMjNbkBobG1VMe+A5AM/t2rX9biJ8V0q5ZP/+/UZrfd76kIuN6upq6ff77aeffuIzR44cvemh"
    "hx6OTE9Pu+LH6zgO4tcBIFNSMovvvvsuuWbN2rbly1f6V6xYUROLoJ4PrQC+f+jQvt9atGjRj7Oy"
    "suZs3rwZk5OTZ+WRE4tO4ymF6elp66mnnrQzMtI+dvBgU+2xYyfq4qHyS0LA0apIYOLA1g/qbU//"
    "cW8gfLcOhOHyuOCAoYWA43Kj4+mHfyKJBsIur6XcXmYik5uTKyfDzsHiq695CIAbiMVVo+FqQ0KY"
    "ioqY9XDm2T87Z8M+n+CNG63m7B7KvOrGZ4NtR5udwe4/y1tn/WlPy3HwcDdcTuiMYy3oNcs8/rjz"
    "RWRkwQwVsCFJQxEriKhkpRAOWAIhMDRZgIiSrmGOTU8SECCwIYjYAREMAMOOBtsQsJVitlJFwazZ"
    "yp2RiaABYLltV1GBGmb5zwuvvOWR2tYBUVXxzpr6Gxoa1Nq1a+3m5uY/7evrT+vu7naEEFZijucN"
    "CixYKSWEENOZmZ6WSx2KSuL1iIf7Ozs7P9nW1mbGx8cpKqogXlcxDACRSMTk5eWJWbOKD5WVlT0c"
    "U/ix48Rzrs+oqKhw4p7IddddV7tly/PmrrvueuTBBx80k5OTdL4QdCzXRsPDw7qtrUNs3br1ax/9"
    "6Ec/XxftDnDO4wlTdXU1xaVMR0dHxapVV+9g5uZAIPCxAwcOYHJykmcS30yRmJigxNtej3FPOH4s"
    "zc3NVFxc3Aageteu7VOtrSe+3draeqYR9s2KkFwuFzwez1s+Dp/PJ1paWuj555+f3dTU+Cc1Nb/R"
    "3d3dNJMY4l5gNK+b5nz845Xqyiuv+s5nP/vZr8feR1VXVxtE1fTscxlLCd/bcfW96upqcf31Nz67"
    "devWQ5blfhXArP3790NrLRIJ6VLlfaurq83+/Q0feuWVrf/0i1/8wp6YmHAlFoDF87mWZfHSpUvF"
    "Zz7zGSxbtuy/7rjjzi9HRUai544EScrq6mpUV1cDiLZfNjc306pVV72glJr7y1/+/Ccej+f3amtr"
    "ORwOq3gkcGYRoDEGlmVRZ+dpuXXr9ty8vIKffv7zG+ZeshxwXGDh2MG9ZelDQ087HUfQ39XDKUSk"
    "pISRDC0ITAr5Obl/5HJ5oZQb0rLAAnDGe6FDYZw63vhtKaUgoQBFMMqjvemp1L6p5kRodOg/4bLc"
    "HrfXCI8VS5i74PJ6SXhTHbpi/a9h9JkL6Z2/rB3AV3r3v+JNEeJTrk5KC3acZK0NSRKXvDKPCdAq"
    "2utHDBjNgIiGwTURbEgQEzwRG4IBFgKGCIaYHRC0YLBhTYKYSYBdKSojN09YqelIyS1AAC4EWdSk"
    "z5kjbaYfFeaWHEK211OQE9VbRryK7B2gsbERABAMBrm/v09MTU1py3KdyQy8USFWPESjlBLv1PJL"
    "4u1h+fLlDABjY2MfHx0dFcYYLQSdN2/PzCYnJ0cZg4aSkpLBjRs3WhdSNBfXX/b5fOquuz706P33"
    "//CV1atX37xz504T34zP5ZHGQpV05MgRWrBg4eIjR47k+ny+sfMVa8Ze44TjNmVlZWp09BQyMjIa"
    "ioqK1sZFYt7oMxOLoN5BeP+sY6mrq1ODg4O8aNGVv54zZ+e3s7OzaXx8HBdorL4tlJSUyA0bNtgP"
    "PvjgA3v37l3S3t6mES0YOuv5jN9fKSXfdNMtqry84tkbbrjh+4gNY6ioqHDeShFawrmbTZs2uW+6"
    "6abOXbu23a+U/KeTJ086Q0NDYmal8cVEvCf/T/7kT4o3bXrm6aeffhp9fX0cv89xYoxFOExZWZn5"
    "nd/5nc6rrrrql+vW3egDIOrq6kR5ebme2QZ6rutQU1Mj/+u//os++cnP/NEjj9SuHRsbu+L555/X"
    "tp1QeptwfxNyzqKxsdHJz88v3rRp04+uuuqqrxUVFQXf7l6o3uym7K/bbIU6Txk+2WpSCMpybChE"
    "c5k61s4zOtTnRPlJn2lGNwRASqUsq8hGLJ9PQEQIjJGE1+0tSs/MWE9KQlouSMuCkApSKbjcXgh3"
    "Gvpqf/xnxu22tSdFREBQQpBHEtwT4/0eT6qQ6RkICDrj5jLzJQ1BawGEhANhBBRLEKLermYnJgJi"
    "AZogNTMcW9vCwHg9QqSmCVgWyOVBVn6ecnu8IHcKRmzAUepgRmGxSkvL+qbLnXoia9HqxteRX0Jx"
    "98VCKBTC5OQUwuFQTEzgrEV2Tu8m/tC7XJaRUsYjGsm5ue8BRkZGxoPBYCwkKc5rNMVzo6mpKZ6Y"
    "V/tWCYEcx6HFi8sevvXWW29paGiwjTHnVPJI8NDkqVMntW3bt0xMTCyvra3dWltbKy/EgIyFg1FR"
    "4R//+c8f+GleXt7a2N+963rE5dEUmOnt7c1MTU1FRkYGJiYmztmTfDEQr2z/xS9+8fv79+9ft3Xr"
    "Vtu2HRXv7018NqPPo+FVq1bTHXfcPvKhD33onoSw/juKkn3wgx8Mb9y40br++vX//MQTj65es+bq"
    "T2zZ8qLDzOpSOTlxXefHH3/sL3fs2IFdu3Y5lmUpIjpTIBXLxaKgoMD55Cc/4br22nXfWbdu3Q/j"
    "leIVFRUXTILxiX1ERKWls/+youLWXzc3N2e1t7dH63fPoQYIREVMAoGAOnr0KPr6ejdY1rV/GWuV"
    "elvdQG+aMJHeFJmSli7GpSSpHUhmKDIw2gEEQUKAIKIl8bGZvIgTMDuAHeLoIAMGmOEiAkgAk8TT"
    "I6SZBRmAQVGRDMNEUkqGsMibln4tpAXypECQgLRcYKXgCML05DRMcBi21iRJQGsDJc6oalwSEAwE"
    "DARiuVxSiGgYIyQbBggCQrkQTs+T6VlZypuairDlQkjQlDcjC2m5eWQb8wOSdNKTme3Jgqu7cNk1"
    "j5y1adb5FFAO1NcbxGaLXgoJNMdxYhumnGnhvT70nlB8kQw5Xx44WwaRMXNfTMzx2baNUCj0tj6n"
    "p6dH+3w+mju3dEdzc3NXcXFxSUdHh+FzsP5rYXBGOBym2KAEDQBvZ1iHy+VKj+c830sopXSccONG"
    "xsxK3ItEvnr37t2/s3379p8999xzHA6HIaWgqPqcnOkBc3p6hrnllpvDy5YtqygvL5c1NTV8sVoC"
    "77vvPr1hwwaaP3/RP1199ZrfOXbsuOro6ODz1QC803MHgGPHjs1/7rlnP7V7924thBDx/ut42Dnm"
    "/ToVFRWuefPm/eSGG274Yaye4W1pNRORqaurU9dff/0LTzzx2P1r1qz5ek9Pjw6FQta5tPATRTx6"
    "enr0vn37zOzZc+4F8MPa2lrxdiKUb0jAdXV1KjM7bSgtkHfcUzp34WBHu2FyLMaZqfTxVn2wAZxY"
    "pCQqvxl1jkhwbEwCwGSgTNSUNcwEJsGMWF6UYNhAACB2YGQYU6Njmg0gpAKMgJDqtRyvEFKzE/sb"
    "giQ6p0N24dQVU904708Jwhh4tA0YAgk3Iiwdd0a6cmWmwEgBpdzIzinCgJUxGBCqtnjWLA+Y9rlL"
    "sx+0beHOm1XCRFlnjUpjn0/A72cwE2priSqqnDO58EswTzQ+xzQqPeecVcRyPqt+5ibjOMm6q/ca"
    "aWlpKW6350yP78x1HxcpiJN1IBAwMe/yLbGZ3+83lZWVcs6chYceeOCnPfn5+aWdnZ36PFXQMU9F"
    "IhAIYGpqSgghIm875cNs3k3hh/MhEnm9oXGRw7DU3NxMhw8fzmlubv6rnTt3ck9Pj45rdc9MMcSq"
    "f83tt98uFy1aULtu3bqDPp9PVFVVXbSDiksvXnHFFYceffSRR8vKFn/q5MmT2u12y4t9LyorK5mI"
    "zHPPPfOnR48ezT958qQjpZRxMZAEuVFevnw5rr322uH58xd9N2YEvqPZxRUVFdrn84mlS5d/p6Wl"
    "5SsNDQ2eEydOsBDiDVMeo6OjfOzYMWvt2mtuB/CDF1988eISMBFxQ8NGmrNkQ/fEwfp/Kym4ZaOd"
    "lomJ0x2gSACWjiASCbE22iECjETUM4xND0KMTKOmOQnAkCAhmBixyiQQm2gVMAMEJ6a9HPt7R4BJ"
    "SSMNNNkwwoBJxDSaJYQGCAqSFSSiFcXizPlHhTcMJUhSvgn5EnTs76zXBjrEVLZYStiwIIwbFglo"
    "jwcmJc3MXbFKDTlOv+1WbdKTSum5ecIx2JLpSbs/Z8Wajhkh5Kk44TbeUyIBYHKyjGNDGAjvcj41"
    "PgHnXIox5wsvxr86jiNiYvjJHPB7hOLiouacnKxVSr1eUi9qpSswA1IqCodtNgaF0b7K/dNvNXWw"
    "fPly9vl8oqCgMLuwsPCMR/JG3p8xhkOhEMLh8N8D+PDbO8t4K+qFPseXjILPEqG4mJ4vEO1F9vv9"
    "TnX1X6W98MLz1+3cuZ2FIKW1iY5MnWGAGGN0SUmJWLlyVd1nPvO535+eDlv33Xef47/IBnt8LvLc"
    "ufN+MmfO3E8ppXgmEb1TxHt+OzuPz3r88U1/2NDQoI0xMh5lSExtKKV43bp1qqSk9N7Vq1cfezOh"
    "l7eQZpGO40zMn7/gYGFh4bXHjh1jl8tFM42eRE8YgGxpadG9vb23vPTSS0tuu+22lrdzPG/oAa9d"
    "uyGufXl/d9Mr/ap0UVleyYKv6bFhcHDapCtVnJmRZgUj0el5yrHBxoF2HGg7AtYOtGNjcmoSMAbh"
    "cPhM2TwBYGMA1nDs6AzguApVdKSIiFIxxSQkmSERbTzXBlAx0QqQjqpecVQHlGOuOSOm63yBXvA5"
    "fOfosbIBQyAtPQOUlQfOykNGVo5Jzy+AOy3zK3OXrvnhzCrDc7z3a+OZ/H5e64d9DvcbNTWVMj9/"
    "OQFA+eAKpouwuGYiPqYtFAq96SYaR6JSDLPB5OSkjpFvMgf8HiEnJ/e7paWzPu3xeJxIJHKWvF68"
    "ICkSsbXWjN7e3khKSsrtkUjkhqqqquff6kYRny385JNP/mdOTs73LctCOBw+58jKxM0qFAphampq"
    "XvJuvTk6Ozu9Dz/86Nd37dppxsbGoJSiuOc7U/FJSmmWLFkiFy5c6BAR19XVXbK54X6/33z0ox9N"
    "y8jIQEpKCiKRyEX1/uvr62VDQwMdOHDsL5uaGtNOnjzpKKXkucK/xcXFWLBgQWDVqlWtPp9PVFZW"
    "vq1zrqmpkc3NzWfkJDds2GAzsxobG3mwtLT0GqWU5vMUVsRbkohAU1NTOHr0aM6iRYv+D4Dfeztp"
    "ljfNASck9p+MHcB/AHABcMYPNX7clepZPT46ZqC1iISmIUwEEdsGOw4c7Qg7HIqI4vnrJYQ7Oy39"
    "umAgCEkCRIDW0VyyiZWPi2hwFEwGEWlH875GQGgBoRnkGBjbhiSCsSPAZA+mh3pgOzZIigTLLDoL"
    "ycR0nN+8PT86qAGJgww56lFHiZwhJWB5LURSLLAy0DosutqOz3HaW7/avvXZTOlK0ThbOs/EPBIG"
    "BCnpdgqvXPufAKY5RlpSWSYWDwazQVXV2b3U8SEPF/Phikv0vRUrPuH3KByOGLfbnXPy5Mn1Cxcu"
    "3H4pJoQk8aZegzh8+PBgWdniQxUVFaueffbZiG3bgpnh9XqRlZWFlJQ0zJo1S2VnZ2Pu3LlSSnnM"
    "tu1jPp9PvJ2NIhqKnd6WlpZ2wWpQ0bm54+FYe0/yxp0D8Z7fysrfuaWvr//Le/c2RKuuzuNpMzMy"
    "MjLkggULOD094ykANHhxlPHOC8uyTGpqKtxuN2IToi6m9+tMTEzkHTy4/4v79u2nSCRizeyfjhkh"
    "9urVqy2Px/vzOXPmnAIgL7TYbGZ7WeJ8ZmYWe/bs/Mxjjz36xdOnO9cdO3bMuFwu63xFWK/pJQg4"
    "jkPNzc163bp118eMlbd8DS5MiIOIDx8+7Apt3x7vKbOZQVmr8SsAv7qAC2ABUJgYLA2NBU2YIq/d"
    "wfCZ/3vtQacIKeEIox3tdrvhDrsRsR3h6KD02ESO4mG3Y+9N73bP7hzqNhHDwhICDpuEnFjcsbzQ"
    "xUKx0HP8X+ZMiFyzwXRgGtR7GmbgNKYgxQhZKCos/hq5PODJMagUL1gQNBjCMEQsFM8QgBAASbQf"
    "3fMF6XY5UlnCURZObnqIYz2cJi0ji8ZGh3+qI+F9mTmF7rCyJuia8i2xENVFJ7nU1FS2LOuMwfJG"
    "jmyCp0ORSMRh5pxTp07dDWBbeXm5TBLwuwe/32+qq6tp1apVp7u7u28NhSL1paWlKwYGBuD1epGd"
    "nY3c3FxkZmZjfHz8lcLCwolZs2ZRUVHRl1auXNn5TrTb8/ML0tLT0y9AESlasRsMBjE0NMSx3HOy"
    "eu8cGBwcZGamF17YtK6xsdH09/eTUgqJIdiZ0ai0tDQxd+5c5OXl1QDgt2tQXTg0XC7XW540dUGb"
    "LsA7duxY29raqqampjhx9ONMFiouLkZp6Swr5vW/GeFRXV2dTE9PpxhnnblGzc0H7uro6FBDQ8Mf"
    "/tnPfnJ7e3v7vJaW4zh48BCGh4fl2SHnGDckTJUCYJRSOj093cyfP98thJBCCHzsYx+jtyrLe8Gy"
    "MStXroycfTnA0XGClYRYWDOOxpYWAoA1ZWXcGFWPsQHYiKqQvGP0NNV9QZtw3sjQkIY2QoKj0WiK"
    "Ti2Kh6BFPP58AcuTQWdEO6Lh8Gh1aTRbbRCxI2Btw4IDoRkKEmMDAw6ZaLyPiWGgwYLAQkeJ3ETl"
    "NqJjDgUpT8osRwoYEDQBTkw8wLJcCCkFb1r6Nz0uLzg8CcvyILL96edH3dlfLly7/uTFakVqid0b"
    "l8ujXyvAAi4kpROPMHR393BHR0dPcvt8bxCLSgkiGmpqaqq45po11/b2Dpj09HROS/PoqACEQllZ"
    "2ZZzeBxvaw3V1dUpIaTweDznlaKcuVYmJiY4FAp5d+3alXH99ddPIZmyeB1i7TByeHjsT3t7e6O1"
    "qm9cYc0ej4cyMzNHMjPdKXgXJpLFFThnhsLfKe6//34JwGhtf3l4eAjj4+NnCs9mriUhQJZlmczM"
    "zP5YJI/PtUaBqJAMAI59BTOnHTmyr/j48c4/1Dq85pVXtt7e29uL5uZmHD9+HD09PSZ2TiIx5B+N"
    "QjAcxzEAm4yMDMrNzZMFBQViwYIFoqioCPPmLewoLi785tu9Hu9oGtKF5ihfK12vJlxoJKq6mmtr"
    "a0XlggWC1q61ASB88thqHu35ynR35x8MnWgBD/dCGA3JURWsMwVgBMBEg8oQ9BZK0+L56YSRhiTA"
    "JGFIARxVs2IdncJEghSMAciAKJ7D1iDYAAyIJFTMFDAMmEDEmFjbsmCGxzAgCCwkAIEJEqwhOcIS"
    "npQU8i5a9gGdVfr84IFtH4LMauMjR/TFygu7XApCyLMk/c730CcuLiGE7OzsoGuuufbjzPxAdXV1"
    "4GLPyEzigkg43ns4CODZ8/2az+eTK1as4MrKSvNW71FMmQixfJmzZ8+e4fN5Zmf/XXSdTE1N2UVF"
    "RSvdbvfHADwQD7cm797r9+GWltap8fGxrHguP1H96azoYCTi5OfnW6FQ6Afz5y9rvxh6xG+GuDTy"
    "pZiGxMz0xBOPTg8ODr1uvnP8M40x7PF4ZCQSGb/hhvXfjr2umZlqa2vPpFTia0tKCcdxXLt27fiT"
    "SCRS8uCDP13T19d/y8GDBzEyMoITJ06YsbExA4Acx6F4vjfeYhQ9BjKxFiiePbtAzp49W8yZMwdL"
    "ly6NFBQUHPV4PP9dWFhI69ev/08iCgNAbW3txW1DupgWe0JM+MJuTHV1vFBED7U0XK/C4Q3jHYc+"
    "gZ5O7/jx49odCEgDG44EJEcpUwgJox0QAZKjQeWoZOgF5jr5bJ84WhtNEG4P0tJzEQnb7LYjjnac"
    "aDECDMENyWQA4RCTiYZrHBeUkYhJQJ+ZLxw15OKBcQMJEyV1HT1GYoYS0UVnTwfQcuCAYxWOLJrj"
    "uXrfUIr+fH5V1a8u1iYmpTQej5uVUmeR8Pn0oOMPAwDR0dEBrfXNANL9fv9U9cwpWUm8m54wxXoQ"
    "Y+G1ZiovL0d8Q7rQzTluJFdXV1N5ebmor6+PD2cAEZlDhw59sru7c0V/fz+01nShG3HMo0imKM6B"
    "hoYGa+3atfauXbu+qrVT0tfXbwshrBnP2+uewfT0dBQUFFDME4T/ErQrvhvo6elhIuKf/OTHGBoa"
    "QigUwrnyvwnGP6am+tzx8HJsfZ4hvYMH91WNj095JifHP/bMM0/d2Nramnv48GEcOnQIg4ODPDEx"
    "wbEZ1DLu7cbHRkYiEY5JmWq324NZs0pVfn4+VqxYgcLCwhNz587dkZ2dPb5w4cJ/WrBgwXTigIt3"
    "kiJUl9tNYWZqvP9+RUR2d3d3Sna4+57xntM/t4f73QMtzTDDg46XjbJUNDcLEW190sxRylQqJhIC"
    "SMTDCRdShvWadRBvYSIQjFBwp2Yid/ZcuHNyyZWVbk1NTiA8NYmp0REM93RBhgPg8LR2a2OEgSTh"
    "EgYKFA9nxELaMrYPRUcUK2gpADZALNAdjVUzCA4kCBJaRfo6dP8h481YtLR6uHnr/pwVNx15Jx7n"
    "mjVrYh6wK23WrFnk9Xph2/abSusl6qIODg6hvb09glhrVRLvLQljRv/hW9mQfT6fKCkpkQmV/Bzf"
    "TDo7O1edOHGirKur42+3bNl81fHjx9HY2Ij4mMD3WiDj/Y7GxkYwM7388svp3d3dYnx8TMfnGZ/v"
    "2rrdbmRlZSEjI8MBXkspvd8QS6How4f3Xbt3b9MdPT09Wkop31iLgDgtrWgkHl4eGRnJdLvdiw8e"
    "bMo/fbrrH1566aU1vb29OH26C93d3ejr6zMDAwM6EAgIj8cjhRBnRhnGPd1oeBmcmpoqi4qKcM01"
    "a9XKlasA0L558+Y7GRkZ/z179uxfl5WVDSYez8aNG62ysjKuqKjQ76QGRl1ON6S2tpZiSi5238Fd"
    "91Hb/j8KjA9fM3HkIJvhPp3KEQkYZUuBCcGQWkCwAQkBISwolwXDDCcUba1FPKz69rY2aBgYEKbC"
    "EZMaMUJ4sg5F8gv/75RIydGe9CB7s+aUzFryleBAv0lz7BxvKCBH+3sxFhhhEk50mrcxsRprI4li"
    "YwejNdBgqKiKETQEMwRxrH9aAIIgIOBlkpM97XZ2cfGi4NDIXxLwOa6vP6e4/QUSsAaArKys3xQX"
    "l/xZcXGxbG9vZyklnc/qnim9Nz09ZQKBgLVt27YNAL7zdlVgknhvUFNTI2tra1FbW2tim4dhZi8A"
    "z7Fjx67o7+//g4aGBv7pT3/2u6dPd6KnpwsnTpxwxsfH4TiOuth9sP/bDahf/vKX0wMDA7BtB1Kq"
    "80aimBkpKSnIzs6G2+02iQb1+w21tbUEwAwO9ueMjIxkh0IhJ06Q56hAJq21ychIT2lubv7CxMTE"
    "w83NzZ/bsWPr73V2nl555MgRdHR0oKWlRQ8PDxuttbBtWwAQSimRlpZ2JpyvtTbGGLYsS2RkZPC8"
    "efPE0qVLkZubOzx//nzh9br/37x5C46vW3fjrxNVxeKEG5/+9matp+8bAmZmQn39mZLy06ebFlt9"
    "fZ9XHUe+Pt7dgdGONp2qI8JrtJSsYQRBU8xbJCAsXOD0HOTm50KFphAeHgRIQWsDlgQVHa0LW0Ur"
    "lCUzJAiG+cwsYWEITCYWdBYJIWMBMgx7ckr0nTym872piyaDpqDs9ru/l3D83wbA/XVbvh4KTmXZ"
    "GUXrStJxXWBiDIq1tCcnYU+OYXKo1xE6RBYbUjBgWGQIJMGQZKJCINHkGbTRAEswDNzawMVM4/39"
    "Wljp0xfhgTcA6Iorrmj6yU/u/01xcfG9J06c0PHeuzcKQScK/IfDITU1NXkrM3+3vr4+uRtfxvD5"
    "fGLFihWUn59PP/jBDzixB/jkyZPX9PR03fjggw98JhgMruzq6nJPTU1RW1sb9u/fzyMjI4bZwOVy"
    "qUslxP+/lXt7enp0Q0ND5t69DbcODAxAvEl5eXz+bXZ2NgCk/E+4CP39fYGuri6OTzI6n+evtRYt"
    "LS2urVtf+f7g4OC/Tk9Pe06dOoWmpiY9NDREtm2zUkoKIWR0NKUV1YzQGo7jsG3b7Ha7zOLFS1RZ"
    "WRlyc3Mxe/ZsysnJ6czOzrr/Yx+r/B4AnTDKUFRWVsrKyko0Nzfzhg0b7Etx/u8pAbPPF1dTckYO"
    "NKyi4NgfB44e/2RksDdn4uhBHYkEyK1YwiJEGCAbECQgDGAZgrZcjqtwrkiZvVikpXkQOHUIdjgE"
    "qQFtJBBvCTICDjMciei/HQOWAo6I9QkTw8S801jjb+w/hpIEL9ug4IgcOtboTVl69X+M7X3pizo1"
    "7fdyl1+3k4gCsfixHwB6mVNzJ4dmm/a2wkhg0h8eHHQ4fTxz4ZKr104P9CA8NggRCSMUCsIOhWzb"
    "sUmRUNEyLIYgBUECjol68EE4xrHcuqSk1B3Mzs28SBuyBGAWLVrwUHFx0b0cW/XxiSPnDv+cPYN1"
    "amoSQ0P9g3Hx/CQuP9TV1an6+nozM0TW1NR0S1dX1zLbDn+1tvY3JZOTk+lNTU3o6GjH6OgYpqam"
    "4sUsyu12ybi4R1SbWOJ8kZIk3tIzSH6/39x55535jmPfPjo6CimFPJ/3Gw+dpqSkkFIqpJTrMABM"
    "Tk6+L/MAVVVVDADp6Xn+kZERYmbxRtX1zIzdu3fj0KFDPDk56ZmYmDChUAjxsLXLFZ0PkqAdbpSS"
    "2utNodLSUlVcXEzLly8XGRmZA6tWXXEoLS31ZytWrGj0eDxj6enp/fHPqayslF/84hepoqLCiUWJ"
    "Lul1UO/NxuBTFRV+h/x+c+jQtoXp08FvBPraP+EZGfBOn2iBMzbiSMFKQYB0VNWHmGCkQAQEQ8pR"
    "aRkqf+kqlVE8HyaiMdXXg/Gh8Wjrj3AgJANGQCsXHMsNo8MgjoCEgE0OCICKznSGEQbE4rWpSgBY"
    "xFubGAoGxjHgiUmMNR/UHq3LvPMXPXN6/84Pzr5y3W6uq1MYHGQsGBVENA3gWOy/8nh43Tmy755Q"
    "RtENmeB7+3u6IhwMFs1L8XqCEyMYHeiFCQdYhwNaR8KwlIrmKdxeIefOFll5he5gVl6rKyf/BwCA"
    "8vJ35IaUl5ejoqLCfPjDe7wFBYXs8XgSBUTefNEoJdvb200k4tx+/PjxZWVlZccAvK8FOZhZ1NfX"
    "i4RrZN5vUps+n09UV1efEYGJF+t1dbXOPniw9TopxWc7Otrknj277j5+/BgOH27GiRMnMDQ0ZLtc"
    "Lqm1JiEECSFUfMN3HIeFEDojI0N4vV4wsxgeHr5gFbUk3swD7Hfa29uN4ziCSLzZGmWv16sCgcDg"
    "Nddc81jiPX4fozQui/tGnj8RYXBwEMxMQghOVKqKDRwxRGSys7NVamoaSkpKxPz580VBQQFmz57V"
    "M3/+wpYVK5b9h5R8tLR0YctMQzV2Ham2tlZfatJ9zwiYfT5RDaCiIlqZ2bP/1S8448M/sPu7MH6q"
    "FWJ4WKc7WrigVYgklJBQ2oCYoSEQBrGdmoachQtUZukc2Gm5P9XCO4+Hu28bPnFS6+mIFCRhRAiG"
    "AWGlwcouhONWkKPD4KlIVG5LRnVyVaxMWYvYtLNYDzEoqgEtpAQbAzIGFhFcxoE9OSKHmpscb2Q6"
    "J2PZ6k2dh7Z/gFbduLeurk5VrK06M9gaAKO2VtTnN8fHgz0J4Elm/rtsEpG+XXuvT09Tdzp93alZ"
    "pYv+3J4YkkqH1fT4CNi2keqykJlfgFBeyQHjSfvBD6+p+Ik/trFeLGJQyi3mzp1L2dnZGBsbu+Ci"
    "GiKitrY2MzExUWKMKQJwLD5Q/f0aoo1d07Ou68XSmn23QsxVVVXa7/eDmd0AzI4d277S39+/6Ikn"
    "nttw4sQJDA4OoLPzNLq7u/X09DQikQg5jkNEZMUL8WIRECOEhFKK09PT5erVq9XixYuRnZ2NY8eO"
    "4ZVXXsHo6OgF9QMn8cbQWoupqSkRvf5v/tyNj4/bGRkZxXv37v5rAH8fr6R+v55/KBQKBgKBM5G3"
    "8+SAz3pda02xcDxLKRkAFxQUyAULFolZs2Zh5cqVXFhYeDAjI+PR1NRUXVFR8W9EFEp4XhQAE1Nn"
    "4wRFrXd9Mb9rBMw1NTLaw0oYPbDr487w4N86J1uunGo7ys7ooHbpsJRsZFhqGAIMMaSJesCSLTjC"
    "pVV2jkhfXEbWrOInMheWfqOn6Mo2+8UnTk62HmRnrE9YUsKAIQXBFgxkZOv0suVienqU7OEReFhA"
    "2zYcS4FYQJloO5ARiFUfEzgm4hG9M1HyJxIQEJBaw0IYFDZq4sQRzTA56QuWvdi5b8cdc666YQ/X"
    "1anYzYy/hUb8revrJOrrDRFFAKDo+rW7AOwCgKn2I/8+PZqzKsXt+kKkv1eHAlOcVZAnlNv9TPqq"
    "xb8mmhNkn09UX6R+25h0HQmhO1JTU3uLiooKhoaGzhRiXQhGR0fR2trKN944OUVEXFNT874lX7/f"
    "b06cOLFYSqwdHh7WWVlZMhiMnFi1atVeZpYXa8TbpfLc4wZZT0/P3K6urt97+umn7zty5DAPDg7O"
    "amlpQUtLC6anpzkYDGrHcQiATCykinkUICLH5XJhzpzZVknJLCxevBglJcV6yZKlj6empr9qjL1C"
    "KbXhlVde0cwskcQ7iUIJv99vMjMzq7XWiERsByB1Pnsm7gU6joPx8XHYtv0/IhlPBG88/3uh09ii"
    "69TNBQX51uzZs2nlypUoKCg4sWhR2X4p5Q+vvPLKY6mpqVN5eXkTMz3dWGTLAXBZtG+pS7gxEBFF"
    "1bKqqgxVVelTjY1z08a7/yJtvPfLHccPItTbya7ABKWwVkIRIkQIkYKRBGWiqlJBKZldqchbvEym"
    "ly3FOLm+VLju1h8AQP/WzQ85/e1FkeF27VW2NMYFwYAQLhiXy2QtWS49c+dj7PAITMiGywAOAw4E"
    "tHBBM0MRwNqBEPEWoNdModcGOhA0EywVrboWJoyUMMvpY0cNBeyM7LIVW4YPbb+LVt24g+t8iirO"
    "7r0kAgNnVFlEa2urtXjxYhuNjRKnThmat7wXQC+AF84dsq9TVFHhXKzxhFVVVbqmpkauWnX13kcf"
    "rTm1YMGC4oMHD2qllEyY9vGGMMZQS0sLDh06+A8NDQ33rl27dhjvM6WjGPnywYMH17z88pYXurq6"
    "chxHw+12ARCBmpqahwFsuBz1rmPPF4jIHDt2bH5XV+eXH3mk9tOnTp0qOnToEMbGRtHV1eVMTk6C"
    "iFRiaBlAwog37bjdHjV//nxceeWV1uLFi+Fyufry8wsGUlLcvrlzF3ReccUVTQCwa9euKzwez4Zk"
    "6PniwRiTo7XBa63S/EYeMCKRCCYmxhEnLaDxfX3+lqU60tLSFimlEAqFzhtViY/XzMvLw0033aSu"
    "uOIKGGOmS0tnn0pPT/vZqlVXPFRcXDyQ+DeJrUJIUMa6nHDJCLi+vl42bNxIVFVlA4Tgnu23T4wO"
    "PBrobs3o6Txq6+CU9LAjpHTAxkAbBUEWXGyBNUMIwrQmbRUUyaxlq4Ci0i2TaVn/Vbj0micB4FT9"
    "1ofco52fGjre5FhmUgmKDhpm9iACly5ZvFra+SWbxzUvFNosMrbNwmiSiBK0N7cYdjhiItPjghwD"
    "wwZx9eaYaAYMA5AuCI8XwZCDEGwodsFiwK0NVCQogi3HzKSjM7JWXvlsV9Mrd9PVt+yIDlGo0ESv"
    "f5piCkZ2YsgzGrauJtSXi3rURy1klAPl9Qao5gsVHX8raG5uZmampqaGX+bm5l4frWW4MPJlZrjd"
    "bnHo0CF7/fquO1esWPlRAPdv3LjRulTVgpdojYq6ujrs2rX931566aWcl156OSSlVMzGWby4LCU9"
    "PfNmKaX+2Mc+dll5e/GQuRACu3fvvmPv3t2Pvfrqq2n79+9HW1ubEwqFpNYaLpdLxYtT4vctEokY"
    "y7I4JSWFCguL6ZprrlGFhYWRkpKS4fT09P+7cuVKOXv27Efy8vK64n/3wAMPeFJTU22tw5kJRS5J"
    "XBxDKiKluKDfJSJMT09jeHgE09PTAgAa36f8W1NTQ1VVVejv7/tmVlbWbbH9ULzRuSuleP369VRe"
    "Xv6tK6+8stvttpoWL162I5Fwe3p6dHV1dXxs4mW/F6lLtUHErY2RAzs/jPHhL413Hrlj8NRJskd7"
    "tEUhS5EBCYZNDCIJMgoSCi5jwWYYJz1Dly4qs2jW3EnOy/vYF471vlxbdYcGgNHt9VdFuns+2n9k"
    "py0jAaWEiFqQAgjD0mmzy0Q4s6S+ZN09d7dte6ZJRjQEDDM5xExg5YanqASZaemi+/gRY4bDwqJI"
    "tAqZo6uADYGFACsXXDm58Kako39wAKHJUXbZEUoD4IIDySymO1q1QyYrbfGS5/ubnv8AXV2xgwE6"
    "n2DGzBxu7HcYSPSy/DO+XorwDzEzP7V///4fzp49Gz09PW/YhpQIIQQmJyfFzp07zdKlS28iovt7"
    "enreN73APp9PVVfD7NuXdltr64l1dXX1tuM4nkgkAiGIRkfHOBgMTABAbW3t5cQ45Pf7zeDgYPrL"
    "L7/42FNPPXH79u3bsWfPHg2AlFJKSgmlFJj5zChJYwxSU1Odq6++Wi1btgwFBQUoKChEScnsn69f"
    "f8O/5ObmdsQq+s/c4pqamvhCsKuqqvSOHa/qJPleXEgpPS6XO+GZe+Mgkm3bGBkZwejoqAMAZWVl"
    "7+sbkptbaOXk5JiZoyzP5QFHIhHOysqiz3zms98jor74Oo13YcSN/3crtBwfawhEC1vP1XHwrhNw"
    "y6ZN7rIPfjDcf/hwUboz+Wk9Nvjd4ZbDGO1ohdsOsQtGGo7maYkFGBaIJLQQCBuG5QW78opEyZrr"
    "hFZp/6W9qd/NXn19G9fUSNTUyM6F+VdFBvtfCp0+5NLBEVjKRWRcYCcER2jjKswBSueenl5c9QEA"
    "JKdDXhMIAgpwjIHWArYQCHi9xlUyqytP8Jy+XeOGjRbMTjTna2LykZAIMyFACgXFJSjKysV0dy+N"
    "9XSaiAkLKQwk2bCMIyc7WoyOhNIz5i1+Yfrk/s892Jn57O8ThS7XHKLf7zc1NTXy9OnTo3PnzvvW"
    "Bz5w51/94hf/bSKRiLwQlaPYxi67u7u5r6/v3sbGxu9fddVVDdXV1eL9UD3s9/u13w+uqVn29y0t"
    "LdbU1JT2eDxIUMohl8sliAiVlZX0blZGvlHYubq6Wn7iE5/IfPXVutq6urqKp556ygQCAVJKyZlR"
    "irjnYFkW8vPzzYc+9CF11VVX9RUVFf1y1qzZv/J6vXrOnDmHEo2S8vJylEdHYJqqqqozG02SKi96"
    "9MXEiOWflJJ3eTweGQwGIMQbD0ZRStHExASCwdDCtrY2z1/+5V/a7+frMHfuLNfo6EKhlDKJa/Zc"
    "a9m2bTQ3H+af/eynv2hsbPy9hoaGwfvuu895L3TofT6fSCzQfLukLy7qBuHzibIPfjA8dmjbQjXS"
    "vcs+3fbdU9vq9VjrEZPqTMNCkKIayBLKAdwOwaMlSAs40qV1bg6nrFpF7hXLDo7m5Hwi44bbv5x9"
    "5fVtDQ0NVn1+PvWtLpvjdezNk6ePZ0wMtLKkMGkD2I4CRBrYSuXcsiUSudn/XlZGYSJyQkNDEYSC"
    "YHagiWEEwWbo/DlzRCQ/9w9Fdu4v51+5VkQYDpNBtLpdRgcwsIFxbEyNj/Pk9BQyC0rG8+YsGSxY"
    "sEzYltsJCSACB4IcpJARoe5unjx6JFV3dzzyofQTjzGzheZmyT6fuBwXf1VVlZ4zZ07wt37rzv+z"
    "cOHiU6Wls0TMK75Q6x19fb3OsWNHTX9//+eJiBsbGy/7zZq5Rvp8THv37v1IV1fXda2trVpKKbTW"
    "Z0LwXq8X6enppLW+bDzg2tpa4ff7nebmQ7+7c+euiscffzwcDAaFMYbixx0vropHMpRSvHjxYnPf"
    "fZ8XH/3oR/++vPyG62+99favL1myZP+cOXMO+Xw+xczEzOT3+52Kioo33NCS+d+LHoVqc7lc5kIL"
    "IKWUamRkBC6X+l3HmS6tra3Vvst0f3kjVFZWmsrKSpmdnb8PwJ68vFxpjNHn0iJJGAIh9u8/4PT2"
    "dt8RDAbX3XfffU59ff27vt/U1dUpv99vXn21/ou1tbU1P/7xj3/zxBNPPLhr1675zPyWnNqLeeOo"
    "vrpcDDa8+gW7rW2z3dIwt3PXC7Z3clBmshZkgJBlwVYKDIaRDFs6CArbBN3SSVlQJguv+S3CnGu/"
    "NOu3PndF6Yqbavjv/14wszh16l9MRUWFg+HeP7W7T+VMdZ2yveQItwnB4jAgJaaEx6QtWs2maM6R"
    "3JXzfsnMNH2o6ZoUNgWhiREjAQIsGCOM1+MhOxhp4WDkZMGNH/5snzv7OU9JkdJsa2IXNLsRlgpG"
    "2EgzQXhGh1j0DxjBsp2uvv6m1IXL985Zs06FPKkmzBZL4YZyDLzsEEaGuG/Lc5GUno67Rhs2b2p1"
    "uYj8fnO5PiT33Xef9Xd/93ciKyv7X9etW09CCCdxoz2fJF78oQiFgnLPnt3U3Hzwyrq6uqK4vu3l"
    "/PDff/+o8PvJDAz0XXPq1ClrYGDACCEo8YF3u11wuz1ObLOgy2DDklVVVebFF1+8q7m52f/kk086"
    "wWDQlXh/EiVD48UsJSUl/OlPf5quu+76L6xbd+M/lJQs6KipqZE+n08ws/D7/Q4R8YV7EXzBaYpL"
    "HA2Y8f370zBwuVypJSUlQgh51nP1RgiFQjh48JDp7Oy3368nTkR8++23i/nz5/dJKVtmzSolIQS/"
    "0f2WUmJ6elocOnTYjIwM3EVEXP4ONRHejvFeUVHhbN36ypdOnTr1X7/4xYOVDz/8q6qHH/7V5w4c"
    "2H/i6NGDq2PHe0H7/UUhBY7OYeRZu8zn05zwD/qOHV7YfeqoYQ5abEIg48BiQBmC1BrEDIcUpmWq"
    "MXnzxIIb71Qpc5e8HMkvXjRn3Y0/AJto5a/fb1BfL6qqanXvzpe/rcbGvtp36KBtObZljBOddqSB"
    "sLQY2TnImLNAuTJS7snIWDpERBwOT92VnubJtyNBLQhkBIFBJj09Q0xPTTUWL42Gtt2ZOf+UO2dB"
    "gNzprBksKDoQAYhqMytAjPT08sRg3xXDJ4/9KKf87ptEwfzHMsquMZRTTAFt2JIGFhwQIhQyU67j"
    "Dbsc3d55e/F47xPc1uap9vuZ+fIj4dtvv91UV1fzkiVLTqxatSqcnp7Bxhg2hs/ayGdapPH8olJK"
    "nDx50uzatfuarq6OBzds2GBXVVVdthY5M4sNGzbYjY2NN/f392/YtWu3CYfDKnHTE0IgJycHWVlZ"
    "aZfRfRIAeGpq8oMdHR3p/f39RilFif2T59iw7Ouvv06Uls759xtuWP+jw4cPu+KhM7/f/74TGjmP"
    "V3TWv99EzfGyQXV1NTMzzZ07d7KgoOBkXl4uYpN6ziuKEyehiYkJp7u7S0xNjX8DAJeXl4v3631k"
    "ZsrOzkovKiqC2+0+p3EXvx5SSjCzbGlpwenT3Z87cuRIGRG9a84NM4vvf79XnTx5MrO7u+drTz75"
    "pN6+fXtk7969zo4dOyIDA/1E5HpLBpG4GBewur7eTJ48WGj19/xVS92L2h7q0W6yBUkbrDSEjKpO"
    "uW2GZQBDwtjuLJ2+4ErhWXztKSdv3uf0snX3zF219mRLS0sGM4uKigqHGxosqqhwOg+88vV0hP6i"
    "u3G37Z2aslyOA5YESAtGeRF2p5gF11wnKD3zW5ll6ztaNn3PDQDB0HhvyA6wFBwbOq+jGtAgKJfH"
    "y8zUmtar5t186/aIN+eVlMIFyhFkBNmw2IZiig5MEAJwQrLzcGMkLTx143jTzorMG+74mGfust/N"
    "u2Idi9w8nnbCDsgB4ICVAxGZUoP7G23q6fjAyEjbU2BWoMuPhKuqqnRtba1Yu3btyxkZ2Z//9Kc/"
    "5UpJSTFx8n2jqug4ERtjxI4dO+y2tvb1dXUvfWr58uV8ueYNq6urRVtbm6ejo/0fX3311by2tlMm"
    "cbaeMQzLslBcXAyAv3O5bFIbNmxwDh8+PGdoaOiD+/fvZ621Fd+Qz+cxZWRkUEnJLLuoqOCwz+cT"
    "g4OD5v2sVnYB6xDvF6lqIuL7779fzZ8/v8/r9f64pKQEjuM48eduppefSMKhUAinT3fxyMhYGRCt"
    "5n8/esJlZWVMROz1pvSkpKQYl8tF59pvEq+HlBKdnZ144YUXxIkTx15l5swVK1bQpY66xQpqzVe+"
    "8pXwkSPNT23fvm3+nj17CIDLsizl8XhUbm4OpaRYb+l9xUXY0Mjv9xvlhFLT7MBcMd4vM8mWLh2C"
    "ZaJTfrRhOAC0EJgkaTi/WCy5sUJ65iz60cKP37sw85qbfjFr1qwAM4tIJBI6c8Jr19pt+3bO49HB"
    "L3Q27nCsyWGZxjaU4wBGIORIRJTHKVy6goPpmf+cee3t/wcAL75rxObe3tSIdqomJoeJDAuKVVYZ"
    "IUCWC3C5mYg4MlXM7PMJkV7kd5cscrTHww4cCB0dCRgfZeiBhgqMyUBnqzU11Pc3zOyZdWPFwzq3"
    "4OO5V64TafOXqJCQBoLhDkeQSg4sM2WdbnrVptOtd/Rse2JzT3e3l+jyC0dXVlaauro69ZnPfPKp"
    "NWvW7rzqqqvYspR+s1BjQriTQqGQeumll1JPnTr1s+rqalVVVXXZ5aZiRoFpbT12y7FjR9dv2bLF"
    "sW1bzdzwlFLIzc1j27YbY9fnvTYaCAB3d3fnnD59esHg4GB8Qsx5hQsikYjJzc1VWuuOG25Y/4Df"
    "7zfvRR9krKCAL9F7nzlfrTXC4Qi0fv/UJMUmGVF2dnZeaWkp0tLScIHPnGptbTWHDh28+cknH/+h"
    "3++P+Hy+9x0BxycLZWZm/9958+aJjIwM+Ub1J3Fytm1b7N69m+vrXy18/PFHv1hVVaWrq6utS0XC"
    "NTU1koj44MGD2Zs3P/fvL7/80rpNmzbx9PS0iK67sFm2bJnIyso6FA6bzlj4+YLW/EXbIFkJZzo0"
    "aTM7kAxYWkKyBcCFCLk4YHl4ypsSSV2+WniXr2kP5xb6i+/8yBf+7ht/K7iuTsUtjJUrV0YAALW1"
    "YqS5cW5KZOoFfbp9fqinQ6RIRzDsmDKVhTDcTsGi5UoVlLxSeM1tf9vWVucBwER+0x8IpFmE2wMT"
    "Y2BowWzATBBKQXlTQJbbAYAVy5cD1dXUbWU0cVbhpqKlK1RECIekAky0kdcYA2IHKcaWY+0tRoz3"
    "3dRT98wtAFBy7W2PRbJmf9Sz8Op6NWshbKGMxxZQDIRNAGRPWt07XnFcXW0V7o7Gpw/3H077pt9v"
    "6qJh+8vGGq+PqnSNz5+/4MO33nqrKigokOFwmGeG+eIPSPxhiM/VNMbQgQMHnFdeecX6zW8eqmFm"
    "y+/38+WUD66trYXf7zednaf99fX1FAwGRdzLf23ji/agFxYWUnFxcdZlcugCALTWvra2NhMOR0xi"
    "CPZ80n2WZSE7O1sYYy5KNCIqAfjWCNKyLCIizzv1cN/sZ47jIBwOwxgjLvaaCwQCZyaFJZ7bhRzX"
    "mxCwAcDp6em7MzMzp3JyckTUi39jXfaYaIXYtGmT6ezs/N3m5ubFQLQy91Isvrj041uNSFwIfD6f"
    "WLJkSWDBgvlH8/Pz4TiOSRz+Er/WQogznrCUEsFgUDz22GP2qVOn/nnbtlf/1u/3R6qrq+XFvvcN"
    "DQ1WVVWVZubs3t7uTbt37/7q5s2b5fDw8JnRiZZl8dKlS8MlJSX/UVZWNhgbq/vuEvCEY5S7uNgK"
    "eVP0tFFOhD2ObTxO0HidSEoOmcJSmnfLra6UFWvbefkNFelrbqnmmhrp9/sNzai8rK2qIqqq0sGR"
    "wS95BvsWTxw7EknRYcGIICIYtlSw/397Xx4eR3Xle+69Vb2o1d3ad1mbLUuW5AXJ2MibhAETwGYJ"
    "EgEy4XsDg8lCXl5m8l7ykpeWwkwSksxkwhCITSBgQsAtwIDBu92yjWywJcuLJNtarH2xpJZaaqmX"
    "qrr3vj+6S7RFS5ZZAsn07/v48CJX1d3O+rvnMJHrY5ORFBE7QMMif8QtFpyeXuKdspb0lNEJp4vI"
    "0tQgGQfgCINWpwcOKBwAoM/jEQAhVlRUJJOI6J+4zVFdJC4JeUBgHPzhPcQBcQYCMNDIbuRoOaWI"
    "E307uo68fSsAwLwVJTvk1MWb45au5WGpmVgWwqjEBMBEAEAUNEgShs7Uy7ij/cbEprZ33tu1S1ta"
    "Wqp8mcK0lZWVjHNOCgsLR/Py8v6luLjYqzb99uderiCJBHqN6n+MMWHPnj3o/fdr7nz33berOOek"
    "oqICfRk8YZvNJlRVVdEXX9z2v06fPr2ssbGBUkpxYJ1Z//+5TqdHUVFRI3FxcQ4AQGVlZV8KFrTX"
    "6zXb7XYsy9IVufnpwk5tOO7vbqV8Ftfg1AbqCMGcakAzxqjJZBQ6OzsbJicnXwUA7K9INGcIguD3"
    "blnQOsmBApoxBi7XJLhcrkm/QfmZGD4WiwVHRkY61Dn9OPkruOExl4I26rosX778daPRpERHRxM/"
    "IW7G56vfwDlH/f39rKmpKaypqel/V1ZWso0bN36m8sRoNKItW7aIGo1I1Xvls407UPH6OCJXN/wB"
    "AJvNZntERPQT+fkFisFgYNMjO4EphgDjH42NjYnbt29Xzp49+68HD+7/iUoonCsB6mryAgCgqKhI"
    "5pyb3n57x97du3evfO6557wDAwPY3wACBEHgUVFRJC4ujqxbd+NfEEJw3333zXmf489AcHOr1UpG"
    "mXFINse8sWDdBmJauFiIzF0qRORcJ8QWrBCEtFwet2zliC4t5+dSdHhpRkZGR4PVqkFBCt1zzkmZ"
    "1YpGag99m4wN/UtX/QlFL7k1Ok6BcQayoAGXoAVqMMsZK1YTMT7pyXnXrzoJJSU4UImjwSFExycx"
    "cUuAOAY0pUwJdrq9XOHwAgBAktMpIQBeW7tFTFhSdE6TPO9nqctXE0lnVij4WhoixIEBB4wIaDlC"
    "bGyEjLWd1UVr4M2J5tPLardsEZNzF7Vwc9zdxtzlDnHBIuLRRVKvouEMCFABQOCSOHCyVpGam0tX"
    "GdnR/vrq8vLycsq/XJ4wRQjB7bdv/Pebb7753J133imKoqioxRxmImUF5mc8Hg/evXu3fPLkyTv3"
    "7dv7VkVFBa+srORfpBKura0VS0tLlYMHD/6v1tbm/3jrrbfEyUnXVNnNQIubUkpjYmKwLMvvLFy4"
    "8LTFYiFfFrKS1ysrbrd7qriGajwE85QIIWh4eJhKkjehru7DWwNC8J8EBCHEZFnWyrI8p05ICGHs"
    "8XiZwWCIM5vNOf5o9Jy8k5KSEgAAMBgMXK/XAyEz7zn/2wAhhPv6+rjdbv9KQ0NDVF5e3ieOvnDO"
    "kc1f172yspI1NTV9z263c6/Xy2f38q703OaihP1sdCEtbd6fEhMTgDHGVOMp2JjVtfevv3Dw4EF6"
    "4cKFRw4ePPjDoqIiuba2VvwsFJDVaiVPPvkk27x5s2y3j4SrHbBmO/8fGeQ+Y4WQOekPBQDQjTfe"
    "+Gpe3qKWJUuWCADAAvknM72Pcw4tLS3C008/rZw48eETO3bs+H+9vb1halUt/9ySucyHxWLBVquV"
    "qGektLRU4ZwbDhw48J0//3nb0TfffHP59u3bFafTqVXLgBJCgFJK16xZC/HxcW8AANu+ffs1eeGf"
    "hfDn5eXlDACcAOjevpOHf67N1JgZpwCCwKKi4ggfd78Vu2rt+2qlHW6xYFReLn18Uq0EIUTHz9bk"
    "hE2OPt1zsZ4xz5iAkQKYMRCwBrxcBBcSlbTcXI1d1L02vGL9M9xqJWhafktyMpmNT4JW9tV+BgDg"
    "wAGLGsyBcM68h8C38orP0tksc4tFgCWGbV2HXTfEZOQ/PHruJNUgmWABgHEMlCIABKAlInLbR2Vv"
    "f7deUYQfFT766P1QnU2iSgt3nj59OiUyk/8hY17G1zvqPgTv2GWmQQwDYqAVqTDU1iAnG/FyGpX0"
    "4+YPPtiDVq4c559Rk4XPAhaLhZeUlAiJiYn/gBDsHh0dTd+3bx8lhBCV+RwsPKYqM0opjI2NiS++"
    "+JLscDhvdzjG3uOc34UQkm02i1BSUkH/WmP1dwkSioqKJJvN9oP6+vpfvfHGG/L4+LgQzKBQrWyt"
    "VgsmUzgGAJSUlPS3es0DDQ0NUafTaRofn1gKAHvAR2ig16KIKioqUHl5udTb255z7tz5FxoaGpm/"
    "ocPslj1GyOl0Mq1WF0cpLQSAk/47m3M2ZkRRZBqNhhNCgLHZPpsDIQRfvnyZNjc35yQnJ3/9gQce"
    "eMrvxSjXMt6qqirs90yV06dPr7Db7f9UU1Pz8I4dO9jw8DAOVLzBlNE0D/mq81TicxyUU6dOHViw"
    "IPt7779fwxVFgWAed7Cw7NDQMHnppZco5/wXu3fvFoqKiv41IGJxzYajWuylvLxc4ZwLNpvtf773"
    "3q7/u3//fnA6nWSm77oy/YGBsbkf8bKyMmy1Wll1dfXPenp6/tLU1MRcLhf2G1ZXGMnB5rqvr4/8"
    "8Y/PKzfeeNPPZFl+5NixY3cUFxefq6ysnCqQYbPZhKGhId7Y2Mj9qTGoqKhAeXl5qLGxEfkNgSkn"
    "sKam5tHXX6/6wfHjxzJqa2vh/PnzzOv1CoSQKXIqxphHRETwoqLCiSVLCn6NEJL8Cpz/NRUwfPRC"
    "DknL1/7fmRfXgisqANA0JqbqHSFUzmwNtvDRnq7fo+YGqtiHkEAoMMyAMwSEIgCOWGR6Cg5Lz5RY"
    "TMzP8z8atA++Q65oNcoPtYKgdbplikVCGPcV2RCJBvT6MKbRGo0AMMIZ+0j5VVTwiooK9q2yTb/U"
    "S9LXeF+7xjXSjTnniCECDGEAxoATERhD4qWmJjl/TXzZWO3+MxGlt/wb37VLi5YuneScPzRSf9iR"
    "sHhZudxxMW6i8xLlSCIuUQGOuDhw7rSsS5cWM9DtPH/+/B11dVs9nHPly6CEVZZsZWXlhXPnam8a"
    "Hi6xORyO1Lq6OooQIqrHNd26VwUTIQRkWQaHwyFWVVllzuErGOM3Oef3+DpBVQb23/wcvd4tYlHR"
    "ZhkApHffffeHJ06c+MW2bS/L/f39AiECYkwBzj8qXKEKFVEUwWQygSAInxt56JNCp9OIPo+QTBlC"
    "s+UjKaW4paWVr1w5kg4A0NjYyOZi7HHOUV1dnYAQkjnncO+9995RXf3+H0+cOBF/9OhR5vV68VzC"
    "0G63G0ZG7Nzj8Yxdyzibm5sRAIBer9eGh4fPmndWx04IAa/Xi44f/4Dm5uZ87/Llyy/Ex8dP2Gw2"
    "4WolAi0WC964cSPx1w6mnHPtqVOn7j579uwLR4++rz9wYD8dHx8nastAVfl9PBLw0Tq43W5wuVzy"
    "HBQwLSsrI4yxmqVLlzatXbs258CBAxQAyGwhbp+HTQEhAoODg6SqqkoJCwt74uBBm5icnPEiQqgd"
    "APCWLVvIXKpF+eoob6Z+pa20trYWvfnmmz9pbW25c+fOndDT0wPMJytnLRfpkw0UJicnwOl0Ep9I"
    "nj0dUFVVxaqqynF29o93pqennV68ePHimpoaRRAEwcfJmPkqpP//aHh4WNi5cyft6OiYl5+f/35V"
    "VVVvTEzMb3Nyclp1Ol17ZGRkRxBDZuqho6OjGQ6HI72trS2tqqrqx01NDfOPHDkCzc3NbHx8HBBC"
    "WF1b//VLQAix1atXi3Fxse8sXlxY57+rf21pls9aSAQNqZaUMADEEQp+CCoqKqC6uhpXbH4wcvji"
    "2bfQYH/xaHcv1YscK4iBghlwLACVCQ+PSoCwzAKYIPqvpmQVnVNzU1d8Q0ODZmiwLZd7HYiDm3NO"
    "QOAACkLg0ggQmRhLIDIiaAiW22wCyi9sHTx6oMywaOmugQ/HuIG5QOQKYM6BIwGAcgjjHDzjTqH7"
    "VC2NXrb8MefA2T9CfMEgt1hwVVU5Ki+venz41JFnTfqIKkWMWuToaWVcGUd6riCtMiFOXjolhxvI"
    "Wo0WdueUPFrq3w0IfQmEvlqmsqCgqO2dd95Zfddd99h0urDMmpoahTEm+JTvx3M0gR6lv0iH+MYb"
    "VYrDMXL70NDlmpqammeLi4utCKEJNV7n/9nPZMyqB1NeXs6LijbLnZ2dec3N539hsx3e+Pbbb/He"
    "3l6Rc7XUH5oSpqrhoHo1BoMBCBGu2XuYdqavarNyzuBa6lmIojgcGxsLgiCA2r93NjIQ51w4f/48"
    "nDvXsLm19fxz8+fn1lVWViIAwDabDassVLWaUHV1Nah5NACQnU5n/N69e57s6el+6PXXX4fz58/z"
    "8fFxPN0jCSwAErgnJicnoL+/DymKlGC1Wonb7SaBHqnVaiWxsbFoujHmrymOIiIiXhdF8fH4+PjE"
    "7u5OjhBCAd7HFZELtVJSX183O378WEZkZORuh8PxUERExKWrLVtlZSXzcyCMp06dLNu27aXHLl68"
    "uLympgYaGhqoLMtEDfuqnnjwsCgAxgg4Z8jpHANJ8kRZLBZcXV09lVMsmdakRW3lWVRUNHb69Ol/"
    "vXx56NW6ulPM6ZwASpWpcG5wpYemUmQDA33CCy88T++4Y+CnS5cue6ylpeXu3NzcY5s3b2abN2+e"
    "dfwAwNU6ypzzlJ073yrdv3/ftuPHj8Hhw4el8fFxUVW+s0RcptZBlhVwuVzgcDgcAAD9/f1IHf8M"
    "hjdvbLTy8nI0WV198JcrV96wvb7+tEIpA1mWAWO44kpksH2PEIKJCSc5fvwY7+hoN7W0NJsWLFiw"
    "9cyZ04AQGnjllW3vZ2fn6BCCA5GR0TsURaGDgwN3S5K84dy5M56nnvptsSBokrq6OqG9vQOamhoV"
    "h8NBfFsOXyHjBEEARVFoTk4uLFtWWB8REf3Qli1bRL+hA1+oAkafwLPpqK7WlJaWeoY/ePeHRvdY"
    "8aXzDV4tZ1pQZBCAA+cEvEQAISxCiczKxlxj3pRSVLrLZrEIgZ2CrP5Q9OWG2tVaYHf0DfVQJCgC"
    "QggIB/DKCtOYwrGTUxvlOofFdyeXT/9+brWSxtXrDxqr3zsUuSCvZLK1gWkkJ0HAgSIAzBmIwIEw"
    "hiaGhsBgH07xdIU/Fh5f8ER1SQkuL61UfOH0tU0AkNe9f9cPU1NSftF3+rgijfZjDQDWgiI42y8w"
    "gz5qyeDgoDY+Pl7mX6K7fGrLwk2bNnV98MEHN+l0un2EkPk1NTUKpYqgVkQKDElPF8Kcc3C5JoVd"
    "u96jHR3tRb29vc9futT2/ZqaI88UF6/ZihBSEEJTDbJVY2wOobOpe3/+ptoAANi/FyjnXHvkSPXj"
    "e/fu/t6pU6eSDxw4QO12O5mt16gaftbr9RAXFwcGg8FwLaHuwUGMvF4XUKrMQVH75kmWJVAUaS4k"
    "NX/XLPGnycnJ94SFheGRkREQRXGm0OfUegwNDcGOHTsUs9lcu2/fnl9t3Hjn//F6vby0tDRwjpWA"
    "3DEoihJ3+LDte6+++sp3m5qaDPv371d6e3uJv2TlFe9RhVPwMTLhwoUm6Orqstx//4NPI4S8jz76"
    "qDg6OsoWLVqEysvLlVmiMKigoKBt27ZtjrS0tKT29jau0cxc5MAnGAlQSnFNTQ1njK2WZan11Km6"
    "Z+Ljo55OSkq/6DdipgYgyzIBAN7c3LxmfNxx4x//uPWenp6egmPHjsHZs2cVj8dNCBGIyoxX5zWY"
    "ERDoBYqiKJw5cwYKCgp+UFFR8QzGuC+wVrDFYhHy8vK46i2VlZUxq9VKlixZ8nZzc/OHhYVFRQcO"
    "HKCEYMIYn/K0ZwvDyrIMvb095OWXX6anT5+Oa21tqXnqqd++lZu76ExJyYotGk1kvyRJRKPRUAAA"
    "SZIEABD9e0t35Ej1twYHBxc8++zv/0d9fT3U1tbyzs5ORinVBJ4TSukUKfPKlMNHtwlkWUbNzRch"
    "MzPzXo1GY9u6dau8devWqR+1Wq3IHwpmH6054larlaxbd+PbDofz2L33lhVbrVaKMSYIwVR6S733"
    "Hiz07y+HjgYG+nl/fx8/ceIDrtXqeHJyckJWVta9x49/ABjjO2RZ/s+PeBVeGBoahPb2DhgdHVVc"
    "LheSJGmqqUlgPj+AAAeRkZH8tttuF/Ly8n6zYcOGSfWq0rXK2S+cAFRbu0XMKCr1DJzcfad7oO/R"
    "ntO1igBejSTIQBgHwgTwMgzjRKtkXLdU9ETG2pLWrN/FrVaCZjjAjiH7RPiYAzFFAcGfRwDEgXFO"
    "zUYjRiC/l5SVNebr3Ruk1V9ZGStAiDLOb+/aM95F7H1R3mE30zCOffWkGSBgAIAAM5lcPt8kLYxJ"
    "qeg/efhEaWnpbl5bKyJUJKs1oNHNt/1yqPYgzVhb8quBc6fA2d0pgyITzrXIOymNh/vLIKIvWcjT"
    "f5dXWLlyZXtDQ8N6SumRyMjItF27dkkTE+MaURRBEIQZySaqwGKMkdOnT7MLFy5AUVFR3qpVq37f"
    "3Nz8w3fffaf29ts33hfYNqyyshIsFovQ39+PEhMTuUrIUT20pKQktHnzZiWIkmacc3N1dfX9zzzz"
    "9I8HBwdTjhw5AidPnlQwxoLq4c7kMaphdb1ejwwGAxNF8aTfG+OzeA5w7NgxXXFxsbu0tPT7k5Mu"
    "8+joqIwxFq8WMvV4PNDf3w8TExNSZWUl27Jliw4AlNkiAllZqYLDkYvNZjMfGRmZ0RMLJJXJsgzd"
    "3d3CSy+9BKWlJf9769ZnHzQYjMeWLl38WlZW9lEA0AwM9K6+eLGVDQwMgCx7fvof//GbxPHx8ehD"
    "hw7BhQsXmNfrFQRBAL1ez6OiopSRkRHR5XJNGS3BPBK/p4A6O7vo0aNHw0VR83JLS8tTCxYsmGof"
    "19p6YdXExGTK0qWF24OQgHBZWRk7efLkE/n5+a8dO/Y+JYRgda+phKBgxByXy4X279/P+vv70cDA"
    "wLd1Ot1DJpPJodfr0Kuv/hkwJuByueC557agsbFx5na7Yiml2sOHD0NDQwPzer1ACBEQ8j0/NzeX"
    "6nQ60tHRAWNjYzBbkRrVKG1uboY9e/YIkiTVvfrqK106nfaXWVnp2vz863YghLzT1otzzhlCyO1w"
    "OG7r6em53NHRgVtamjnGPq9/NmZ0oMHlcrnIiRMn+JkzZ1BhYeFdly8P3nX06NHvPPnkzz1btvwB"
    "Pfvs7zljDP7rv36HJEkCWZa5y+XCGOPErq4uOH78OPT19XH/+IggCBAZGUk558TpdMLVwuLqeT9+"
    "/AOIiYn91osv/ulOjPFps9n4QnJyslBYuNyqNvmY/oiysjKGEJI45ze5XK59vb29qw8dOsQIwVd4"
    "obORANV9xzlHkiSDJMnQ1NTEzp49y/zrhQVBUHP5jDHO/OIXi6IoqGmoYPsKwHfdzWw2K3fffbew"
    "fHnRhQ0bNlRxzjHG+BPdNPhCPS5usWBUWcn6Thy8RzvheK3/1DGRDg9wPVaQxD2gowIIXAdOpGW6"
    "nBxM5i+wGdYs2BRXXeWBIGQeq9VKysvLaefhfa962xu+5m2upzruJQwwaCkHF+NyxLIVIsnMqYy7"
    "4fZKsNnITB67mqds/+C9H0VMjv6866hN0rrdGl+zYAocyb62hVQEyg1Ml5gO4Tn59fqMtFtNOUXD"
    "gbk2dZz2+iOrqH3oV1qPp7jlbD1odVoIn1/gTV++LholJk5+WULQQYQhKS8vpx0dHYlnzpw5fPTo"
    "0QU7drxB7XY7+FmGsyq1wBAvY4wZDAaemppKli5dCrm5i86kpKT2OJ3On2ZkZEjp6ek0IyPj/NW+"
    "aWBgYPHg4CAMDg4q4+PjUZTSysuX+1N6e3uyDxw4CG1tbQqllCCEkGq9zuZJ+BnQPCEhAd17772u"
    "++67P26pL59/Rc6Uc46rq6txYChtcLAj8dChmt9arVX32Ww2hTEmXM0DppTy9PR0/v3vf3946dLr"
    "vrZs2TJb4I/YbDaiNhJXQ+tr1qzRHTv2/htVVVW37Nu3jwa+Z6Ym5oFXxjQaDU9NTUUFBYvh+uuX"
    "g8fjvsAY18XGxqYPDQ1BX18fNDQ0QHt7O9jtdtn/fOT/ZrZ27Vp88803w+uvv84aGxuxJEnAGAvq"
    "EQWONSYmhpeWlqLFixdznU5XTQjBLpebeTzu4pSUFK3RaHry9tvv+KG6z9R0Qnl5OX7qqad0L774"
    "4t5du94trq+vZz6PCAX1gtR1DvTmEUJKQkKCoNfrQRAEiIqKAkIIjI2NwcTEBLjdbpicnITx8XEZ"
    "fGxvrCiKn8VLICkpiX33u4/j3t4+eOONN9Rc6JQCDraX1Obxer0eEhMTITU1FfLz8yExMREcDkdd"
    "SkrKeGxs/OF58+b9Ii8vb8qYtFqtpKysjB06dOifT52q//Wzz/5eGRmxT7FqZztnwdIAnHOKEGJm"
    "s1k0Go2g1+un/r3X6wVJkoBzDm63G8bGxhT/dwvqGAAAYmJi+KZNm1B3dzecOnUK7Hb7rIZl4PeF"
    "hYWx1NRUnJGRAdnZ2RAXFwdO58TR9PR5LDY29uX16295wc8XZIHRpMrKSr53796w1tbWhldeeWVe"
    "U1MTB2BELbgyEwEu8BwEpsICr1BOT5kF/lkwbktg1MN/t1656667hE2bNjZkZWWXvvrqqyMVFRX8"
    "k6bRvjAP2E8NRwMPbLhbM2zfPnimnoDdzgycYc5lEBAAxwS8iqDEZuZiJTGzrgOjO1ej/InZiCSc"
    "c9Rte3ep4hoHxHz8ZwYcGOOAiQBIEEDmnCEAXusne8xAjmC1tbVibKruT3RA82Bcdn7ecMM5pqUe"
    "DH7/lyEKIsKgVTzYO9BLeUJs4cR42KGxsaZSqKgYVb0ZVFnJuNVK0LK1NY7Os3fQofFvmhh83xym"
    "DWN680uQkOD9pKzFv2Y4Oj09vb+2tvam22677Z9SUpJ+YrPZ4OTJk9THWAQ0XQioijfw9xhj7Ha7"
    "oaWlhbe1tbGYmMNLMjKylqSkpNze29sDra2t9Nlnn92KEOKEENDpdEin0wGlFCYnJ5GiKLJGo0k+"
    "evToPXa7HZxOJ3R2dkJHRwe0t1+C3t4ehRBCKKXC9PBY4LcEC5crisJ0Oj2kpKSEabXUDACTQYQd"
    "83vbGofDEXbqVO2/HDly4tEzZ87GnjlzhnPOhasVIvD/HRoYGIAdO3bEOZ3Od957b+cfsrIWPLlw"
    "4UIPQmgiUMEjhHhZWRmUl5dPHjx48Oe5ubm31NXVof7+fvDLy6DvEEUR1CsT/n6qqK2tjXd0dPB3"
    "393JIiMjczjnMD4+rqi5OwCOKaXIH570CQlBoGvWrCHr1q29lJeX94eRkZFf9ff3Q39//xVhuZmU"
    "0dDQENqxYwfds2cviYyMKFXnxm6306ysLOVrX7tvSWBvVXXMNpsNJSYmTu7du7djYKDvhnPnzvFA"
    "w2I6pt9X949d6O7u5qoADby/HliAhTEmBhppgkBYWlo6KysrE/Ly8rbb7SNf9Xg8giqoZwrFBp6B"
    "yclJaG9v521tbbympoZhjMFkMhYmJCTAN77xP0oLCgr+ghBq8SsdVl5eTsvKykhVVdVvduzYoZSV"
    "3fvbqqoqZWhoiGCM0UyGR5CIkzo/BACIw+HgDodjirwXqLwwxiDLMhBChECjgjEGZrOZrl69mmzY"
    "cMs7x49/sKm2tpYhf1hger4/cF3U309OTuLm5mbe3NzMDx48yBBCEB0dvSYtLQ0efPCBGwDgFYSQ"
    "BwKaIKv1CBBCk0eOHPnVI4888sxzzz2nXLjQRGVZJjPt90DDR71HPL12QTClqu6B6WunPitwn2i1"
    "Wrm8vFy86aabzqSlZd6ak5MzbLFY8KfhsAhfjPIFBFABPT33JHr7+9+ULjQAH+xhBiphAAUoBxBA"
    "BA+IXIiLhYjcPOwxRv56deENTu6/nxfsuWX+g9vx3p/HRSaD7FsVQAT7ylByAEQEAOwbduHszay5"
    "TqdD4fH5A131BzZoYpIPamKGFyj9XUyDEUacAiMcgCtAOAKReklncwNNnBdToDi8Caiy0s4rYCrH"
    "jMrLKbdaCUpbPAoAP+/tvbA1zhwpIGPCAHAO/Etey1UtLVlUVNQFAP/v2LFjLZmZGY8WFBSsev31"
    "12FgYIAihDClvkI2gYcy8BAEMhcVRSG9vb308uXLoNVqkV6vh/DwcBIbG/tNQRBAFEUQRdEvOBj4"
    "qpkBTE5OwuDgZSZJElIUhXs8HnC53FxRZIQxFmazzAOtX3+4ijPGOCGEZmZmiDfddDPMm5d6MiZm"
    "ntOfm+WqZV5RUcHPnz+fNzk5ueKVV155uLe3N6+3t9t05swZ6Ozs5GNjY2gm5TAdfsGHjh07xvv6"
    "+sKXLFn6L1FR0Y+sWLECvf3226/Fx8efzMjI2B4fH68anNRiseD169cfqaqyHsjLW3Szn6w0dZ9Z"
    "FbDTWbrTDCMkSRJCCOGenh7m9/SEQIGqRiv8HiQtLS0ld955Z1d6euYtq1evbrNaX/1aZmbmdf39"
    "/QwhhKeHZKeTY/w5R+J2u7ndPswCSVOdnZ1Mkrzr5s1LKi4vLz8a6AWXlJRQzjkaGhr6WW9v94Pz"
    "58+HS5cuXbGnghXGmF5JSTUOp3MWAvN6vqtOTPV02KJFefib3/wmTk1N/kFx8ZoXjx37oJxSyhlj"
    "KNDDms3wwBiD/8oWkiQJM8ZgYmKCulxubrfbkUaj4UEYwXTLli3i3Xff/Z9VVdsZAP/dG2+8SYeH"
    "h323MQK+eaac8PS/Rz5cUcVO/RlKKYiiOK0aHHCz2Yw2bdpISktLj+fn5z9fW1u3SZIk5udafKwo"
    "x0yNMDjniFKK/BXKoLu7WxEEAU1MuJyzRIiozWYT1q5d++zRo0fZww//4x9efvklqK8/PcUODzZ+"
    "dQ3VP5+NqR8YVg70ftX5U88SIQQIITwyMpLffffd4po1axsXLcq/Zf78+YOBe/VvSgFDtS/02/Z+"
    "4T9rHEMw1t2pmKgkIEZBQQAMieBlAsfR8TSpaKXg1Okeiy28YXuw+75TE2q1EigvZ8Pn3v+qxt6f"
    "3z42omgZJwAB5AkigKjVAmiuXhnPn5ehzc3N2nnZ2b39Hxz9ddwi+GPPqEOingmNiDkonALHDATC"
    "ADMJkETwxPAQg7gUX46nYtozfSXNEEAVRihnWN2gCCH+ZQw9ByPHcM5RdXU1KS4u3sY5f81ojLgt"
    "Li7uz/X19Ya6ujro6+sDRVE+1iJvOoMxgNlLfFWMXDA5OQnDw8Nw6dIlBSGkei1oWlhJLbAg+K35"
    "qRxZMIJGIDs7UOgyH52VJSQkiKmpqWj58uU4PT3jTFJSwq/1evfbsbGxU5EWhBBUVFTwF198UTsw"
    "MHC4s7MzqqGhAUZHR2F4eJC63W6sEpTUnPhclDDnHCRJQi0tLbyjo4NGR8dGHD16FCIjIzcXFxdv"
    "LiwsXAsAD1VUVBAAUPLy8pDFYsHZ2ZlP3HPPV2+220dYU1MTlmUZBQqfq10RUhWsTqfDgfMzPawZ"
    "FhZGN27cSFavXtOdmZlSsnLl6nar1Uri4qIqcnJyX6+rqwNFUTSBgitwzqcrYkIIwhgT9RqHP5dL"
    "GQP9xMRk+AxnEMfGxrZkZmb+5KGHHvrXZ555Runv7xeChV9nCsVPN3wCPPsrPD6fsBXoDTesJLff"
    "fnvXvHnznl+9eu1vWlpa8ihVkCzLfPr+ClaRbLrBpxpDhBAQRZHo9XomCALW6XRBN8nmzZtlfwrs"
    "Kav1NR4ebnxq+/bXoKend+oqYGCYeLrRM5d0xPQUUUCaiKamppCSktLJW2+99fjy5Vn39PePrQz0"
    "DIMZWDN5xIHz7FdsxHeXdnbdU1paqthsNmHNmjVbDh8+jL797W9/b/fuPQv37dsHbrebM8ZQoBJV"
    "xzWdVDmXcxhMPqmsd0IIXbJkCVm//ia2YsXyPfPnL3wgLS1t9LNQvl+IAuacI6ioYI5z72d5e7sf"
    "GG5poXrZiwlVgAkYZITBzQWOIxJY4pIiwRVm3Jyw8sat3GYTZmVYx8YiBMBHJDmRUDmMelwyCbgI"
    "zzkAwhg4wgBzDPX6rzdRm80mJK5c83xn9c4liUuXPN7/4QmZMCYyRIALCBjHIHIEoqKA6PFgQfFG"
    "zKbY/Uxd5LcmKfwNwf/9ip/1JwHAW7W1tWvT0tIfLS5edXdjY0NcbW0tNDU1MUmSmJ/Gj6cLpUDC"
    "SqBg8P+dELzPrS9Jrv75dGJVoGerhgdVz4wxxmVZpgghbDKZ0PXXX08KCgpIVFTUWHx8/OWYmLh/"
    "27Bhw8tXVFPz//qnP/2pgBBSDh48+KMzZ85EVVVVeSYmJjSiKCCMEQko+/gxIX81Mpb/+5CiKEJv"
    "by/v7e0Fj8cjDwwMoPnz56cFSwUsWVJ0tLq6+uH77//a83/604u0sbEREEIkMNx+tXermB6SZIxx"
    "QRDowoULyc0330xWrVrdEB+fvnHZstwOi8UiNDY2svLyyp0vv7ztYnd3V8GePXsUjUYjKIrysZza"
    "TMrP7/1zvV4vb9q0UZOYmHCYEE29r1NYGZse8veHaP9t9+7d/Bvf+Ma/bdu2Te7r68PqmAMNq7mC"
    "MQ7qj1NKOaWUxcbG8nvuuUcoLLzu5JIl627LyUketlgsOCIigno8HnC73TN6Uh/t4Y+ut/kUAAZF"
    "USillKveo8lkwkZjOLhcLmUmHk5paamybt06obz8a/+1c+fb3GQyPrFv3/6I48ePM0mSOCGEBBqV"
    "gWclmHc4/e8DjVL/M7hGo2GrV68m69evH8nOXnBkw4bb7gYAOHOmTu92u0GSpI8VAgmmcKdHm2RZ"
    "pn7DGQMgiI6Oxkaj2d/vdebIX2lpqWKxWIR169b9gXO+jRDx9cTExFuqq6tJU1MTVRQFiaKIg3m0"
    "M+3Bmc6Daryrtbi9Xi+NjY3F69ffRNauXTOyYMH8b69atfa1AKfpM5Hbf3UFrHZP6l9f/DgeHY2j"
    "Q8PUCAwDMJARBoUIXDRG85glhUQ2mh9JXLnh+ebmXVqUXeqdy/MVj9tFJ8Y55hQIMEDc14pQ9S8J"
    "IeBReNi1fHNJSQmzWq2EpMY9TbRwX0xaRtz4pRbGiIBloL7mExyBXuEsXOFEHmE/A4DbIC9PJU+Q"
    "6QumMh/hbxR+6w9ZrVZcVFR0CgAe45w/brMd/Mfrr1/+h1On6nF7ezu+ePEi2O12P73fixDCQAjB"
    "hGBQWaZX81w+Okg8qEAJVO6+XsaM+4UF02q1oNNpISYmVsjIyBDi4+MhNzcXjEbTrtzcnMbly5c8"
    "o9NFdE0nwgTL6xBCXIIgKAaDQZQkCfvyTT45ooZ+GaOAMZnToZ8+Zj97E/R6PdFoNJgx5popH19S"
    "UvLC3r272SOPPPynPXv2wrFjx2Sn00kIIYgQEkTy8I9xL6fNG9NqtSwtLU247rrrhBtuuAFychY8"
    "fsMNa3+PEOJ+JaioJfuWL19y3/j42MHR0dHEkydPKBgTIVAZqHnSwHKZfmHMRVFkGRkZ+K677tQs"
    "Xbq05p57yu5ACM3I7fCXMtV85Stf+fl77+1UHnts85Nvv/0OnDt3TvF4PKoiQtP3ijq9wTw2hBgw"
    "xjlCiCUlJZFly5aRlStXwuLFS+qLipbfZjKZhhsarJr8/HLpO9/5Dni9XvB6vVMM2emelqIo6lnm"
    "Go2GazQa0Go1oNXqICUlWfB1OzKC0WiE+fOzZKPR9PWEhIQOq9WKZ/KkDh8+rJSVlZGNG+98mnP+"
    "mlarfz03N3ddbe1JOH/+guJyuTDnfGq9g0chANSKUoHhZ5UIyBjlhBA2f/584fbbbyMZGVlvP/zw"
    "I//Q3d2t/O53v9N+97vflRoazjBJ8oIsS36jV/UY+dSvKaXMJ87Q1Ph1Oh3o9XpISUkREhMTwWg0"
    "QmRkBMybl2YXBHEjAEhqE55Zom6qse8CgNuqq6tLs7Iy3zxz5mzEyZMnobm5WZEkL+YckCiKaHre"
    "VxWxqvEeTN9zzoBShSsK5Xq9nqWkpKD8/HxSXFwMqakpb95111cfQghNzCYb/mYUcGVlJQcAkNyT"
    "98iX+7mOSoggGaiAYZJjLpiiWUL+dcRrivrHlOJb/8QbGjQoO39W5csBEJSWUs65frBm74/HBgcR"
    "YowwTABxBsA4UECcaDRIUqhb0JAGAAAYGprTRCKE1ApCzYMndt2qTcnYp52ciPSMdjENUEyAAUMM"
    "EEeAPAooXnm6B4xm8Sb/lsFVhVBeXs7914m2dHR07Fu4cGFBc3PzN8+fv5gnCEJqZ2cHXLp0Cez2"
    "Ebh8eYA5HA7KGAdBEAS1wILPcwjOSPQVrUDBFBiXJEnBGINWq4X4+HgxIiIChYeHQ3JyCo6KioTI"
    "yEjQ6fQsKSlp37x5qcdKSm78CwBcCpx/f8iPBhOGeXl53L9eAzqdTtDpdLLBYGA+r8f3zT7hN/cu"
    "MNO3hE8w+srcud1uajKZdBqNJmIm48f/vS9++OExHh0d+91FixZdd/r0GWhvb4eOjnaZMebn1aCp"
    "96nz68ulc2CMUgDE4uJiyaJFi/DSpctwQkKCJysr842CghxLWlp2WwAzVa2QxqxWK8nIWHj+8GHb"
    "UxiTxw0GQ9LJk7XU6XQiAJ+X8xHBxXdCGeNMFEWalTVfvO66ZSQ/v8CzYsXyN1avXvcwQuhqJERe"
    "WVkp+cf8q+PHj1ODwfjIhQvnc06dOgWXLrXDwEC/jBAmvowEnqZ0VFIOAOeMM8aU8HAjLFiwQFy8"
    "eDExmcxDq1evqs3PX/SKKOreNJlMbs45rqioUALXR2VHM8aBMcYURaaCIIDBYIDk5GQxJiYGoqNj"
    "wGw2g8lkhIiICIiKiobJycna/Py84fBw04cFBQUvc+5mkZGJ7Q888ADAVUqDVlVVUb8CGuac33L4"
    "8OFbU1NTflBcPLj6woUL0N7eDl1dXfLExAQQIggIAVKjPwjhqdrMairOPxdUq9WxpKQkMSsrC+Xl"
    "5ePExMTR665bVrV2bcm3VGfBXw+b19fX8okJF5NlCgAquQ1xWZYVrVYDJpMZxcfHCxEREZCUlAhm"
    "cwQYjeFgNkeA2WwCj8dbvWzZEik83PhGQcGSgxMTE+7Y2Ni+IJbhjMa+X/5CSUmJ7eLFi/nz5qV9"
    "Jycn5+u9vX0pFy6ch5aWVujt7ZFdLhcAAMGY4ODcDzZlmPm8XapotTqekpIspqamotzcHJyZmQkZ"
    "GZmHVq5c/kRMTMJhf7GUzyTk/OXIAQMAaDDGWsJlYIrEKPYA5jwmBkcvXk68puiHU1bd+iduswhI"
    "bU94NVYXQpyPjenA7U2Ux5yAGCAZA2CEAVEOnAPX6A2Cc9LdPX/Jmr+oOdlrCb36Pdn6yx8cvUmk"
    "ntOowS7DhIdRRmECc5AJ5kZRSzhi/wEAAI2NquBW4O8YgRvTYrHg9PT0dgBoB4B3RkZG5o2MjORf"
    "unQp/vLlyxXnzp1TnE5n5ujoKHY4HDAyMgJO5wTIsuQvDUlBlmXOphX/FUVBEAQyFeoTBAEIIWAw"
    "GJDJZBJ9QjAFOEeO2NgYe15evmg2RzyFsdAUExOh12q1A7m5uccCn1lbWysWFRUpAMBnK41ZVuYL"
    "i3o8nt1RUVFt999/f9bHPVh0tYjazOZjgIXuFxgi+ArS/9Jfm5cFC89ZrVayYkXxS5zzVxYsyC5b"
    "vHhpZV1dfczly/2Rra2tYLePAICva4vXK3FKFarRaIher0eCIEBERCRJS5tHMjMzYeHChaeWLVv2"
    "hCAIHQsXLjwdqHinl3FUja5160p/efr06e1Go+n1efPSrzt//jz09PSAer+UEAFpNCJotVqIiorC"
    "eXl5ODs7e2zhwoXWmJikZ6+/fml9QEjvqtEgdcw33HDDv4+MjPyxra3t/oKCpT84ceLD2JGREWNn"
    "ZyfY7XZQFAqyLDGv18s456DT6QR//hUMBgNKSkoSc3JyID8/fygjI+P5lJSUF+bPn98SmCZTQ98A"
    "ALIsI4QI6PUGJSIiQhRFDUREmHF0dDSOiIiAtLQ00Ov1vcuWLZM0Gs1et1veaTKF6SIiIpTo6Gia"
    "lpa2L/Cu+3Sj5hoUkAQA7yCE3jl+/PidK1ZctrS2tiZ2dXUltLa2wuDgIExOTvrX2wuUUsXPhhdU"
    "MqNGo4XY2BiSnZ1N0tMzIDk5paawcNkvCNFfWrgw47yfsHZFJIIxpAsPN2JR1HhjY2MFjUYDBoMB"
    "xcfHi7GxsZCcnAyiKHbk5uYio9H8J0rlk0ajUWcwmGhYmFbKy8vbPX1M1zL+QGfF376wFwB+1NbW"
    "+6yiTBS0tLRubm9v/0pbW6vY3t4Og4ODMDY2Dr7GIQp4PJ7AQjNYFEWs1WpBq9VCWlqakJ6eDuHh"
    "4a5ly5ZdjomJ+Z3ZbD6/bNmyfVfSd8o/l1ThX515q058x7H9v4nwOv/Z2XYBZNc44LBw0CSlAUTF"
    "PpS08tZtvLZWREVFc+qurW4YzrnQ8+ZLQ47GDyOEyRGOOENYIACMAVOAaRNTMJ+f05Gx6Z9yAED6"
    "JB4ot1oJuq+c9hx462WjNP710e52QEwBLuqAafQQlpBGTZlpmWFZRV3q/V/4b4aAyk54WpFzjDFm"
    "NTU1D8qyPO/SpUsUY/y4y+VK7O/vZ7IsY0mSOGNM0Ov1V+SUPB4XEEIUURSRXq8Hs9nMzWYzIoR0"
    "KoryXFRUFE5NTWWiaKoqKMhuCxb2n1Zx65rv7lksFvyzn/2MffDBB48RQiICWbGfIUeCRkVFEbvd"
    "fmLVqlWHriaoplvmFy9ezBwYGCjv6Oi40e123zg4OMgmJiaQ1+sVwsPDQZZlCA8Pl1JSUgjn/M30"
    "9PRTZrO5q6io6C/T5+lqAjLw3YcOHbpvbGzsW/39/Te0tbWJkiQBxliJiIjgmZmZWK/XH09LS3s3"
    "MjJyZ3Z2dlNg1AGukYA4fcznzp1b4nA4vtLd3X332NhYYVdXFwUAjcfjAUVR/IpHI0dHR2Oz2Txs"
    "NBp/m5ubi1JTU7epnpjFYhH8hs7UvlC98tbW1vz9+/efPnv2LAkLC+MpKSk0LCysLjw8/K3o6GgS"
    "GRlJr7/++j9gjB0zRUB81a+aeGPjoqnrNp+IPwMVCKHKqTrlDocjqrm5eXNbWxtzuVyPeb3e1K6u"
    "LuZ2u0W9Xg8AvrrUer1eMRqNPD09HWOMt6WnpzeLonihqKjorcBvDDyv6jnetGlTps1m29PT05OF"
    "MaapqakQGxsri6L4m7i4uEmDweBevnz504IgUBq8ByH23yJgauW6yk8hFwNKzk697OzZs+s456su"
    "XrwoO53OuxhjK3p7e+no6Kig1+uxeq3I6/UCAChxcXEoOTkZ9Hr9rxYsWDCOMT6Un59/IsDoJlar"
    "lQfuh78LBRy4sQfrj/5E9EwWjA5dlqMTkjQOBd5LK775peZdu7TZt93mvbaNCeD48Eip1Nm0c/Ds"
    "cX2Y4gVKKQKCABADRUFMl5qB0fy89vTbHloIAJ+o+QEHQFVlVlxeVU6HThz8d0GWU8ado0zUGzAn"
    "2omw5LgneifFvvy5eO7/TZRxSUkJHhoa4sGsyLGxseixsTFdV1cXp5QSj8fjFQTh1oiIiFudTifl"
    "nKOwsDA8Pj7+DGOsWaPRaLRaLdfr9ZCamgrR0dEuhNBokH3Ny8rKyKJFi1BJSQnM9P4vdXx/jnfD"
    "gwkkzrlgt9vjW1tbEaVUkmW5MDk5+R96e3urTCbT+zk5OTqTydSt3hNW8/nTIxpzebfK6+CcawHA"
    "fPjw4YWU0lRFUQ6YzWZNQUEBDwsLG1YrQPmFMXyae+8zjDl8ZGQksrGx0cMY+0eNRrPEH7l4Xpbl"
    "c0lJSdro6GhPUlLS0HQlM4tCQDabjVBK4wHgSVEUn8nOzm5NSEhw+WuaB1U2AfwRKCkp+aw7gCGL"
    "xUIqKyuvMF4451EtLS1h7e3tEsY4Kzw8/HFFUUCSpCd1Ol0/QkhcsmQJMhgMPYHP8lcd48HWQ3Vu"
    "3nzzzazo6OgnEEIVUVFR40ajEaelpfXNYHxjdexq5OLzki15eXnI340vcB70drs9qqWlRZmYmDCI"
    "ovhTzrloMpnQ+Pi4TZKkt5OSkrQRERE8NTW1Z7pxF6w++d+VAg5c2E8qdK74N1u2iGjzZrln/xtb"
    "jaN9/9Rb+75sACYyAFA4A0QYeClhutRsTBbkd6fden/apz0QX6YWgn9rsNlsgtrxRi0C/2mwZcsW"
    "MTs7mxuNRrRz505a+TlHHWw2m2A0Gj+Xs1NXVweFhYVw6dIl9kkMBs452rp1qzDXeVXn7tMKnOne"
    "UzDU1taKn8f6WCwWnJSUROY6ZnX9CgsLP1UHMrXBwGfxrE8jh+rq6oS6urprOUuotrZWmOsem0nW"
    "WSwWQW3X+Vmc408Dq9VKMjMz8TXOw9S/KywspF9EIaQvthQltxKojkUA1QBQAlBSzWbqmDSr91tR"
    "gQbLyuI8fRffdV2ouw71XeIaRcIIi0C5AgQr4OV6pk1fhLW5Sy8mr78rh7NPP9fcaiUQGzs1h9VQ"
    "DSUllVd0OwlhbtGLQFRVVeHR0dEr7pT09fXRioqKoPMaMoTmNK+orq6OXLp0iak5bXRt7Zvm/E61"
    "z6r6nr/WOgWOuaKigqjKYfre+USRL79hE/isL9u+mzZ+lJSURP47jf9q8wAAEBkZyQL35Rc9BvR3"
    "IWxsNgGMRtQ11rsNei+WOy+cplpKRQQEGJdBwBRcVK9ELrqe4HlZ9yWuv6Pq82K1hRBCCCGEEMJc"
    "gP8uRlFSQlFRkRyRlvpTMMUgWRcuSIhwmSPgRABFEGGSA+d6I1J0Bg4AEBvguYYQQgghhBBCSAF/"
    "Ui/YaiXm+YUtOCz8mbjUTFAwkWQAkDinHs6kyKRUQRK1bTqd6ZjVWkbUZuQhhBBCCCGEEFLAnxAV"
    "FRWocREQAA6ayOgTSXmLUUxGthaFGxVDXCIxp2dr4nMLkCkm9j9jCwv7yqAslDcMIYQQQgjhC8Xf"
    "VRiWWyzY/uCD4QY0eZ3Tbv+hx+XZ4HW730eGsA6d3vhEyoq1zSH2cgghhBBCCCF8zuhoOnu7lfMp"
    "Blzg/bwQQgghhBBCCOGz9II5R5xbMLdaiU/pAuY2mxDsuksIIYQQQgghhPA5wOpXwiGEEEIIIYQQ"
    "QgghhBBCCCGEEEIIIYQQQgghhBBCCCGEEEIIIYQQQgghhBBCCCGEEEIIIYQQQgghhBBCCCGEEEII"
    "IYQQQgghhBBCCCGEEEIIIYQQQgghhBBCCCF8CfH/ATJc2Hd/Ew/8AAAAAElFTkSuQmCC"
)

# Decode once at module load — these are reusable BytesIO factories.
def _logo_bytes(b64: str) -> bytes:
    """Decode a base64 logo string to raw PNG bytes."""
    return base64.b64decode(b64)

PICKLERICK_LOGO_BYTES = _logo_bytes(_PICKLERICK_LOGO_B64)
CLAUDE_LOGO_BYTES = _logo_bytes(_CLAUDE_LOGO_B64)

def get_picklerick_logo_image():
    """Return a BytesIO of the Pickle Rick logo for st.image()."""
    return BytesIO(PICKLERICK_LOGO_BYTES)

def get_claude_logo_image():
    """Return a BytesIO of the Claude logo for st.image()."""
    return BytesIO(CLAUDE_LOGO_BYTES)

# Backwards-compat shims so existing code that references LOGO_PATH still works.
# The page_icon needs a path or PIL Image, not BytesIO — write the logo to a temp
# file once at startup so st.set_page_config can use it.
import tempfile as _tempfile
_logo_tmp = _tempfile.NamedTemporaryFile(suffix=".png", delete=False)
_logo_tmp.write(PICKLERICK_LOGO_BYTES)
_logo_tmp.close()
LOGO_PATH = _logo_tmp.name
CLAUDE_LOGO_PATH = "embedded"  # truthy — UI code only checks for truthiness

st.set_page_config(
    page_title="Pickle Rick — Your Pickleball Rules Assistant",
    page_icon=LOGO_PATH if LOGO_PATH else "🥒",
    layout="wide",
    initial_sidebar_state="expanded",
)


# =============================================================================
# THEME — Creme background, black text, vibrant pickle-green + court-blue accents
# =============================================================================

# Pulled directly from the Pickle Rick logo:
GREEN       = "#22C55E"   # Bright pickle green — primary accent
GREEN_DARK  = "#15803D"   # Deeper pickle green — header/text accent
GREEN_LIGHT = "#4ADE80"   # Light pickle green — hover/highlight
BLUE        = "#1E56A0"   # Court / paddle blue — secondary accent
BLUE_LIGHT  = "#60A5FA"
CREAM       = "#F5F5DC"   # Spec'd creme background
CARD        = "#FFFFFF"
BORDER      = "#E5E7CB"   # Soft creme border
TEXT        = "#1F2937"   # Black-ish for body text
MUTED       = "#4B5563"
AMBER       = "#92400E"
RED         = "#991B1B"

# Subtle pickleball-court-inspired SVG background pattern
_SVG = (
    "<svg xmlns='http://www.w3.org/2000/svg' width='140' height='140'>"
    "<rect width='140' height='140' fill='none'/>"
    # Court sidelines
    "<line x1='0' y1='70' x2='140' y2='70' stroke='%2322C55E' stroke-width='0.4' opacity='0.10'/>"
    "<line x1='70' y1='0' x2='70' y2='140' stroke='%2322C55E' stroke-width='0.4' opacity='0.10'/>"
    # NVZ rectangle suggestion
    "<rect x='40' y='55' width='60' height='30' fill='none' stroke='%231E56A0' stroke-width='0.35' opacity='0.07'/>"
    # Pickleball with holes
    "<circle cx='70' cy='70' r='6' fill='none' stroke='%2322C55E' stroke-width='0.5' opacity='0.13'/>"
    "<circle cx='68' cy='68' r='0.5' fill='%2322C55E' opacity='0.20'/>"
    "<circle cx='72' cy='68' r='0.5' fill='%2322C55E' opacity='0.20'/>"
    "<circle cx='70' cy='72' r='0.5' fill='%2322C55E' opacity='0.20'/>"
    "</svg>"
)
BG_URL = "data:image/svg+xml," + urllib.parse.quote(_SVG)

# ── CSS Layer 1: Buttons, selectbox, sidebar, dark text ──────────────────────
st.markdown(f"""
<style>
    /* Buttons — creme bg, green border, dark text */
    .stButton button, button, .stButton>button {{
        color: {TEXT} !important;
        background-color: #FAFAF0 !important;
        border: 2px solid {GREEN_DARK} !important;
        font-weight: 600;
    }}
    .stButton button:hover {{
        background-color: {GREEN_LIGHT} !important;
        border-color: {GREEN_DARK} !important;
        color: {TEXT} !important;
    }}
    .stButton button:disabled, .stButton>button:disabled {{
        background-color: #F1F5F1 !important;
        color: #94A3B8 !important;
        border-color: #94A3B8 !important;
    }}
    /* Download buttons stand out a tad */
    .stDownloadButton button {{
        border-color: {BLUE} !important;
    }}
    /* Selectbox + multiselect */
    .stSelectbox > div, .stMultiSelect > div,
    .stSelectbox > div > div, .stMultiSelect > div > div {{
        color: {TEXT} !important;
        background-color: #FAFAF0 !important;
        border: 2px solid {GREEN_DARK} !important;
    }}
    .stSelectbox label, .stMultiSelect label,
    [data-baseweb="select"] span, [data-baseweb="select"] div,
    [data-baseweb="popover"] li, [data-baseweb="menu"] li {{
        color: {TEXT} !important;
        background-color: #FAFAF0 !important;
    }}
    /* Sidebar */
    [data-testid="stSidebar"] {{ background-color: {CARD} !important; }}
    [data-testid="stSidebar"] .stMarkdown,
    [data-testid="stSidebar"] label,
    [data-testid="stSidebar"] .stSelectbox,
    [data-testid="stSidebar"] .stTextInput,
    [data-testid="stSidebar"] .stMarkdown h1,
    [data-testid="stSidebar"] .stMarkdown h2,
    [data-testid="stSidebar"] .stMarkdown p,
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] span,
    [data-testid="stSidebar"] div {{ color: {TEXT} !important; }}
    [data-testid="stSidebar"] [data-baseweb="select"] > div,
    [data-testid="stSidebar"] [data-baseweb="select"] span {{
        color: {TEXT} !important; background-color: #FAFAF0 !important;
    }}
    /* Tabs */
    .stTabs [data-baseweb="tab"] {{ color: {TEXT} !important; }}
    .stTabs [aria-selected="true"] {{
        color: {GREEN_DARK} !important;
        border-bottom: 3px solid {GREEN} !important;
    }}
    /* Dark text everywhere */
    .stMarkdown, .stMarkdown p, .stMarkdown li, .stMarkdown h1,
    .stMarkdown h2, .stMarkdown h3, .stMarkdown h4,
    .stMarkdown span, .stMarkdown strong, .stMarkdown em,
    p, span, label, h1, h2, h3, h4, h5,
    div[data-testid="stMarkdownContainer"] p,
    div[data-testid="stMarkdownContainer"] li,
    div[data-testid="stMarkdownContainer"] span,
    .stChatMessage p, .stChatMessage span, .stChatMessage li,
    [data-testid="stChatMessageContent"] p,
    [data-testid="stChatMessageContent"] span,
    [data-testid="stChatMessageContent"] li {{ color: {TEXT} !important; }}
</style>
""", unsafe_allow_html=True)

# ── CSS Layer 2: Inputs, alerts, expanders ────────────────────────────────────
st.markdown(f"""
<style>
.stRadio label, .stRadio label span, .stRadio label p,
div[data-testid="stRadio"] label span,
div[data-testid="stRadio"] label p,
.stRadio > div > label > div > p {{
    color: {TEXT} !important;
    font-size: 0.95rem !important;
}}
.stChatInput textarea, .stChatInput input {{
    color: {TEXT} !important; background-color: #FFFFFF !important;
}}
.stTextArea textarea, .stTextInput input {{
    color: {TEXT} !important; background-color: #FFFFFF !important;
}}
[data-baseweb="select"] span, [data-baseweb="select"] div {{ color: {TEXT} !important; }}
.stAlert p, .stAlert span, .stAlert div {{ color: {TEXT} !important; }}
.streamlit-expanderHeader p, .streamlit-expanderHeader span {{ color: {GREEN_DARK} !important; }}
.stCaption, .stCaption p {{ color: {MUTED} !important; }}
</style>
""", unsafe_allow_html=True)

# ── CSS Layer 3: Full theme styling ───────────────────────────────────────────
st.markdown(f"""
<style>
html, body, [data-testid="stAppViewContainer"] {{
    background-color: {CREAM};
    background-image: url("{BG_URL}");
    background-repeat: repeat;
    color: {TEXT};
    font-family: 'Inter', 'Segoe UI', sans-serif;
}}
.main .block-container {{ background: transparent; padding-top: 0.5rem; max-width: 1100px; }}

/* Sidebar */
[data-testid="stSidebar"] {{
    background-color: {CARD}; border-right: 2px solid {GREEN_DARK};
    box-shadow: 2px 0 8px rgba(34,197,94,0.10);
}}

/* Hero header */
.pr-hero {{ text-align: center; padding: 1.5rem 1.5rem 1.0rem 1.5rem; }}
.pr-hero-title {{
    color: {GREEN_DARK} !important; font-size: 2.8rem; font-weight: 900;
    letter-spacing: -1.0px; margin: 0.6rem 0 0.2rem 0; line-height: 1.1;
    text-shadow: 1px 1px 0 rgba(0,0,0,0.05);
}}
.pr-hero-slogan {{ color: {MUTED}; font-size: 1.05rem; font-weight: 500; margin: 0 0 1.2rem 0; }}

/* Tabs */
/* Tabs — force creme background with high specificity to defeat Streamlit defaults */
.stTabs, .stTabs > div, .stTabs [data-baseweb="tab-list"],
[data-testid="stTabs"], [data-testid="stTabs"] > div,
[data-testid="stTabs"] [data-baseweb="tab-list"] {{
    background-color: {CREAM} !important;
    background: {CREAM} !important;
}}
.stTabs [data-baseweb="tab-list"] {{
    border-bottom: 2px solid {GREEN_DARK} !important;
    border-radius: 8px 8px 0 0;
    gap: 2px; padding: 0 0.4rem;
}}
.stTabs [data-baseweb="tab"] {{
    color: {MUTED} !important; font-weight: 600; font-size: 0.95rem;
    padding: 0.55rem 1rem; border-radius: 6px 6px 0 0;
}}
.stTabs [aria-selected="true"] {{
    color: {GREEN_DARK} !important; background-color: {CREAM} !important;
    border-bottom: 3px solid {GREEN} !important;
}}

/* Cards */
.pr-card {{
    background: {CARD}; border: 1px solid {BORDER}; border-radius: 10px;
    padding: 1.2rem 1.4rem; margin-bottom: 1rem;
    box-shadow: 0 2px 8px rgba(34,197,94,0.06); color: {TEXT};
}}
.pr-card-green {{
    background: {CARD}; border-left: 4px solid {GREEN};
    border-radius: 0 10px 10px 0; padding: 1rem 1.2rem; margin-bottom: 0.8rem;
    box-shadow: 0 2px 8px rgba(34,197,94,0.08); color: {TEXT};
}}
.pr-card-blue {{
    background: {CARD}; border-left: 4px solid {BLUE};
    border-radius: 0 10px 10px 0; padding: 1rem 1.2rem; margin-bottom: 0.8rem;
    box-shadow: 0 2px 8px rgba(30,86,160,0.08); color: {TEXT};
}}

/* Quiz cards */
.quiz-question-card {{
    background: {CARD}; border: 2px solid {BORDER}; border-radius: 12px;
    padding: 1.5rem 1.8rem; margin-bottom: 1rem;
    box-shadow: 0 3px 12px rgba(34,197,94,0.10); color: {TEXT};
}}
.quiz-question-text {{
    font-size: 1.05rem; font-weight: 600; color: {TEXT} !important;
    line-height: 1.55; margin-bottom: 0.5rem;
}}
.quiz-result-correct {{
    background: #F0FDF4; border: 2px solid {GREEN}; border-radius: 8px;
    padding: 1rem 1.2rem; margin-top: 0.8rem; color: #14532D !important;
}}
.quiz-result-wrong {{
    background: #FFF1F2; border: 2px solid #F87171; border-radius: 8px;
    padding: 1rem 1.2rem; margin-top: 0.8rem; color: #7F1D1D !important;
}}
.quiz-explanation {{
    background: #ECFDF5; border-left: 4px solid {GREEN}; border-radius: 0 8px 8px 0;
    padding: 1rem 1.2rem; margin-top: 0.8rem; font-size: 0.92rem;
    line-height: 1.65; color: {TEXT} !important;
}}

/* Pills */
.pill-ok {{
    display: inline-block; background: #DCFCE7; color: {GREEN_DARK};
    font-weight: 700; font-size: 0.78rem; border-radius: 20px;
    padding: 2px 10px; border: 1px solid {GREEN};
}}
.pill-warn {{
    display: inline-block; background: #FEF3C7; color: {AMBER};
    font-weight: 700; font-size: 0.78rem; border-radius: 20px;
    padding: 2px 10px; border: 1px solid #FCD34D;
}}
.pill-err {{
    display: inline-block; background: #FEE2E2; color: #991B1B;
    font-weight: 700; font-size: 0.78rem; border-radius: 20px;
    padding: 2px 10px; border: 1px solid #F87171;
}}
.pill-blue {{
    display: inline-block; background: #DBEAFE; color: {BLUE};
    font-weight: 700; font-size: 0.78rem; border-radius: 20px;
    padding: 2px 10px; border: 1px solid {BLUE_LIGHT};
}}
.pill-green {{
    display: inline-block; background: #DCFCE7; color: {GREEN_DARK};
    font-weight: 700; font-size: 0.78rem; border-radius: 20px;
    padding: 2px 10px; border: 1px solid {GREEN};
}}

/* Misc */
.streamlit-expanderHeader {{
    background-color: #ECFDF5 !important; color: {GREEN_DARK} !important;
    font-weight: 600 !important; border-radius: 8px !important;
}}
.pr-footer {{
    text-align: center; color: {MUTED}; font-size: 0.78rem;
    border-top: 1px solid {BORDER}; padding-top: 1rem; margin-top: 2.5rem;
}}
.pickle-log-card {{
    background: #ECFDF5; border: 1px solid {GREEN};
    border-left: 4px solid {GREEN_DARK}; border-radius: 0 8px 8px 0;
    padding: 0.9rem 1.1rem; font-size: 0.88rem; color: {TEXT};
}}
#MainMenu {{ display: none !important; }}
[data-testid="stMainMenu"] {{ display: none !important; }}
footer {{ visibility: hidden; }}
[data-testid="stHeader"] {{ background: transparent !important; }}
[data-testid="stAppDeployButton"], .stAppDeployButton {{ display: none !important; }}
[data-testid="stToolbarActions"] > *:nth-child(n+2) {{ display: none !important; }}
[data-baseweb="select"] {{ background-color: {CARD} !important; }}
.stAlert {{ border-radius: 8px !important; font-size: 0.88rem !important; }}

/* Inputs / file uploader */
.stTextArea textarea, .stTextInput input {{
    background-color: {CARD} !important; color: {TEXT} !important;
    border: 1.5px solid {BORDER} !important; border-radius: 8px !important;
    font-size: 0.92rem !important;
}}
.stTextArea textarea:focus, .stTextInput input:focus {{
    border-color: {GREEN} !important; box-shadow: 0 0 0 3px rgba(34,197,94,0.15) !important;
}}
[data-testid="stFileUploader"] {{
    border: 2px dashed {GREEN} !important; border-radius: 10px !important;
    background-color: #F0FDF4 !important; padding: 0.5rem;
}}
.stChatMessage {{
    background: {CARD} !important; border: 1px solid {BORDER} !important;
    border-radius: 10px !important; margin-bottom: 0.5rem;
    box-shadow: 0 1px 4px rgba(34,197,94,0.07);
}}
</style>
""", unsafe_allow_html=True)


# =============================================================================
# SESSION STATE INITIALIZATION
# =============================================================================

def _s(key, default):
    """Initialize a session state key with a default if not set."""
    if key not in st.session_state:
        st.session_state[key] = default

# Chat
_s("messages", [])
_s("uploaded_files_content", [])
_s("api_key", "")
_s("model_choice", "claude-sonnet-4-6")

# Video Analyzer
_s("va_frames", [])
_s("va_frame_count", 0)
_s("va_video_name", "")
_s("va_fps_used", 1.0)
_s("va_analysis_result", "")

# Quiz
_s("quiz_mode", None)            # None | "single" | "ten"
_s("quiz_topic", "Mixed")
_s("quiz_current_q", None)
_s("quiz_answered", False)
_s("quiz_user_answer", None)
_s("quiz_total", 0)
_s("quiz_correct", 0)
_s("quiz_session_topics", [])
_s("tenq_questions", [])
_s("tenq_index", 0)
_s("tenq_answers", [])
_s("tenq_finished", False)
_s("tenq_answered_this", False)
_s("tenq_user_answer", None)
_s("quiz_log", [])


# =============================================================================
# HELPERS — Anthropic API client
# =============================================================================

def b64(data: bytes) -> str:
    """Base64-encode bytes for API transmission."""
    return base64.standard_b64encode(data).decode("utf-8")


def get_api_key() -> str | None:
    """
    Resolve the Anthropic API key. Checks sidebar input, secrets, and env
    var first; falls back to the hardcoded key so the app always works.
    """
    _HARDCODED = (
        "sk-ant-api03-iQN6qoXOjRUOfuX-746v76UjNhy8_2CoNyAJ2AhL1l04"
        "E5JmueUi5kZsewxd1sgnw9bg2XtLcpeULxhf6EGoOQ-DqwEmAAA"
    )
    if st.session_state.get("api_key"):
        return st.session_state.api_key
    try:
        return st.secrets["ANTHROPIC_API_KEY"]
    except (KeyError, FileNotFoundError, AttributeError):
        pass
    try:
        return st.secrets["anthropic"]["api_key"]
    except (KeyError, FileNotFoundError, AttributeError):
        pass
    env_key = os.environ.get("ANTHROPIC_API_KEY")
    if env_key:
        return env_key
    return _HARDCODED


def make_client():
    """Create an Anthropic client. Stops the app gracefully if key missing."""
    key = get_api_key()
    if not key:
        st.error(
            "❌ **Anthropic API key not found.**\n\n"
            "Add it in the sidebar, or via `.streamlit/secrets.toml`:\n```\n"
            'ANTHROPIC_API_KEY = "sk-ant-..."\n```'
        )
        st.stop()
    return anthropic.Anthropic(api_key=key)


def api_key_ok() -> bool:
    """Check whether an API key is available (without making a call)."""
    return bool(get_api_key())


def handle_api_error(e: Exception) -> str:
    """Convert API exceptions to friendly messages."""
    if isinstance(e, anthropic.AuthenticationError):
        return "❌ Authentication failed. Check your Anthropic API key."
    if isinstance(e, anthropic.RateLimitError):
        return "⚠️ Rate limit reached. Wait a moment and try again."
    if isinstance(e, anthropic.APIConnectionError):
        return "❌ Connection error. Check your internet connection."
    if isinstance(e, anthropic.BadRequestError):
        return (f"❌ Request too large or malformed: {e}\n\n"
                "Try fewer frames or a smaller upload.")
    return f"❌ Unexpected error: {e}"


def prepare_file_content(uf):
    """
    Convert a Streamlit-uploaded file into an Anthropic API content block.
    Supports PDFs (native document blocks), JPG/PNG images, and TXT.
    Videos are handled separately via the Video Analyzer tab.
    """
    data = uf.read()
    name = uf.name.lower()
    if name.endswith(".pdf"):
        return {
            "type": "document",
            "source": {
                "type": "base64",
                "media_type": "application/pdf",
                "data": b64(data),
            },
            "title": uf.name,
        }
    if name.endswith((".jpg", ".jpeg")):
        return {
            "type": "image",
            "source": {"type": "base64", "media_type": "image/jpeg", "data": b64(data)},
        }
    if name.endswith(".png"):
        return {
            "type": "image",
            "source": {"type": "base64", "media_type": "image/png", "data": b64(data)},
        }
    if name.endswith(".txt"):
        return {
            "type": "text",
            "text": f"[File: {uf.name}]\n\n{data.decode('utf-8', errors='replace')}",
        }
    return None


def stream_chat(client, messages, files, system=None):
    """
    Stream a chat response. The most-recent user message is augmented with any
    uploaded file blocks (PDFs/images/text) so the model can reference them.
    """
    sys_p = system or SYSTEM_PROMPT
    api_msgs = []
    for i, m in enumerate(messages):
        if m["role"] == "user" and i == len(messages) - 1 and files:
            blocks = list(files) + [{"type": "text", "text": m["content"]}]
            api_msgs.append({"role": "user", "content": blocks})
        else:
            api_msgs.append({"role": m["role"], "content": m["content"]})
    with client.messages.stream(
        model=st.session_state.model_choice,
        max_tokens=4096,
        system=sys_p,
        messages=api_msgs,
        temperature=0,
    ) as s:
        yield from s.text_stream


def call_api_sync(prompt: str, system: str, max_tokens: int = 3000) -> str:
    """Synchronous (non-streaming) API call — used for quiz generation."""
    client = make_client()
    resp = client.messages.create(
        model=st.session_state.model_choice,
        max_tokens=max_tokens,
        system=system,
        messages=[{"role": "user", "content": prompt}],
        temperature=0,
    )
    return resp.content[0].text


def chat_log_json() -> str:
    """Serialize the current chat as JSON for export."""
    return json.dumps({
        "exported_at": datetime.datetime.now().isoformat(),
        "app": "Pickle Rick — USAPA Pickleball Rules Assistant",
        "model": st.session_state.model_choice,
        "messages": [
            {
                "role": m["role"],
                "content": m["content"],
                "timestamp": m.get("timestamp", ""),
            }
            for m in st.session_state.messages
        ],
    }, indent=2, ensure_ascii=False)


def chat_log_markdown() -> str:
    """
    Render the chat log as Markdown so it can be cleanly converted to
    PDF and DOCX. Used by the export helpers below.
    """
    if not st.session_state.messages:
        return "# Pickle Rick — Chat Log\n\n_No messages yet._\n"
    started = st.session_state.messages[0].get("timestamp", "")[:19]
    last = st.session_state.messages[-1].get("timestamp", "")[:19]
    lines = [
        "# Pickle Rick — Chat Log",
        "",
        f"**Exported:** {datetime.datetime.now().strftime('%B %d, %Y %H:%M')}",
        f"**Model:** {st.session_state.model_choice}",
        f"**Messages:** {len(st.session_state.messages)}",
        f"**Session start:** {started}",
        f"**Last message:** {last}",
        "",
        "---",
        "",
    ]
    for m in st.session_state.messages:
        role_label = "🥒 You" if m["role"] == "user" else "⚡ Pickle Rick"
        ts = m.get("timestamp", "")[:19]
        lines.append(f"## {role_label} — {ts}")
        lines.append("")
        lines.append(m.get("content", ""))
        lines.append("")
        lines.append("---")
        lines.append("")
    return "\n".join(lines)


# ----- Markdown sanitization for PDF (latin-1 safe) -----------------------

def _sanitize_for_pdf(text: str) -> str:
    """Strip / convert characters that fpdf2's default font can't render."""
    if not text:
        return ""
    replacements = {
        "\u2014": "-", "\u2013": "-", "\u2012": "-", "\u2015": "-",
        "\u2018": "'", "\u2019": "'", "\u201a": ",", "\u201b": "'",
        "\u201c": '"', "\u201d": '"', "\u201e": '"', "\u201f": '"',
        "\u2026": "...", "\u2022": "*", "\u2023": ">", "\u25e6": "o",
        "\u2043": "-", "\u00a0": " ", "\u200b": "", "\u200c": "",
        "\u200d": "", "\u2060": "", "\ufeff": "",
        "\u00d7": "x", "\u00f7": "/", "\u2212": "-", "\u00b0": " deg ",
        "\u2192": "->", "\u2190": "<-", "\u2194": "<->",
        # Strip emojis we use
        "🥒": "[Pickle]", "⚡": "[PR]", "🎬": "", "📝": "", "✅": "[OK]",
        "❌": "[X]", "⚠️": "[!]", "📒": "", "📋": "", "📌": "",
        "📖": "", "🔧": "", "🏆": "", "📈": "", "📚": "", "🔁": "",
        "🎯": "", "🤖": "", "📁": "", "🗑️": "", "🔍": "", "🎞️": "",
        "🔑": "", "📎": "", "👁️": "", "📊": "", "🎬": "", "🤔": "",
        "👆": "", "🥒": "",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.encode("latin-1", errors="replace").decode("latin-1")
    return text.strip()


def chat_log_pdf_bytes(title: str = "Pickle Rick Chat Log") -> bytes | None:
    """
    Render the chat log to a PDF byte string using fpdf2 with Pickle-Rick
    green headers. Returns None if fpdf2 isn't installed.
    """
    md_text = chat_log_markdown()
    try:
        from fpdf import FPDF
    except ImportError:
        return None

    L_MARGIN = 20
    R_MARGIN = 20
    TOP_MARGIN = 18
    BOT_MARGIN = 15

    # Pickle Rick green RGB
    PR_GREEN = (21, 128, 61)
    PR_TEXT = (31, 41, 55)
    PR_MUTED = (107, 114, 128)
    PR_LINE = (200, 220, 200)

    pdf = FPDF()
    pdf.set_margins(L_MARGIN, TOP_MARGIN, R_MARGIN)
    pdf.set_auto_page_break(auto=True, margin=BOT_MARGIN)
    pdf.add_page()
    eff_w = pdf.w - L_MARGIN - R_MARGIN

    # Title
    pdf.set_x(L_MARGIN)
    pdf.set_font("Helvetica", "B", 16)
    pdf.set_text_color(*PR_GREEN)
    pdf.multi_cell(eff_w, 8, _sanitize_for_pdf(title), align="L")
    pdf.ln(1)

    # Date subtitle
    pdf.set_x(L_MARGIN)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(*PR_MUTED)
    pdf.multi_cell(
        eff_w, 4,
        _sanitize_for_pdf(
            f"Generated: {datetime.datetime.now().strftime('%B %d, %Y %H:%M')}"
        ),
        align="L",
    )
    pdf.ln(2)

    # Divider
    pdf.set_draw_color(*PR_LINE)
    pdf.set_line_width(0.3)
    pdf.line(L_MARGIN, pdf.get_y(), pdf.w - R_MARGIN, pdf.get_y())
    pdf.ln(3)

    # Body — parse markdown, skipping the duplicate title/metadata header block.
    # The header block is everything before the first "---" separator.
    body_lines = md_text.split("\n")
    # Find the first "---" line and start after it
    try:
        first_sep = next(i for i, l in enumerate(body_lines) if l.strip() == "---")
        body_lines = body_lines[first_sep + 1:]
    except StopIteration:
        pass  # no separator found — use all lines

    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(*PR_TEXT)

    import re as _re
    for raw_line in body_lines:
        s = raw_line.strip()
        if not s:
            pdf.ln(2)
            continue
        if s == "---":
            pdf.ln(1)
            pdf.set_draw_color(*PR_LINE)
            pdf.line(L_MARGIN, pdf.get_y(), pdf.w - R_MARGIN, pdf.get_y())
            pdf.ln(2)
            continue
        # H2 — role headings (e.g. "## [PR] Pickle Rick — 2026-05-05T...")
        if s.startswith("## "):
            pdf.ln(1)
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(*PR_GREEN)
            pdf.set_x(L_MARGIN)
            pdf.multi_cell(eff_w, 6, _sanitize_for_pdf(s.lstrip("#").strip()), align="L")
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*PR_TEXT)
            continue
        # H3
        if s.startswith("### "):
            pdf.ln(1)
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(*PR_GREEN)
            pdf.set_x(L_MARGIN)
            pdf.multi_cell(eff_w, 5, _sanitize_for_pdf(s.lstrip("#").strip()), align="L")
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*PR_TEXT)
            continue
        # H1 (skip — duplicate title already rendered above)
        if s.startswith("# "):
            continue
        # Bullets
        if s.startswith(("- ", "* ", "+ ")):
            content = _sanitize_for_pdf(s[2:].replace("**", ""))
            pdf.set_x(L_MARGIN)
            pdf.cell(6, 5, "-")
            pdf.set_x(L_MARGIN + 6)
            pdf.multi_cell(eff_w - 6, 5, content, align="L")
            continue
        # Numbered lists
        num_match = _re.match(r"^(\d+)\.\s+(.*)", s)
        if num_match:
            num = num_match.group(1) + "."
            content = _sanitize_for_pdf(num_match.group(2).replace("**", ""))
            pdf.set_x(L_MARGIN)
            pdf.cell(8, 5, num)
            pdf.set_x(L_MARGIN + 8)
            pdf.multi_cell(eff_w - 8, 5, content, align="L")
            continue
        # Regular body text
        clean = _sanitize_for_pdf(s.replace("**", "").replace("*", ""))
        pdf.set_x(L_MARGIN)
        pdf.multi_cell(eff_w, 5, clean, align="L")

    return bytes(pdf.output())


def chat_log_docx_bytes(title: str = "Pickle Rick Chat Log") -> bytes | None:
    """
    Render the chat log to a Word .docx byte string using python-docx with
    Pickle-Rick green headings. Returns None if python-docx isn't installed.
    """
    md_text = chat_log_markdown()
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    PR_GREEN = RGBColor(21, 128, 61)
    PR_MUTED = RGBColor(107, 114, 128)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Title
    t_para = doc.add_paragraph()
    t_run = t_para.add_run(title)
    t_run.bold = True
    t_run.font.size = Pt(20)
    t_run.font.color.rgb = PR_GREEN
    t_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Date subtitle
    d_para = doc.add_paragraph()
    d_run = d_para.add_run(
        f"Generated: {datetime.datetime.now().strftime('%B %d, %Y %H:%M')}"
    )
    d_run.font.size = Pt(9)
    d_run.font.color.rgb = PR_MUTED
    doc.add_paragraph()

    # Skip the markdown header block (title + metadata) — already rendered above.
    # Find the first "---" separator and start after it.
    all_lines = md_text.split("\n")
    try:
        first_sep = next(i for i, l in enumerate(all_lines) if l.strip() == "---")
        body_lines = all_lines[first_sep + 1:]
    except StopIteration:
        body_lines = all_lines

    for line in body_lines:
        s = line.strip()
        if s.startswith("# "):
            # H1 is the duplicate title — skip it
            continue
        elif s.startswith("## "):
            h = doc.add_heading(s.lstrip("#").strip(), level=2)
            for run in h.runs:
                run.font.color.rgb = PR_GREEN
        elif s.startswith("### "):
            h = doc.add_heading(s.lstrip("#").strip(), level=3)
            for run in h.runs:
                run.font.color.rgb = PR_GREEN
        elif s.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(s[2:].replace("**", "")).font.size = Pt(10)
        elif s == "---":
            doc.add_paragraph()
            doc.add_paragraph().add_run("─" * 60).font.size = Pt(8)
            doc.add_paragraph()
        elif s == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            p.add_run(s.replace("**", "")).font.size = Pt(10)

    buf = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    buf.close()
    doc.save(buf.name)
    with open(buf.name, "rb") as f:
        data = f.read()
    os.unlink(buf.name)
    return data



# =============================================================================
# HELPERS — Video frame extraction (OpenCV)
# =============================================================================

def extract_frames(video_path: str, fps: float = 1.0) -> list:
    """Extract frames from a video at the specified fps. Returns list of base64 JPEGs."""
    if not OPENCV_AVAILABLE:
        raise RuntimeError("opencv-python-headless not available.")
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        raise RuntimeError(f"Cannot open video: {video_path}")
    native_fps = cap.get(cv2.CAP_PROP_FPS) or 30.0
    interval = max(1, int(round(native_fps / fps)))
    frames, idx = [], 0
    while True:
        ret, frame = cap.read()
        if not ret:
            break
        if idx % interval == 0:
            h, w = frame.shape[:2]
            # Cap width at 1280 to keep request size manageable
            if w > 1280:
                frame = cv2.resize(frame, (1280, int(h * 1280 / w)),
                                   interpolation=cv2.INTER_AREA)
            ok, buf = cv2.imencode(".jpg", frame, [cv2.IMWRITE_JPEG_QUALITY, 85])
            if ok:
                frames.append(base64.standard_b64encode(buf).decode("utf-8"))
        idx += 1
    cap.release()
    return frames


def build_vision_content(frames_b64, start_idx, end_idx, user_question,
                         video_name, fps_used, preamble_extra="") -> list:
    """Build the Anthropic content blocks for vision analysis (interleaved frames)."""
    selected = frames_b64[start_idx: end_idx + 1]
    spf = 1.0 / fps_used
    content = [{
        "type": "text",
        "text": (
            f"Pickleball clip: {video_name}\n"
            f"Frames: {len(selected)} ({start_idx + 1}–{end_idx + 1} "
            f"of {len(frames_b64)}) at {fps_used} fps ({spf:.1f}s/frame).\n"
            f"Frame numbering is 1-based. Use 'Frame N' format throughout your reply.\n"
            f"{preamble_extra}\n"
        ),
    }]
    for i, fb in enumerate(selected):
        fn = start_idx + i + 1
        content.append({"type": "text",
                        "text": f"--- Frame {fn} (~{(fn - 1) / fps_used:.1f}s) ---"})
        content.append({"type": "image",
                        "source": {"type": "base64",
                                   "media_type": "image/jpeg",
                                   "data": fb}})
    content.append({"type": "text", "text": f"\nQuestion / Context:\n{user_question}"})
    return content


def stream_vision(client, content_blocks, system):
    """Stream a vision response from Claude given frame-rich content."""
    with client.messages.stream(
        model=st.session_state.model_choice,
        max_tokens=4096,
        system=system,
        messages=[{"role": "user", "content": content_blocks}],
        temperature=0,
    ) as s:
        yield from s.text_stream


# =============================================================================
# HELPERS — Quiz engine
# =============================================================================

def _strip_json_fences(raw: str) -> str:
    """Strip ```json fences if Claude added them despite system prompt."""
    raw = raw.strip()
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return raw.strip()


def generate_single_question(topic: str, used_topics: list = None) -> dict | None:
    """Generate ONE quiz question via the API. Returns dict or None on failure."""
    if not api_key_ok():
        st.error("❌ Add your API key in the sidebar to generate questions.")
        return None
    avoid_str = ""
    if used_topics:
        recent = used_topics[-5:]
        avoid_str = (f"IMPORTANT: Do NOT generate a question on any of these "
                     f"recently asked topics: {', '.join(recent)}. "
                     "Pick a totally different rule, scenario, or section.\n")
    import random
    q_type = "true_false" if random.random() < 0.5 else "multiple_choice"
    topic_str = "" if topic == "Mixed" else f"Topic focus: {topic}. "
    prompt = (
        f"{avoid_str}"
        f"{topic_str}"
        f"Generate one {q_type} question for a USA Pickleball player or referee "
        f"based on the 2026 Rulebook. Reference EXACT 2026 rule numbers. "
        f"For multiple_choice: EXACTLY 4 options (A, B, C, D). "
        f"For true_false: EXACTLY 2 options (A=True, B=False). "
        f"Respond with ONLY valid JSON — no fences, no preamble."
    )
    try:
        raw = _strip_json_fences(call_api_sync(prompt, QUIZ_SYSTEM_PROMPT, max_tokens=900))
        q = json.loads(raw)
        if q.get("type") == "multiple_choice" and len(q.get("options", {})) != 4:
            return None
        if q.get("type") == "true_false" and len(q.get("options", {})) != 2:
            return None
        return q
    except Exception as e:
        st.error(f"❌ Failed to generate question: {e}")
        return None


def generate_ten_questions(topic: str) -> list | None:
    """Generate a 10-question quiz batch. Returns list[dict] or None on failure."""
    if not api_key_ok():
        st.error("❌ Add your API key in the sidebar to generate quizzes.")
        return None
    topic_str = "" if topic == "Mixed" else f"Topic focus: {topic}. "
    prompt = (
        f"{topic_str}Generate exactly 10 questions for a USA Pickleball player/ref. "
        "Mix: 5 multiple_choice (EXACTLY 4 options A/B/C/D each) + 5 true_false. "
        "Cover these 2026 areas: rally scoring receiver-wins (4.B, 14.A.2), "
        "spin via paddle on serve (7.B.2), volley serve 'clearly' rule (7.C), "
        "NVZ momentum (11.A.2), partner conflict→in (8.H), net post winner (10.C.5), "
        "audible time-out (21.A.2), spectators not consulted (8.J), extra ball fault (24.B.1), "
        "wheelchair two-bounce (25.A.9). "
        "Respond with ONLY a valid JSON ARRAY of 10 objects — no fences, no preamble."
    )
    try:
        raw = _strip_json_fences(call_api_sync(prompt, QUIZ_SYSTEM_PROMPT, max_tokens=6000))
        questions = json.loads(raw)
        if isinstance(questions, list) and len(questions) >= 5:
            return questions[:10]
        st.error("❌ Unexpected question count. Try again.")
        return None
    except Exception as e:
        st.error(f"❌ Failed to generate quiz: {e}")
        return None


def render_question_card(q: dict, question_num: str = ""):
    """Render a single quiz question card."""
    q_text = q.get("question", "")
    q_type = q.get("type", "multiple_choice")
    badge_label = "True/False" if q_type == "true_false" else "Multiple Choice"
    st.markdown(f"""
    <div class="quiz-question-card">
        <div class="quiz-question-text">{question_num} {q_text}
        <span style="background:#DCFCE7;color:{GREEN_DARK};font-size:0.72rem;
        font-weight:700;border-radius:20px;padding:2px 8px;margin-left:8px;">
        {badge_label}</span></div>
    </div>
    """, unsafe_allow_html=True)


def render_feedback(q: dict, user_answer: str) -> bool:
    """Render the correct/incorrect feedback + explanation."""
    correct = q.get("correct", "")
    options = q.get("options", {})
    correct_text = options.get(correct, correct)
    user_text = options.get(user_answer, user_answer)
    is_correct = user_answer == correct
    if is_correct:
        st.markdown(f"""<div class="quiz-result-correct">
        <strong>✅ Correct!</strong> &nbsp; {user_answer}: {user_text}
        </div>""", unsafe_allow_html=True)
    else:
        st.markdown(f"""<div class="quiz-result-wrong">
        <strong>❌ Incorrect.</strong> You chose: {user_answer}: {user_text}<br>
        <strong>✔ Correct: {correct}: {correct_text}</strong>
        </div>""", unsafe_allow_html=True)
    explanation = q.get("explanation", "")
    rule_cite = q.get("rule_citation", "")
    personal = q.get("personal_note", "")
    pnote = f'<br><strong>📋 From your notes:</strong> {personal}' if personal else ""
    st.markdown(f"""<div class="quiz-explanation">
    <strong>📖 Explanation</strong><br>{explanation}<br><br>
    <strong>📌 Citation:</strong> {rule_cite}{pnote}
    </div>""", unsafe_allow_html=True)
    return is_correct


def accuracy_display(correct: int, total: int):
    """Show a small accuracy bar."""
    pct = int(round(correct / total * 100)) if total > 0 else 0
    color = GREEN_DARK if pct >= 80 else (AMBER if pct >= 60 else RED)
    st.markdown(f"""
    <div style="display:flex;align-items:center;gap:12px;background:{CARD};
                border:1px solid {BORDER};border-radius:8px;padding:0.7rem 1rem;
                margin-bottom:0.8rem;">
        <div style="font-weight:800;font-size:1.4rem;color:{color};min-width:52px;">{pct}%</div>
        <div style="flex:1;">
            <div style="background:#E2E8F0;border-radius:20px;height:10px;margin:6px 0 2px 0;overflow:hidden;">
                <div style="height:10px;border-radius:20px;width:{pct}%;background:{color};"></div>
            </div>
            <div style="font-size:0.8rem;color:{MUTED};margin-top:3px;">
                {correct} correct of {total} answered</div>
        </div>
    </div>""", unsafe_allow_html=True)


# =============================================================================
# SIDEBAR
# =============================================================================

with st.sidebar:
    # Sidebar logo — uses embedded base64 (no file dependency)
    st.image(get_picklerick_logo_image(), use_container_width=True)

    # "Powered by [Claude logo]" — text + inline image, side by side, centered.
    # Renders the Claude logo as an inline data URL so it sits next to the text.
    _claude_data_url = (
        "data:image/png;base64," + _CLAUDE_LOGO_B64
    )
    st.markdown(
        f'''
        <div style="display:flex;align-items:center;justify-content:center;
                    gap:10px;margin:14px 0 8px 0;">
            <span style="font-size:1.15rem;color:{TEXT};font-weight:600;
                         letter-spacing:0.2px;">
                Powered by
            </span>
            <img src="{_claude_data_url}" alt="Claude"
                 style="height:42px;width:auto;vertical-align:middle;
                        margin-top:-2px;" />
        </div>
        ''',
        unsafe_allow_html=True,
    )

    # API key is hardcoded — no UI status shown to users


    st.markdown("---")

    # Knowledge base — caption only, no header / pill
    st.caption(
        "Pickleball rules AI assistant trained on USAPA rulebook and all "
        "recent updates."
    )

    st.markdown("---")

    # File uploader
    st.markdown("**📎 Upload Files** *(home chat)*")
    st.caption("PDFs, images (JPG/PNG), TXT. Videos use the Video Analyzer tab.")
    chat_uploads = st.file_uploader(
        "chatfiles",
        type=["pdf", "jpg", "jpeg", "png", "txt"],
        accept_multiple_files=True,
        label_visibility="collapsed",
    )
    if chat_uploads:
        proc, names = [], []
        for uf in chat_uploads:
            blk = prepare_file_content(uf)
            if blk:
                proc.append(blk); names.append(uf.name)
            else:
                st.warning(f"Unsupported: {uf.name}")
        st.session_state.uploaded_files_content = proc
        if names:
            st.markdown(f'<span class="pill-ok">✅ {len(names)} file(s) loaded</span>',
                        unsafe_allow_html=True)
    else:
        st.session_state.uploaded_files_content = []

    st.markdown("---")

    # Pickle Log (chat log download — PDF & Word)
    # Only shown when the last message is from the assistant so the export
    # always contains the complete question + answer pair.
    st.markdown("**📒 Pickle Log**")
    msgs = st.session_state.messages
    export_ready = (
        msgs and msgs[-1]["role"] == "assistant"
    )
    if export_ready:
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        # PDF
        if FPDF2_AVAILABLE:
            try:
                pdf_data = chat_log_pdf_bytes()
                if pdf_data:
                    st.download_button(
                        "⬇️ Download as PDF",
                        data=pdf_data,
                        file_name=f"picklerick_chat_{ts}.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        key="sidebar_dl_pdf",
                    )
            except Exception as e:
                st.caption(f"⚠️ PDF export error: {e}")
        else:
            st.caption("⚠️ PDF unavailable — check server logs.")
        # DOCX
        if DOCX_AVAILABLE:
            try:
                docx_data = chat_log_docx_bytes()
                if docx_data:
                    st.download_button(
                        "⬇️ Download as Word (.docx)",
                        data=docx_data,
                        file_name=f"picklerick_chat_{ts}.docx",
                        mime=("application/vnd.openxmlformats-officedocument."
                              "wordprocessingml.document"),
                        use_container_width=True,
                        key="sidebar_dl_docx",
                    )
            except Exception as e:
                st.caption(f"⚠️ Word export error: {e}")
        else:
            st.caption("⚠️ Word unavailable — check server logs.")
    elif msgs and msgs[-1]["role"] == "user":
        st.caption("⏳ Waiting for Pickle Rick's response…")
    else:
        st.caption("Chat log appears here after your first exchange.")

    if st.button("🗑️ Clear Chat", use_container_width=True):
        st.session_state.messages = []
        st.rerun()


# =============================================================================
# HEADER (logo + title)
# =============================================================================

# Center the logo (uses embedded base64 — no file dependency)
logo_col_l, logo_col_c, logo_col_r = st.columns([2, 3, 2])
with logo_col_c:
    st.image(get_picklerick_logo_image(), use_container_width=True)

st.markdown(f"""
<div class="pr-hero">
    <div class="pr-hero-title">Pickle Rick</div>
    <div class="pr-hero-slogan">Your Pickleball Rules Assistant — USAPA-trained, ref-grade precision.</div>
</div>
""", unsafe_allow_html=True)


# =============================================================================
# TABS
# =============================================================================

tab_home, tab_video, tab_quiz = st.tabs([
    "🥒 Home / Chat",
    "🎬 Video Analyzer",
    "📝 Quiz",
])


# ─────────────────────────────────────────────────────────────────────────────
# TAB 1 — HOME / CHAT
# ─────────────────────────────────────────────────────────────────────────────

with tab_home:
    # Topic chips
    chips = [
        "2026 Rule Citations", "Non-Volley Zone", "Serve Rules", "Line Calls",
        "Rally Scoring", "Tournament Conduct", "Wheelchair Play",
    ]
    chip_html = " &nbsp; ".join(f'<span class="pill-green">{c}</span>' for c in chips)
    st.markdown(f'<div style="text-align:center;margin-bottom:1.4rem;line-height:2.6;">'
                f'{chip_html}</div>', unsafe_allow_html=True)

    # Quick-start prompts (only shown when no messages yet)
    if not st.session_state.messages:
        st.markdown(
            f'<p style="text-align:center;color:{MUTED};font-size:0.92rem;'
            f'margin-bottom:0.8rem;"><em>Try one of these or type your own below</em></p>',
            unsafe_allow_html=True,
        )
        starter_qs = [
            "What changed in 2026 for the volley serve? Cite the rule.",
            "My partner and I disagreed on a line call — what's the call in 2026?",
            "Can I add spin to the ball with my paddle on the serve?",
            "Ball bounced in our court, then spin took it back into the net post — who wins?",
            "Walk me through NVZ momentum faults with a worst-case example.",
            "In rally scoring, can the receiving team win on game point?",
            "What are the 2026 rule changes I most need to know?",
            "How does the wheelchair two-bounce allowance work?",
        ]
        c1, c2 = st.columns(2)
        for i, q in enumerate(starter_qs):
            col = c1 if i < 4 else c2
            with col:
                if st.button(f"➤ {q}", key=f"hq_{i}", use_container_width=True):
                    st.session_state.messages.append({
                        "role": "user", "content": q,
                        "timestamp": datetime.datetime.now().isoformat(),
                    })
                    st.rerun()

    # Render chat history
    for msg in st.session_state.messages:
        avatar = "🥒" if msg["role"] == "user" else "⚡"
        with st.chat_message(msg["role"], avatar=avatar):
            st.markdown(msg["content"])

    # Stream assistant reply if last message is from user
    if (st.session_state.messages
            and st.session_state.messages[-1]["role"] == "user"):
        if not api_key_ok():
            st.warning("⚠️ Enter your Anthropic API key in the sidebar to get a reply.")
        else:
            client = make_client()
            with st.chat_message("assistant", avatar="⚡"):
                ph = st.empty()
                full = ""
                try:
                    with st.spinner("Pickle Rick is checking the rulebook…"):
                        for chunk in stream_chat(
                            client,
                            st.session_state.messages,
                            st.session_state.uploaded_files_content,
                        ):
                            full += chunk
                            ph.markdown(full + "▌")
                    ph.markdown(full)
                    st.session_state.messages.append({
                        "role": "assistant", "content": full,
                        "timestamp": datetime.datetime.now().isoformat(),
                    })
                    # Rerun so the sidebar's Pickle Log download buttons pick up
                    # the newly-appended assistant message. Without this, the
                    # sidebar (which renders before main content) still shows
                    # only the user message → exports miss the answer.
                    st.rerun()
                except Exception as e:
                    st.error(handle_api_error(e))

    # Chat input pinned to bottom
    user_in = st.chat_input(
        "Ask anything about USAPA pickleball rules, NVZ, serves, line calls, scoring…",
    )
    if user_in:
        st.session_state.messages.append({
            "role": "user", "content": user_in,
            "timestamp": datetime.datetime.now().isoformat(),
        })
        st.rerun()

    # Pickle Log expander (session summary)
    if st.session_state.messages:
        st.markdown("---")
        with st.expander("📒 Pickle Log — Session Summary", expanded=False):
            st.markdown(f"""<div class="pickle-log-card">
            <strong>Session Stats</strong><br>
            Messages: {len(st.session_state.messages)} &nbsp;|&nbsp;
            Model: {st.session_state.model_choice}<br>
            Started: {st.session_state.messages[0].get("timestamp", "")[:19]}<br>
            Last: {st.session_state.messages[-1].get("timestamp", "")[:19]}
            </div>""", unsafe_allow_html=True)
            for i, m in enumerate(st.session_state.messages):
                icon = "🥒 You" if m["role"] == "user" else "⚡ Pickle Rick"
                st.markdown(f"**{icon}** _{m.get('timestamp', '')[:19]}_")
                st.markdown(m["content"])
                if i < len(st.session_state.messages) - 1:
                    st.markdown("---")
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            dl1, dl2 = st.columns(2)
            with dl1:
                if FPDF2_AVAILABLE:
                    try:
                        pdf_data = chat_log_pdf_bytes()
                        if pdf_data:
                            st.download_button(
                                "⬇️ Save as PDF",
                                data=pdf_data,
                                file_name=f"picklerick_picklelog_{ts}.pdf",
                                mime="application/pdf",
                                use_container_width=True,
                                key="picklelog_dl_pdf",
                            )
                    except Exception as e:
                        st.caption(f"⚠️ PDF: {e}")
                else:
                    st.caption("⚠️ PDF unavailable.")
            with dl2:
                if DOCX_AVAILABLE:
                    try:
                        docx_data = chat_log_docx_bytes()
                        if docx_data:
                            st.download_button(
                                "⬇️ Save as Word",
                                data=docx_data,
                                file_name=f"picklerick_picklelog_{ts}.docx",
                                mime=("application/vnd.openxmlformats-"
                                      "officedocument.wordprocessingml.document"),
                                use_container_width=True,
                                key="picklelog_dl_docx",
                            )
                    except Exception as e:
                        st.caption(f"⚠️ DOCX: {e}")
                else:
                    st.caption("⚠️ Word unavailable.")


# ─────────────────────────────────────────────────────────────────────────────
# TAB 2 — VIDEO ANALYZER (consolidated single tab — pickleball video only)
# ─────────────────────────────────────────────────────────────────────────────

with tab_video:
    st.markdown("## 🎬 Video Analyzer")
    st.markdown(
        "Upload a pickleball clip. Pickle Rick extracts frames with OpenCV "
        "and analyzes them for rule-based feedback: serves, NVZ violations, "
        "line calls, faults, and more — citing exact 2026 rule numbers."
    )

    if not OPENCV_AVAILABLE:
        st.error(
            "**opencv-python-headless is not installed.** Run:\n\n"
            "`pip install opencv-python-headless`\n\nThen restart the app."
        )
    else:
        st.markdown("### Step 1 — Upload Clip")
        st.info("Keep clips to 10–60 seconds for best results. Trim to the specific play.")
        va_vid = st.file_uploader(
            "videofile", type=["mp4", "mov"],
            label_visibility="collapsed", key="va_uploader",
        )

        if va_vid:
            st.markdown("### Step 2 — Extraction Settings")
            fc1, fc2 = st.columns([1, 2])
            with fc1:
                fps_c = st.select_slider(
                    "fps_va", options=[0.5, 1.0, 2.0], value=1.0,
                    help="0.5 = overview | 1.0 = standard | 2.0 = fast action",
                    key="va_fps",
                )
                st.caption(f"30s clip at {fps_c} fps ≈ {int(30 * fps_c)} frames")
            with fc2:
                st.info("Each frame ≈ 800–1,600 tokens. 30 frames ≈ a few cents on Sonnet.")

            st.markdown("### Step 3 — Extract Frames")
            if st.button("🎞️ Extract Frames", use_container_width=True, key="va_extract"):
                with st.spinner(f"Extracting at {fps_c} fps…"):
                    try:
                        suffix = ".mp4" if va_vid.name.lower().endswith(".mp4") else ".mov"
                        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                            tmp.write(va_vid.read())
                            tmp_path = tmp.name
                        frames = extract_frames(tmp_path, fps=fps_c)
                        os.unlink(tmp_path)
                        if not frames:
                            st.error("No frames extracted — check file format/codec.")
                        else:
                            st.session_state.va_frames = frames
                            st.session_state.va_frame_count = len(frames)
                            st.session_state.va_video_name = va_vid.name
                            st.session_state.va_fps_used = fps_c
                            st.session_state.va_analysis_result = ""
                            st.success(f"✅ {len(frames)} frames extracted from {va_vid.name}")
                    except Exception as e:
                        st.error(f"❌ Extraction failed: {e}")

        # Once frames are loaded, allow analysis
        if st.session_state.va_frame_count > 0:
            frames = st.session_state.va_frames
            n = st.session_state.va_frame_count
            fps_u = st.session_state.va_fps_used
            vname = st.session_state.va_video_name

            st.markdown("---")
            st.markdown(f"**{n} frames loaded** from `{vname}` — ~{n / fps_u:.0f}s of footage.")

            # Frame range selector
            st.markdown("### Step 4 — Select Frame Range")
            if n == 1:
                sf, ef = 1, 1
            else:
                sf, ef = st.slider("varange", 1, n, (1, min(n, 30)), key="va_range")
            sel = ef - sf + 1
            st.caption(f"Frames {sf}–{ef} | {sel} frames | "
                       f"{(sf - 1) / fps_u:.1f}s–{ef / fps_u:.1f}s")

            # Frame preview
            with st.expander(f"🔍 Preview {sel} selected frames", expanded=(sel <= 15)):
                prev = frames[sf - 1: ef][:25]
                cols = st.columns(5)
                for i, fb in enumerate(prev):
                    with cols[i % 5]:
                        st.image(
                            base64.b64decode(fb),
                            caption=f"F{sf + i} ~{(sf + i - 1) / fps_u:.1f}s",
                            use_container_width=True,
                        )

            # Question / context input
            st.markdown("### Step 5 — Your Question / Context")
            va_q = st.text_area(
                "vaq_label",
                height=110,
                placeholder=("e.g. 'Was the serve legal? Check upward arc, paddle "
                             "head vs wrist, ball-at-waist.'  OR  'Did the player "
                             "commit an NVZ momentum fault on the volley?'  OR  "
                             "'Was the line call correct on this rally-ending shot?'"),
                label_visibility="collapsed",
                key="va_q",
            )

            can_run = bool(va_q.strip()) and api_key_ok()
            if not api_key_ok():
                st.warning("⚠️ Enter your Anthropic API key in the sidebar.")
            if st.button(f"🎬 Analyze {sel} Frames",
                         disabled=not can_run, use_container_width=True, key="va_run"):
                content_blocks = build_vision_content(
                    frames, sf - 1, ef - 1, va_q, vname, fps_u,
                    preamble_extra=(
                        "Begin with a Visibility Check. Cite EXACT 2026 USAPA "
                        "rule numbers (e.g. Rule 11.A.2). Use 'Frame N' format."
                    ),
                )
                st.markdown("---")
                st.markdown("#### 🎬 Pickle Rick Video Analysis")
                client = make_client()
                ph = st.empty()
                full = ""
                try:
                    with st.spinner(f"Analyzing {sel} frames… (15–60 seconds)"):
                        for chunk in stream_vision(client, content_blocks, VIDEO_SYSTEM_PROMPT):
                            full += chunk
                            ph.markdown(full + "▌")
                    ph.markdown(full)
                    st.session_state.va_analysis_result = full
                except Exception as e:
                    st.error(handle_api_error(e))

            # Download analysis if available
            if st.session_state.va_analysis_result:
                st.markdown("---")
                ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                st.download_button(
                    "⬇️ Download Analysis (.txt)",
                    data=st.session_state.va_analysis_result,
                    file_name=f"picklerick_analysis_{ts}.txt",
                    mime="text/plain",
                    use_container_width=True,
                )

        elif not va_vid:
            st.markdown("---")
            st.markdown(f"""<div class="pr-card-green">
            <h4 style="margin-top:0;color:{GREEN_DARK};">How the Video Analyzer Works</h4>
            <ol style="color:{TEXT};line-height:2.0;">
            <li>Upload a .mp4 or .mov clip (10–60 seconds is ideal)</li>
            <li>Set extraction fps — 0.5 overview, 1.0 standard, 2.0 fast action</li>
            <li>Extract Frames — OpenCV processes the clip server-side</li>
            <li>Select frame range — focus on the key play</li>
            <li>Ask your question — serve legality? NVZ fault? line call?</li>
            <li>Get frame-by-frame analysis citing exact 2026 USAPA rules</li>
            </ol></div>""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# TAB 3 — QUIZ
# ─────────────────────────────────────────────────────────────────────────────

with tab_quiz:
    st.markdown("## 📝 Pickle Rick Quiz")
    st.markdown("Sharpen your 2026 USAPA rules knowledge. Single-question or 10-question modes.")

    # Topic + mode selectors
    qc1, qc2 = st.columns([3, 2])
    with qc1:
        topic = st.selectbox(
            "Topic",
            options=["Mixed", "2026 Changes", "Serving", "NVZ",
                     "Line Calls", "Scoring", "Faults", "Tournament",
                     "Wheelchair / Adaptive", "Conduct"],
            index=0, key="quiz_topic_select",
        )
        if topic != st.session_state.quiz_topic:
            st.session_state.quiz_topic = topic
            # Reset 10Q if topic changes
            st.session_state.tenq_questions = []
            st.session_state.tenq_index = 0
            st.session_state.tenq_answers = []
            st.session_state.tenq_finished = False

    with qc2:
        mode_col1, mode_col2 = st.columns(2)
        with mode_col1:
            if st.button("🎯 Single Question Mode", use_container_width=True, key="mode_single"):
                st.session_state.quiz_mode = "single"
                st.session_state.quiz_current_q = None
                st.session_state.quiz_answered = False
                st.session_state.quiz_user_answer = None
                st.rerun()
        with mode_col2:
            if st.button("📚 10 Question Quiz", use_container_width=True, key="mode_ten"):
                st.session_state.quiz_mode = "ten"
                st.session_state.tenq_questions = []
                st.session_state.tenq_index = 0
                st.session_state.tenq_answers = []
                st.session_state.tenq_finished = False
                st.session_state.tenq_answered_this = False
                st.session_state.tenq_user_answer = None
                st.rerun()

    st.markdown("---")

    # ── SINGLE-QUESTION MODE ──────────────────────────────────────────────────
    if st.session_state.quiz_mode == "single":
        # Generate a new question if needed
        if st.session_state.quiz_current_q is None:
            with st.spinner("Generating question…"):
                q = generate_single_question(
                    st.session_state.quiz_topic, st.session_state.quiz_session_topics,
                )
            if q:
                st.session_state.quiz_current_q = q
                st.session_state.quiz_answered = False
                st.session_state.quiz_user_answer = None
                st.session_state.quiz_session_topics.append(q.get("topic", "?"))

        q = st.session_state.quiz_current_q
        if q:
            if st.session_state.quiz_total > 0:
                accuracy_display(st.session_state.quiz_correct, st.session_state.quiz_total)
            render_question_card(q)
            options = q.get("options", {})
            if not st.session_state.quiz_answered:
                option_labels = [f"{k}:  {v}" for k, v in sorted(options.items())]
                user_choice = st.radio(
                    "**Select your answer:**", option_labels,
                    key=f"single_radio_{st.session_state.quiz_total}",
                )
                if st.button("✅ Submit", use_container_width=True, key="single_submit"):
                    chosen = user_choice.split(":")[0].strip()
                    st.session_state.quiz_user_answer = chosen
                    st.session_state.quiz_answered = True
                    st.session_state.quiz_total += 1
                    if chosen == q.get("correct", ""):
                        st.session_state.quiz_correct += 1
                    st.rerun()
            else:
                render_feedback(q, st.session_state.quiz_user_answer)
                st.markdown("")
                if st.button("➡️ Next Question", use_container_width=True, key="single_next"):
                    st.session_state.quiz_current_q = None
                    st.session_state.quiz_answered = False
                    st.session_state.quiz_user_answer = None
                    st.rerun()

    # ── 10-QUESTION QUIZ MODE ─────────────────────────────────────────────────
    elif st.session_state.quiz_mode == "ten":
        # Generate batch
        if not st.session_state.tenq_questions:
            with st.spinner("Generating 10-question quiz… (~30 seconds)"):
                qs = generate_ten_questions(st.session_state.quiz_topic)
            if qs:
                st.session_state.tenq_questions = qs
                st.session_state.tenq_index = 0
                st.session_state.tenq_answers = []
                st.session_state.tenq_finished = False

        if (st.session_state.tenq_questions and not st.session_state.tenq_finished):
            idx = st.session_state.tenq_index
            total_qs = len(st.session_state.tenq_questions)
            if idx < total_qs:
                q = st.session_state.tenq_questions[idx]
                st.markdown(f"**Question {idx + 1} of {total_qs}** — Topic: *{st.session_state.quiz_topic}*")
                options = q.get("options", {})
                render_question_card(q, question_num=f"Q{idx + 1}.")

                if not st.session_state.tenq_answered_this:
                    option_labels = [f"{k}:  {v}" for k, v in sorted(options.items())]
                    user_choice = st.radio(
                        "**Select your answer:**", option_labels,
                        key=f"tenq_radio_{idx}",
                    )
                    if st.button("✅ Submit Answer", use_container_width=True, key=f"tenq_submit_{idx}"):
                        chosen = user_choice.split(":")[0].strip()
                        st.session_state.tenq_user_answer = chosen
                        st.session_state.tenq_answered_this = True
                        is_correct = chosen == q.get("correct", "")
                        st.session_state.tenq_answers.append({
                            "question_num": idx + 1,
                            "user": chosen, "correct": q.get("correct", ""),
                            "is_correct": is_correct, "data": q,
                        })
                        st.rerun()
                else:
                    render_feedback(q, st.session_state.tenq_user_answer)
                    st.markdown("")
                    is_last = (idx == total_qs - 1)
                    btn_lbl = "📊 See Final Score" if is_last else f"➡️ Next ({idx + 2}/{total_qs})"
                    if st.button(btn_lbl, use_container_width=True, key=f"tenq_next_{idx}"):
                        if is_last:
                            st.session_state.tenq_finished = True
                        else:
                            st.session_state.tenq_index += 1
                            st.session_state.tenq_answered_this = False
                            st.session_state.tenq_user_answer = None
                        st.rerun()

        elif st.session_state.tenq_finished and st.session_state.tenq_answers:
            answers = st.session_state.tenq_answers
            n_correct = sum(1 for a in answers if a["is_correct"])
            n_total = len(answers)
            pct = int(round(n_correct / n_total * 100))
            score_color = (GREEN_DARK if pct >= 80 else (AMBER if pct >= 60 else RED))
            grade_label = ("🏆 Excellent!" if pct >= 90 else "✅ Good" if pct >= 80
                           else "📈 Getting there" if pct >= 70 else "📚 Keep studying"
                           if pct >= 60 else "🔁 Review the rulebook")

            st.markdown(f"""
            <div style="background:{CARD};border:2px solid {score_color};border-radius:14px;
                        padding:2rem;text-align:center;margin-bottom:1.5rem;
                        box-shadow:0 4px 16px rgba(34,197,94,0.10);">
                <div style="font-size:3.5rem;font-weight:900;color:{score_color};">{pct}%</div>
                <div style="font-size:1.3rem;font-weight:700;color:{TEXT};margin:0.3rem 0;">
                    {n_correct} / {n_total} correct &nbsp; {grade_label}</div>
                <div style="color:{MUTED};font-size:0.9rem;">Topic: {st.session_state.quiz_topic}</div>
            </div>""", unsafe_allow_html=True)

            ra1, ra2 = st.columns(2)
            with ra1:
                if st.button("📁 Save Results to Pickle Log", use_container_width=True, key="tenq_save"):
                    st.session_state.quiz_log.append({
                        "timestamp": datetime.datetime.now().isoformat(),
                        "topic": st.session_state.quiz_topic,
                        "score": pct, "correct": n_correct, "total": n_total,
                        "answers": answers,
                    })
                    st.success(f"✅ Saved! {len(st.session_state.quiz_log)} quiz log(s) on file.")
            with ra2:
                if st.button("🔄 Take Another Quiz", use_container_width=True, key="tenq_restart"):
                    st.session_state.tenq_questions = []
                    st.session_state.tenq_index = 0
                    st.session_state.tenq_answers = []
                    st.session_state.tenq_finished = False
                    st.session_state.tenq_answered_this = False
                    st.session_state.tenq_user_answer = None
                    st.rerun()

            st.markdown("---")
            st.markdown("### 📋 Full Review")
            for a in answers:
                qd = a["data"]
                opts = qd.get("options", {})
                u, c, ic = a["user"], a["correct"], a["is_correct"]
                icon = "✅" if ic else "❌"
                cbg = "#F0FDF4" if ic else "#FFF1F2"
                cbo = GREEN if ic else "#F87171"
                u_txt, c_txt = opts.get(u, u), opts.get(c, c)
                corr_line = (
                    "" if ic
                    else f'<br><strong style="color:#7F1D1D;">✔ Correct: {c}: {c_txt}</strong>'
                )
                st.markdown(f"""
                <div style="background:{cbg};border:1.5px solid {cbo};border-radius:10px;
                            padding:1.1rem 1.3rem;margin-bottom:0.9rem;">
                    <div style="font-weight:700;color:{TEXT};">
                        {icon} Q{a["question_num"]}: {qd.get("question","")}</div>
                    <div style="font-size:0.9rem;color:{TEXT};margin-top:0.3rem;">
                        <strong>Your answer:</strong> {u}: {u_txt}{corr_line}</div>
                </div>""", unsafe_allow_html=True)
                with st.expander(f"📖 Explanation — Q{a['question_num']}", expanded=False):
                    p = qd.get("personal_note", "")
                    pnote = f'<br><strong>📋 From your notes:</strong> {p}' if p else ""
                    st.markdown(f"""<div class="quiz-explanation">
                    {qd.get("explanation","")}<br><br>
                    <strong>📌 Citation:</strong> {qd.get("rule_citation","")}{pnote}
                    </div>""", unsafe_allow_html=True)

            if st.session_state.quiz_log:
                st.markdown("---")
                ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                st.download_button(
                    f"⬇️ Download All Quiz Results ({len(st.session_state.quiz_log)} saved)",
                    data=json.dumps(st.session_state.quiz_log, indent=2, ensure_ascii=False),
                    file_name=f"picklerick_quiz_{ts}.json",
                    mime="application/json",
                    use_container_width=True,
                )

    elif st.session_state.quiz_mode is None:
        st.markdown(f"""<div class="pr-card" style="text-align:center;padding:1.5rem;">
        <p style="color:{MUTED};margin:0;">👆 Pick Single Question Mode or 10-Question mode above.</p>
        </div>""", unsafe_allow_html=True)


# =============================================================================
# FOOTER
# =============================================================================

st.markdown(f"""
<div class="pr-footer">
    <strong style="color:{GREEN_DARK};">Pickleball rules AI assistant trained on USAPA rulebook
    and all recent updates.</strong><br>
    Built for picklers, by a pickler 🥒 &nbsp;|&nbsp;
    Pickle Rick v1.0 &nbsp;|&nbsp;
    2026 USAPA Official Rulebook + 2026 Change Document<br>
    <span style="font-size:0.72rem;">
    Not official USAPA interpretation — confirm with a certified referee or tournament director if needed.
    </span>
</div>
""", unsafe_allow_html=True)
